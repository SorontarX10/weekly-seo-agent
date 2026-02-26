from __future__ import annotations

from datetime import date
import re

from docx import Document

from weekly_seo_agent.models import AnalysisResult, DateWindow, ExternalSignal, KeyDelta, MetricSummary
from weekly_seo_agent.reporting import (
    _build_reasoning_hypotheses,
    _marketplace_timeline_rows,
    build_markdown_report,
    enforce_manager_quality_guardrail,
    write_docx,
)


def _delta_row(key: str, current: float, previous: float, yoy: float) -> KeyDelta:
    delta_prev = current - previous
    delta_yoy = current - yoy
    pct_prev = (delta_prev / previous) if previous else 0.0
    pct_yoy = (delta_yoy / yoy) if yoy else 0.0
    return KeyDelta(
        key=key,
        current_clicks=current,
        previous_clicks=previous,
        yoy_clicks=yoy,
        click_delta_vs_previous=delta_prev,
        click_delta_pct_vs_previous=pct_prev,
        click_delta_vs_yoy=delta_yoy,
        click_delta_pct_vs_yoy=pct_yoy,
        current_ctr=0.1,
        previous_ctr=0.1,
        yoy_ctr=0.1,
        current_position=2.0,
        previous_position=2.2,
        yoy_position=2.4,
    )


def test_report_contains_reasoning_and_upcoming_trends(tmp_path):
    run_date = date(2026, 2, 10)
    windows = {
        "current_28d": DateWindow("Current 28 days", date(2026, 1, 13), date(2026, 2, 9)),
        "previous_28d": DateWindow("Previous 28 days", date(2025, 12, 16), date(2026, 1, 12)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 1, 14), date(2025, 2, 10)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=100_000, impressions=1_000_000, ctr=0.1, position=3.2),
        "previous_28d": MetricSummary(clicks=95_000, impressions=1_020_000, ctr=0.093, position=3.4),
        "yoy_52w": MetricSummary(clicks=110_000, impressions=1_050_000, ctr=0.104, position=3.0),
    }

    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[_delta_row("fajerwerki", 200, 6000, 5000)],
        top_winners=[_delta_row("licytacje wosp", 8000, 300, 1200)],
        findings=[],
    )

    external_signals = [
        ExternalSignal(
            source="Google Search Status",
            day=date(2026, 2, 5),
            title="February 2026 Discover update",
            details="Core update for discover.",
            severity="medium",
        ),
        ExternalSignal(
            source="Public Holidays",
            day=date(2026, 2, 14),
            title="Holiday: Test holiday",
            details="Test",
            severity="info",
        ),
        ExternalSignal(
            source="Campaign tracker Allegro (news.google.com)",
            day=date(2026, 2, 8),
            title="Allegro Days campaign is live",
            details="Allegro starts Allegro Days with new promotion mechanics.",
            severity="medium",
        ),
        ExternalSignal(
            source="Campaign tracker Competitors (news.google.com)",
            day=date(2026, 2, 7),
            title="Amazon Black Friday sale plan leaked",
            details="Competitor campaign event mention.",
            severity="high",
        ),
    ]

    ferie_context = {
        "source": "OpenHolidays API",
        "source_url": "https://openholidaysapi.org",
        "profiles_ranked": [
            {
                "code": "PL-MZ",
                "name": "Mazowieckie",
                "population_m": 5.52,
                "avg_salary_pln": 9860.0,
                "gmv_share": 0.20,
            }
        ],
        "window_stats": {
            "current_28d": {
                "days_with_ferie": 21,
                "avg_daily_gmv_share": 0.39,
                "peak_daily_gmv_share": 0.52,
            },
            "previous_28d": {
                "days_with_ferie": 0,
                "avg_daily_gmv_share": 0.00,
                "peak_daily_gmv_share": 0.00,
            },
            "yoy_52w": {
                "days_with_ferie": 21,
                "avg_daily_gmv_share": 0.35,
                "peak_daily_gmv_share": 0.49,
                "voivodeship_days": {"PL-MZ": 14},
            },
        },
        "missing_years": [],
        "yoy_comparison": {
            "avg_daily_delta_pp": 4.0,
            "rows": [
                {
                    "code": "PL-MZ",
                    "name": "Mazowieckie",
                    "current_days": 14,
                    "yoy_days": 7,
                    "delta_days": 7,
                    "contribution_pp": 1.4,
                }
            ],
        },
        "errors": [],
    }
    ferie_trends = [
        (
            date(2026, 2, 16),
            "2026-02-16 - 2026-03-01",
            "Winter break: Lubelskie, Lodzkie",
            "Approx. 20.0% of GMV proxy is in regions currently on winter break.",
            "Account for regional seasonality.",
        )
    ]
    additional_context = {
        "ga4": {
            "enabled": True,
            "property_id": "123456789",
            "country_code": "PL",
            "summary": {
                "current": {
                    "sessions": 150000.0,
                    "users": 120000.0,
                    "engaged_sessions": 100000.0,
                    "transactions": 5000.0,
                    "revenue": 2500000.0,
                },
                "previous": {
                    "sessions": 140000.0,
                    "users": 115000.0,
                    "engaged_sessions": 95000.0,
                    "transactions": 4800.0,
                    "revenue": 2400000.0,
                },
                "yoy": {
                    "sessions": 130000.0,
                    "users": 110000.0,
                    "engaged_sessions": 90000.0,
                    "transactions": 4500.0,
                    "revenue": 2200000.0,
                },
            },
            "top_landing_pages": [],
            "errors": [],
        },
        "allegro_trends": {
            "enabled": True,
            "country_code": "PL",
            "source": "Allegro Trends API",
            "from": "2026-02-03",
            "till": "2026-02-09",
            "interval": "day",
            "rows": [
                {
                    "query": "pellet",
                    "visit": 12345.0,
                    "pv": 16789.0,
                    "offers": 560.0,
                    "deals": 42.0,
                    "gmv": 250000.0,
                    "points": 7,
                    "http_code": 200,
                }
            ],
            "top_rows": 10,
            "errors": [],
        },
        "market_event_calendar": {
            "enabled": True,
            "source": "GDELT DOC API",
            "country_code": "PL",
            "top_rows": 5,
            "events": [
                {
                    "date": "2026-02-09",
                    "event_type": "Campaign/Promotions",
                    "impact_level": "HIGH",
                    "impact_direction": "Upside potential",
                    "confidence": 82,
                    "source": "example.com",
                    "title": "Allegro Days campaign starts",
                    "gmv_reason": "Likely demand uplift from campaign traffic.",
                }
            ],
            "errors": [],
        },
        "macro": {
            "nbp_fx": {
                "eur_pln": {
                    "avg_current": 4.20,
                    "avg_previous": 4.18,
                    "delta_pct_vs_previous": 0.48,
                    "latest": 4.21,
                },
                "usd_pln": {
                    "avg_current": 3.60,
                    "avg_previous": 3.62,
                    "delta_pct_vs_previous": -0.55,
                    "latest": 3.58,
                },
            },
            "imgw_warnings_total": 3,
            "imgw_high_severity_count": 2,
            "imgw_warnings": [
                {
                    "severity": "2",
                    "event": "Heavy snowfall",
                    "areas": "Mazowieckie",
                    "to": "2026-02-11",
                }
            ],
        },
        "product_trends": {
            "enabled": True,
            "source": "Google Sheets product trend trackers",
            "top_rows": 5,
            "horizon_days": 31,
            "top_yoy_non_brand": [
                {
                    "trend": "prezent walentynki",
                    "current_value": 1800.0,
                    "previous_value": 900.0,
                    "delta_value": 900.0,
                    "delta_pct": 100.0,
                    "sheet": "Current",
                }
            ],
            "upcoming_31d": [
                {
                    "date": "2026-02-20",
                    "trend": "prezent walentynki",
                    "value": 1300.0,
                    "sheet": "Upcoming",
                }
            ],
            "current_non_brand": [
                {
                    "date": "2026-02-10",
                    "trend": "prezent walentynki",
                    "value": 1500.0,
                    "sheet": "Current",
                }
            ],
            "errors": [],
        }
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=external_signals,
        weather_summary={
            "avg_temp_current_c": -5.0,
            "avg_temp_previous_c": -1.0,
            "avg_temp_diff_c": -4.0,
            "precip_current_mm": 10.0,
            "precip_previous_mm": 12.0,
            "precip_change_pct": -16.0,
        },
        ferie_context=ferie_context,
        upcoming_ferie_trends=ferie_trends,
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error="Auth error",
        query_filter_stats={"query": {"current": 1, "previous": 1, "yoy": 1}},
    )

    assert "## Integrated reasoning" not in report
    assert "## Root cause matrix" not in report
    assert "## Executive summary" in report
    assert "`YoY` (GA4):" not in report
    assert "`Data policy`: GA4 deltas are temporarily excluded from decision narrative" not in report
    assert "## What is happening and why" in report
    assert "Marketplace timeline (market events + promo calendar)" in report
    assert "Weekly market storyline on one timeline:" in report
    assert "Reasoning ledger (facts -> hypotheses -> validation)" in report
    assert "Seasonality decomposition (plain language):" in report
    assert "Plain-language trend read:" in report
    assert "Confirmed facts vs plausible drivers vs open questions" in report
    assert "Open questions for next run:" in report
    assert "## Appendix" not in report
    assert "## Date windows" not in report
    assert "### Market event calendar (API)" not in report
    assert "### Campaign events (Allegro vs competitors)" not in report
    assert "### Product trends (non-brand, Sheets)" not in report

    out_path = tmp_path / "report.docx"
    write_docx(out_path, "test", report)
    assert out_path.exists()
    assert out_path.stat().st_size > 0


def test_write_docx_renders_markdown_bold(tmp_path):
    out_path = tmp_path / "report_bold.docx"
    markdown = "1. **Wazny punkt** i doprecyzowanie"
    write_docx(out_path, "test-bold", markdown)

    doc = Document(out_path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert paragraphs

    first = paragraphs[0]
    assert first.text == "Wazny punkt i doprecyzowanie"
    assert any(run.text == "Wazny punkt" and run.bold for run in first.runs)


def test_marketplace_timeline_canonical_dedup_merges_same_event_same_day() -> None:
    additional_context = {
        "market_event_calendar": {
            "enabled": True,
            "country_code": "PL",
            "events": [
                {
                    "date": "2026-02-22",
                    "title": "Planned campaign: MegaRaty February",
                    "source": "Trade plan sheet",
                    "impact_level": "MEDIUM",
                    "impact_direction": "Upside potential",
                    "confidence": 62,
                },
                {
                    "date": "2026-02-22",
                    "title": "MegaRaty February - external mention",
                    "source": "Market feed",
                    "impact_level": "HIGH",
                    "impact_direction": "Upside potential",
                    "confidence": 78,
                },
            ],
        }
    }
    campaign_context = {
        "allegro_events": [
            ExternalSignal(
                source="Campaign tracker Allegro",
                day=date(2026, 2, 22),
                title="MegaRaty February campaign live",
                details="Campaign mention",
                severity="medium",
            )
        ],
        "competitor_events": [],
    }

    rows = _marketplace_timeline_rows(
        additional_context=additional_context,
        campaign_context=campaign_context,
        report_country_code="PL",
        max_rows=20,
    )

    assert len(rows) == 1
    first = rows[0]
    assert first["country"] == "PL"
    assert "MegaRaty" in str(first.get("event", ""))
    assert "HIGH" in str(first.get("impact", "")).upper()


def test_write_docx_supports_level4_markdown_heading(tmp_path):
    out_path = tmp_path / "report_h4.docx"
    markdown = "#### Seasonality and Winter Queries"
    write_docx(out_path, "test-h4", markdown)

    doc = Document(out_path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    assert paragraphs
    assert paragraphs[0].text == "Seasonality and Winter Queries"
    assert "#" not in paragraphs[0].text


def test_report_hides_missing_ga4_metrics_and_shows_channel_split():
    run_date = date(2026, 2, 11)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 2), date(2026, 2, 8)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 1, 26), date(2026, 2, 1)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 3), date(2025, 2, 9)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=100, impressions=1000, ctr=0.1, position=3.0),
        "previous_28d": MetricSummary(clicks=90, impressions=900, ctr=0.1, position=3.1),
        "yoy_52w": MetricSummary(clicks=110, impressions=1100, ctr=0.1, position=2.9),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "ga4": {
            "enabled": True,
            "property_id": "151393439",
            "country_code": "PL",
            "top_rows": 5,
            "metric_availability": {
                "sessions": True,
                "users": True,
                "engaged_sessions": True,
                "transactions": False,
                "revenue": True,
            },
            "summary": {
                "current": {
                    "sessions": 1000.0,
                    "users": 850.0,
                    "engaged_sessions": 700.0,
                    "transactions": None,
                    "revenue": 50000.0,
                },
                "previous": {
                    "sessions": 950.0,
                    "users": 820.0,
                    "engaged_sessions": 680.0,
                    "transactions": None,
                    "revenue": 48000.0,
                },
                "yoy": {
                    "sessions": 1100.0,
                    "users": 900.0,
                    "engaged_sessions": 730.0,
                    "transactions": None,
                    "revenue": 45000.0,
                },
            },
            "channels": {
                "yoy_deltas": [
                    {
                        "channel": "Paid Search",
                        "sessions_current": 400.0,
                        "sessions_yoy": 250.0,
                        "delta_vs_yoy": 150.0,
                        "delta_vs_yoy_pct": 60.0,
                    },
                    {
                        "channel": "Organic Search",
                        "sessions_current": 350.0,
                        "sessions_yoy": 500.0,
                        "delta_vs_yoy": -150.0,
                        "delta_vs_yoy_pct": -30.0,
                    },
                ]
            },
            "cannibalization": {
                "potential": True,
                "note": "Organic Search down while paid channels up YoY; validate possible channel cannibalization.",
            },
            "top_landing_pages": [
                {
                    "landing_page": "https://allegro.pl/",
                    "sessions": 500.0,
                    "transactions": None,
                    "revenue": 30000.0,
                }
            ],
            "errors": ["GA4 bug: missing metrics in API response: transactions"],
        }
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "| Transactions |" not in report
    assert "Potential channel cannibalization (GA4): Organic Search down while paid channels up YoY" not in report
    assert "| Channel | Sessions current | Sessions YoY | Delta vs YoY | Delta % vs YoY |" not in report
    assert "Top channel growth YoY:" not in report
    assert "Paid Search (+150)" not in report


def test_report_includes_daily_gsc_storyline_with_context_links():
    run_date = date(2026, 2, 24)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=1_420_000, impressions=30_000_000, ctr=0.047, position=6.3),
        "previous_28d": MetricSummary(clicks=1_390_000, impressions=29_500_000, ctr=0.047, position=6.2),
        "yoy_52w": MetricSummary(clicks=1_590_000, impressions=31_500_000, ctr=0.050, position=4.9),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    external_signals = [
        ExternalSignal(
            source="Campaign tracker Allegro (news.google.com)",
            day=date(2026, 2, 22),
            title="Planned campaign: MegaRaty February",
            details="Campaign timing can influence demand allocation.",
            severity="medium",
        ),
        ExternalSignal(
            source="Weekly SEO digest (seo)",
            day=date(2026, 2, 21),
            title="Google Search Ranking Volatility Beginning To Cool?",
            details="Volatility context for this week.",
            severity="info",
        ),
    ]
    additional_context = {
        "country_code": "PL",
        "gsc_daily_rows": {
            "enabled": True,
            "days_total": 7,
            "days_with_data": 7,
            "days_with_previous_weekday_data": 7,
            "weekly_clicks_sum": 1_420_000.0,
            "rows": [
                {
                    "date": "2026-02-20",
                    "clicks": 220000.0,
                    "impressions": 4100000.0,
                    "ctr": 0.0537,
                    "position": 6.1,
                    "previous_weekday_has_data": True,
                    "delta_clicks_vs_previous_weekday": -26000.0,
                    "delta_pct_vs_previous_weekday": -10.6,
                    "yoy_day_has_data": True,
                    "delta_pct_vs_yoy_day": -15.1,
                },
                {
                    "date": "2026-02-22",
                    "clicks": 245000.0,
                    "impressions": 4350000.0,
                    "ctr": 0.0563,
                    "position": 6.0,
                    "previous_weekday_has_data": True,
                    "delta_clicks_vs_previous_weekday": 31000.0,
                    "delta_pct_vs_previous_weekday": 14.5,
                    "yoy_day_has_data": True,
                    "delta_pct_vs_yoy_day": -9.2,
                },
            ],
        },
        "trade_plan": {
            "enabled": True,
            "campaign_rows": [
                {
                    "campaign": "MegaRaty February",
                    "first_date": "2026-02-20",
                    "last_date": "2026-02-24",
                }
            ],
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=external_signals,
        weather_summary={
            "daily_current": [
                {"date": "2026-02-20", "temp_c": -1.0, "precip_mm": 4.5},
                {"date": "2026-02-22", "temp_c": 2.0, "precip_mm": 1.0},
            ],
            "daily_previous": [
                {"date": "2026-02-13", "temp_c": 2.5, "precip_mm": 0.5},
                {"date": "2026-02-15", "temp_c": -1.5, "precip_mm": 7.0},
            ],
        },
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Daily GSC pulse (day-by-day)" in report
    assert "Daily trend view (GSC by day)" in report
    assert "trade-plan campaign active: MegaRaty February" in report


def test_daily_weather_hint_is_directional_non_causal_and_confidence_tagged():
    run_date = date(2026, 2, 24)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=1_200_000, impressions=25_000_000, ctr=0.048, position=6.1),
        "previous_28d": MetricSummary(clicks=1_180_000, impressions=24_800_000, ctr=0.047, position=6.2),
        "yoy_52w": MetricSummary(clicks=1_300_000, impressions=26_000_000, ctr=0.050, position=5.7),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "gsc_daily_rows": {
            "enabled": True,
            "days_total": 7,
            "days_with_data": 7,
            "days_with_previous_weekday_data": 7,
            "weekly_clicks_sum": 1_200_000.0,
            "rows": [
                {
                    "date": "2026-02-20",
                    "clicks": 185000.0,
                    "impressions": 3700000.0,
                    "ctr": 0.05,
                    "position": 6.0,
                    "previous_weekday_has_data": True,
                    "delta_clicks_vs_previous_weekday": -22000.0,
                    "delta_pct_vs_previous_weekday": -10.2,
                    "yoy_day_has_data": False,
                    "delta_pct_vs_yoy_day": 0.0,
                },
            ],
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={
            "daily_current": [
                {"date": "2026-02-20", "temp_c": -3.0, "precip_mm": 8.0},
            ],
            "daily_previous": [
                {"date": "2026-02-13", "temp_c": 2.0, "precip_mm": 1.0},
            ],
        },
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "weather timing hint:" in report
    assert "non-causal" in report
    assert "confidence:" in report
    assert "/100" in report


def test_report_includes_trend_yoy_seasonality_decomposition_in_plain_language():
    run_date = date(2026, 2, 24)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=1_420_000, impressions=30_000_000, ctr=0.047, position=6.3),
        "previous_28d": MetricSummary(clicks=1_390_000, impressions=29_500_000, ctr=0.047, position=6.2),
        "yoy_52w": MetricSummary(clicks=1_590_000, impressions=31_500_000, ctr=0.05, position=4.9),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "product_trends": {
            "enabled": True,
            "top_rows": 8,
            "horizon_days": 31,
            "top_yoy_non_brand": [
                {"trend": "prezent walentynki", "current_value": 1900.0, "previous_value": 1000.0, "delta_value": 900.0},
                {"trend": "black friday promocja tv", "current_value": 300.0, "previous_value": 1000.0, "delta_value": -700.0},
                {"trend": "buty do biegania", "current_value": 800.0, "previous_value": 500.0, "delta_value": 300.0},
            ],
            "upcoming_31d": [
                {"date": "2026-03-01", "trend": "wielkanoc dekoracje", "value": 1200.0},
                {"date": "2026-03-02", "trend": "megaraty smart week", "value": 900.0},
            ],
            "current_non_brand": [
                {"date": "2026-02-24", "trend": "wiosenne ogrodnictwo", "value": 1400.0},
            ],
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Seasonality decomposition (plain language):" in report
    assert "Seasonal/calendar-driven" in report
    assert "Campaign/event-driven" in report
    assert "Evergreen/base-demand" in report
    assert "Plain-language trend read:" in report


def test_quality_guardrail_enforces_markers_and_compacts_optional_lines():
    long_optional = "\n".join(
        [
            "- Macro context: extended optional explanation line for manager readability control."
            for _ in range(60)
        ]
    )
    raw = (
        "# Weekly SEO Intelligence Report (2026-02-25 | PL)\n\n"
        "## Executive summary\n"
        "- **What changed**: sample text.\n\n"
        "## What is happening and why\n"
        "- **WoW diagnosis**: sample.\n"
        f"{long_optional}\n\n"
        "## Hypothesis protocol\n"
        "| Hypothesis | Confidence | What would disprove it | Validation metric | Check date |\n"
        "|---|---|---|---|---|\n"
        "| Sample | Medium | - | - | - |\n\n"
        "## Governance and provenance\n"
        "- Sample.\n\n"
        "## Evidence ledger\n"
        "| ID | Source | Date | Note |\n"
        "|---|---|---|---|\n"
        "| E1 | Test | 2026-02-25 | note |\n"
    )

    guarded = enforce_manager_quality_guardrail(raw, max_words=160)
    lowered = guarded.lower()
    assert "falsifier" in lowered
    assert "validation metric" in lowered
    assert "validation date" in lowered
    assert len(re.findall(r"\b\w+\b", guarded, flags=re.UNICODE)) <= 160


def test_report_surfaces_serp_appearance_and_13m_context_blocks():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=14_200_000, impressions=300_000_000, ctr=0.0472, position=6.27),
        "previous_28d": MetricSummary(clicks=14_000_000, impressions=297_000_000, ctr=0.0470, position=6.29),
        "yoy_52w": MetricSummary(clicks=15_900_000, impressions=315_000_000, ctr=0.0500, position=4.92),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "gsc_feature_split": {
            "enabled": True,
            "rows_weekly": [
                {"feature": "MERCHANT_LISTINGS", "delta_clicks_vs_previous": 55800.0},
                {"feature": "PRODUCT_SNIPPETS", "delta_clicks_vs_previous": -28750.0},
            ],
            "rows_mom": [
                {"feature": "MERCHANT_LISTINGS", "delta_clicks_vs_previous": 121000.0},
                {"feature": "REVIEW_SNIPPET", "delta_clicks_vs_previous": -8300.0},
            ],
            "feature_overview": [
                {"feature": "MERCHANT_LISTINGS", "yoy_delta_clicks": 210000.0},
                {"feature": "PRODUCT_SNIPPETS", "yoy_delta_clicks": -165000.0},
            ],
        },
        "google_updates_timeline": {
            "enabled": True,
            "scan_start": "2025-01-25",
            "scan_end": "2026-02-25",
            "summary": {
                "count_current_30d": 6,
                "count_previous_30d": 4,
                "count_yoy_30d": 5,
                "total_count_13m": 72,
                "latest_update_date": "2026-02-21",
                "latest_update_title": "Search ranking volatility update",
            },
        },
        "serp_case_studies": {
            "enabled": True,
            "scan_start": "2025-01-25",
            "scan_end": "2026-02-25",
            "summary": {
                "topic_counts_13m": {
                    "CTR and click distribution": 18,
                    "SERP layout and features": 13,
                    "Merchant and free listings": 9,
                },
                "count_current_30d": 5,
                "count_previous_30d": 3,
                "total_count_13m": 44,
                "latest_case_date": "2026-02-20",
                "latest_case_title": "Merchant listings CTR shift after SERP change",
            },
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "SERP appearance mix (WoW/MoM/YoY)" in report
    assert "Google updates timeline (13 months)" in report
    assert "External case-study context (13 months)" in report
    assert "SERP appearance deltas (GSC searchAppearance)" in report
    assert "Google update timeline (13M):" in report
    assert "SERP case-study scanner (13M):" in report


def test_report_surfaces_daily_serp_shifts_and_source_reliability():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=14_200_000, impressions=300_000_000, ctr=0.0472, position=6.27),
        "previous_28d": MetricSummary(clicks=14_000_000, impressions=297_000_000, ctr=0.0470, position=6.29),
        "yoy_52w": MetricSummary(clicks=15_900_000, impressions=315_000_000, ctr=0.0500, position=4.92),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "gsc_feature_split": {
            "enabled": True,
            "rows_weekly": [
                {"feature": "MERCHANT_LISTINGS", "delta_clicks_vs_previous": 55800.0},
                {"feature": "PRODUCT_SNIPPETS", "delta_clicks_vs_previous": -28750.0},
            ],
            "rows_mom": [
                {"feature": "MERCHANT_LISTINGS", "delta_clicks_vs_previous": 121000.0},
                {"feature": "REVIEW_SNIPPET", "delta_clicks_vs_previous": -8300.0},
            ],
            "rows_unified": [
                {
                    "feature": "MERCHANT_LISTINGS",
                    "wow_delta_clicks": 55800.0,
                    "mom_delta_clicks": 121000.0,
                    "yoy_delta_clicks": 210000.0,
                },
                {
                    "feature": "PRODUCT_SNIPPETS",
                    "wow_delta_clicks": -28750.0,
                    "mom_delta_clicks": -10300.0,
                    "yoy_delta_clicks": -165000.0,
                },
            ],
        },
        "daily_serp_feature_shifts": {
            "enabled": True,
            "anomalies": [
                {
                    "date": "2026-02-20",
                    "feature": "MERCHANT_LISTINGS",
                    "delta_share_pp_vs_previous_weekday": 2.1,
                }
            ],
        },
        "daily_gsc_anomalies": {
            "enabled": True,
            "rows": [
                {
                    "date": "2026-02-20",
                    "delta_clicks_vs_previous_weekday": -26200.0,
                    "delta_ctr_pp_vs_previous_weekday": -0.31,
                    "markers": ["campaign/news marker: MegaRaty February"],
                }
            ],
        },
        "external_source_reliability": {
            "enabled": True,
            "summary": {
                "weighted_score": 83.4,
                "sources_count": 9,
                "tier_counts": {"official": 3, "public-institution": 1, "high-quality media": 4},
            },
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Daily SERP feature shifts" in report
    assert "Daily KPI anomalies" in report
    assert "SERP feature split table (aligned windows):" in report
    assert "Source reliability score" in report


def test_external_source_hypotheses_are_supporting_only():
    totals = {
        "current_28d": MetricSummary(clicks=1000, impressions=10000, ctr=0.1, position=3.2),
        "previous_28d": MetricSummary(clicks=1000, impressions=10000, ctr=0.1, position=3.2),
        "yoy_52w": MetricSummary(clicks=1000, impressions=10000, ctr=0.1, position=3.2),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "google_updates_timeline": {
            "enabled": True,
            "summary": {
                "count_current_30d": 5,
                "count_previous_30d": 1,
                "count_yoy_30d": 2,
                "latest_update_date": "2026-02-20",
                "latest_update_title": "Search volatility update",
            },
        },
        "serp_case_studies": {
            "enabled": True,
            "summary": {
                "topic_counts_13m": {"CTR and click distribution": 7},
                "count_current_30d": 3,
                "count_previous_30d": 1,
                "total_count_13m": 18,
                "latest_case_date": "2026-02-18",
                "latest_case_title": "SERP layout and CTR",
            },
        },
    }
    ferie_context = {"yoy_comparison": {"avg_daily_delta_pp": 0.0, "rows": []}}

    hypotheses = _build_reasoning_hypotheses(
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        ferie_context=ferie_context,
        segment_diagnostics=None,
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    external_rows = [
        row
        for row in hypotheses
        if isinstance(row, dict)
        and str(row.get("category", "")).strip().lower() in {"algorithm/serp context", "serp behavior context"}
    ]
    assert external_rows
    for row in external_rows:
        assert bool(row.get("supporting_context_only"))
        assert int(row.get("confidence", 0)) <= 63


def test_report_includes_trade_plan_overlap_intensity_and_yoy_fallback_message():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=120000, impressions=2100000, ctr=0.057, position=5.8),
        "previous_28d": MetricSummary(clicks=115000, impressions=2050000, ctr=0.056, position=5.9),
        "yoy_52w": MetricSummary(clicks=130000, impressions=2200000, ctr=0.059, position=5.5),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "trade_plan": {
            "enabled": True,
            "windows": {
                "current": {"start": "2026-02-16", "end": "2026-02-22"},
                "previous": {"start": "2026-02-09", "end": "2026-02-15"},
            },
            "channel_split": [
                {"channel": "Paid Search", "current_spend": 120.0, "previous_spend": 90.0, "delta_spend": 30.0},
            ],
            "campaign_rows": [
                {
                    "campaign": "MegaRaty February",
                    "category": "Promo",
                    "first_date": "2026-02-19",
                    "last_date": "2026-02-23",
                    "current_spend": 80.0,
                },
                {
                    "campaign": "Warehouse Cleaning Sellout",
                    "category": "Promo",
                    "first_date": "2026-02-17",
                    "last_date": "2026-02-20",
                    "current_spend": 40.0,
                },
            ],
            "yoy_availability": {
                "requested": True,
                "applied": False,
                "fallback_used": False,
                "message": "YoY comparator skipped: no prior-year trade-plan sheet/tab configured.",
            },
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Overlap intensity:" in report or "Trade-plan overlap intensity:" in report
    assert "Trade-plan YoY availability" in report
    assert "YoY comparator skipped" in report


def test_report_surfaces_marketplace_event_yoy_density_context():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=150000, impressions=2500000, ctr=0.06, position=5.2),
        "previous_28d": MetricSummary(clicks=149000, impressions=2490000, ctr=0.06, position=5.2),
        "yoy_52w": MetricSummary(clicks=158000, impressions=2590000, ctr=0.061, position=5.1),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "market_event_calendar": {
            "enabled": True,
            "counts": {
                "current": 8,
                "yoy": 5,
                "delta_vs_yoy": 3,
                "delta_pct_vs_yoy": 60.0,
            },
            "events": [
                {
                    "date": "2026-02-20",
                    "country_code": "PL",
                    "title": "Allegro campaign event",
                    "source": "example.com",
                    "event_type": "Campaign/Promotions",
                    "impact_level": "MEDIUM",
                    "impact_direction": "Upside potential",
                    "confidence": 60,
                    "gmv_reason": "Promo demand shift.",
                }
            ],
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Marketplace events YoY context:" in report
    assert "8 events now vs 5 YoY" in report


def test_report_includes_reasoning_engine_sections_for_f_block():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=14_200_000, impressions=300_000_000, ctr=0.0472, position=6.27),
        "previous_28d": MetricSummary(clicks=14_000_000, impressions=297_000_000, ctr=0.0470, position=6.29),
        "yoy_52w": MetricSummary(clicks=15_900_000, impressions=315_000_000, ctr=0.0500, position=4.92),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    external_signals = [
        ExternalSignal(
            source="Google Search Status",
            day=date(2026, 2, 20),
            title="Search volatility update",
            details="Update note",
            severity="medium",
        ),
    ]
    additional_context = {
        "country_code": "PL",
        "trade_plan": {
            "enabled": True,
            "campaign_rows": [
                {
                    "campaign": "MegaRaty February",
                    "first_date": "2026-02-20",
                    "last_date": "2026-02-24",
                    "current_spend": 120.0,
                }
            ],
        },
        "product_trends": {
            "enabled": True,
            "top_rows": 5,
            "horizon_days": 31,
            "top_yoy_non_brand": [
                {"trend": "wielkanoc dekoracje", "current_value": 1200.0, "previous_value": 600.0, "delta_value": 600.0},
                {"trend": "smart week promocja", "current_value": 900.0, "previous_value": 1400.0, "delta_value": -500.0},
            ],
        },
        "market_event_calendar": {
            "enabled": True,
            "events": [
                {
                    "date": "2026-02-21",
                    "title": "Campaign event signal",
                    "source": "example.com",
                    "impact_level": "MEDIUM",
                    "impact_direction": "Upside potential",
                    "confidence": 60,
                }
            ],
        },
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=external_signals,
        weather_summary={"avg_temp_diff_c": -2.6, "precip_change_pct": 112.0},
        additional_context=additional_context,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "## Evidence coverage check" in report
    assert "## Validation plan (next week)" in report
    assert "## Counterfactual checks" in report
    assert "## Causality guardrail" in report
    assert "## Escalation rule" in report
    assert "| Hypothesis | Priority | Confidence | Falsifier | Validation metric | Validation date |" in report
    assert "Technical SEO escalation gate:" in report


def test_contradiction_check_adds_reconciliation_sentence():
    run_date = date(2026, 2, 25)
    windows = {
        "current_28d": DateWindow("Current week (Mon-Sun)", date(2026, 2, 16), date(2026, 2, 22)),
        "previous_28d": DateWindow("Previous week (Mon-Sun)", date(2026, 2, 9), date(2026, 2, 15)),
        "yoy_52w": DateWindow("YoY aligned (52 weeks ago)", date(2025, 2, 17), date(2025, 2, 23)),
    }
    totals = {
        "current_28d": MetricSummary(clicks=14_200_000, impressions=300_000_000, ctr=0.0472, position=6.27),
        "previous_28d": MetricSummary(clicks=14_000_000, impressions=297_000_000, ctr=0.0470, position=6.29),
        "yoy_52w": MetricSummary(clicks=15_900_000, impressions=315_000_000, ctr=0.0500, position=4.92),
    }
    query_analysis = AnalysisResult(
        summary_current=totals["current_28d"],
        summary_previous=totals["previous_28d"],
        summary_yoy=totals["yoy_52w"],
        top_losers=[],
        top_winners=[],
        findings=[],
    )
    additional_context = {
        "country_code": "PL",
        "google_trends_brand": {
            "enabled": True,
            "summary": {"delta_pct_vs_previous": 1.2, "delta_pct_vs_yoy": -0.3},
        },
    }
    segment_diagnostics = {
        "page_template": [
            {
                "segment": "home",
                "current_clicks": 2_460_000.0,
                "previous_clicks": 2_520_000.0,
                "yoy_clicks": 2_500_000.0,
                "delta_vs_previous": -60_000.0,
                "delta_pct_vs_previous": -0.0238,
                "delta_vs_yoy": -40_000.0,
                "delta_pct_vs_yoy": -0.016,
            }
        ],
        "brand_non_brand": [
            {
                "segment": "brand",
                "current_clicks": 2_600_000.0,
                "previous_clicks": 2_700_000.0,
                "yoy_clicks": 2_650_000.0,
                "delta_vs_previous": -100_000.0,
                "delta_pct_vs_previous": -0.037,
                "delta_vs_yoy": -50_000.0,
                "delta_pct_vs_yoy": -0.019,
                "current_impressions": 3_300_000.0,
                "previous_impressions": 3_350_000.0,
                "yoy_impressions": 3_360_000.0,
                "impressions_delta_vs_previous": -50_000.0,
                "impressions_delta_pct_vs_previous": -0.015,
                "impressions_delta_vs_yoy": -60_000.0,
                "impressions_delta_pct_vs_yoy": -0.018,
                "current_ctr": 0.7878,
                "previous_ctr": 0.8060,
                "yoy_ctr": 0.7890,
            }
        ],
    }

    report = build_markdown_report(
        run_date=run_date,
        report_country_code="PL",
        windows=windows,
        totals=totals,
        scope_results=[("query", query_analysis)],
        external_signals=[],
        weather_summary={},
        additional_context=additional_context,
        segment_diagnostics=segment_diagnostics,
        senuto_summary=None,
        senuto_error=None,
    )

    assert "Contradiction check:" in report
    assert "Reconciliation:" in report
