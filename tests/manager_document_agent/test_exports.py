from __future__ import annotations

import json
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

import weekly_seo_agent.manager_document_agent.api as api_module
from weekly_seo_agent.manager_document_agent.api import create_app
from weekly_seo_agent.manager_document_agent.exporting import build_drive_export_name, export_markdown_like_to_docx


def _seed_oauth_drive_integration(
    client: TestClient,
    *,
    folder_name: str = "Manager Documents",
    folder_id: str = "",
    token_json: str = '{"token":"fake"}',
) -> None:
    client.app.state.document_service.upsert_integration_setting(
        provider="google_drive",
        config_json=json.dumps(
            {
                "credential_mode": "oauth_quick_connect",
                "client_id": "client-id",
                "client_secret": "client-secret",
                "token_json": token_json,
                "folder_name": folder_name,
                "folder_id": folder_id,
            }
        ),
    )


def _create_payload(title: str, content: str = "# Title\n\n- point") -> dict:
    return {
        "title": title,
        "doc_type": "MANAGEMENT_BRIEF",
        "target_audience": "Management",
        "language": "pl",
        "objective": "Summarize progress",
        "tone": "formal",
        "constraints": "Keep it concise",
        "current_content": content,
    }


def test_docx_export_and_export_history(tmp_path):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    with TestClient(app) as client:
        created = client.post("/documents", json=_create_payload("Export me"))
        document_id = created.json()["id"]

        exported = client.post(f"/documents/{document_id}/export/docx")
        assert exported.status_code == 200
        body = exported.json()
        assert body["status"] == "SUCCESS"
        assert body["export_type"] == "DOCX"
        assert Path(body["file_path"]).exists()

        history = client.get(f"/documents/{document_id}/exports")
        assert history.status_code == 200
        records = history.json()
        assert len(records) == 1
        assert records[0]["export_type"] == "DOCX"


def test_google_drive_export_requires_settings_and_records_error(tmp_path):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    with TestClient(app) as client:
        created = client.post("/documents", json=_create_payload("Drive me"))
        document_id = created.json()["id"]

        drive = client.post(f"/documents/{document_id}/export/drive")
        assert drive.status_code == 400
        assert "Configure /integrations/google-drive first" in drive.json()["detail"]

        history = client.get(f"/documents/{document_id}/exports")
        records = history.json()
        assert len(records) == 1
        assert records[0]["export_type"] == "GOOGLE_DRIVE"
        assert records[0]["status"] == "FAILED"


def test_google_drive_settings_persistence_and_actionable_error(tmp_path):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    with TestClient(app) as client:
        _seed_oauth_drive_integration(
            client,
            token_json="",
        )

        saved = client.post(
            "/integrations/google-drive",
            json={
                "folder_name": "Manager Documents",
                "folder_id": "",
            },
        )
        assert saved.status_code == 200

        fetched = client.get("/integrations/google-drive")
        assert fetched.status_code == 200
        assert fetched.json()["config"]["credential_mode"] == "oauth_quick_connect"
        assert fetched.json()["config"]["folder_name"] == "Manager Documents"
        assert fetched.json()["config"]["connected"] is False

        created = client.post("/documents", json=_create_payload("Drive fail"))
        document_id = created.json()["id"]
        drive = client.post(f"/documents/{document_id}/export/drive")
        assert drive.status_code == 502
        assert "quick connect oauth configuration is incomplete" in drive.json()["detail"].lower()


def test_google_drive_settings_require_quick_connect_first(tmp_path):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    with TestClient(app) as client:
        saved = client.post(
            "/integrations/google-drive",
            json={
                "folder_name": "Manager Documents",
                "folder_id": "",
            },
        )
        assert saved.status_code == 400
        assert "run quick connect first" in saved.json()["detail"].lower()


def test_google_drive_settings_update_folder_for_oauth_mode(tmp_path):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    with TestClient(app) as client:
        _seed_oauth_drive_integration(client, folder_name="Old Folder", folder_id="")

        saved = client.post(
            "/integrations/google-drive",
            json={
                "folder_name": "Updated Folder",
                "folder_id": "folder-123",
            },
        )
        assert saved.status_code == 200
        body = saved.json()
        assert body["config"]["credential_mode"] == "oauth_quick_connect"
        assert body["config"]["folder_name"] == "Updated Folder"
        assert body["config"]["folder_id"] == "folder-123"
        assert body["config"]["connected"] is True


def test_google_drive_quick_connect_start_does_not_include_granted_scopes(tmp_path, monkeypatch):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    captured_kwargs: dict[str, str] = {}

    class FakeFlow:
        redirect_uri = ""

        def authorization_url(self, **kwargs):
            captured_kwargs.update({str(k): str(v) for k, v in kwargs.items()})
            return "https://accounts.google.com/o/oauth2/auth?mock=1", "state-123"

    monkeypatch.setattr(
        api_module.Flow,
        "from_client_config",
        lambda *args, **kwargs: FakeFlow(),
    )

    with TestClient(app) as client:
        response = client.post(
            "/integrations/google-drive/quick-connect/start",
            json={
                "folder_name": "Manager Documents",
                "client_id": "client-id",
                "client_secret": "client-secret",
            },
        )
        assert response.status_code == 200
        assert response.json()["state"] == "state-123"
        assert captured_kwargs.get("access_type") == "offline"
        assert captured_kwargs.get("prompt") == "consent"
        assert "include_granted_scopes" not in captured_kwargs


def test_build_drive_export_name_is_readable_and_stable():
    name = build_drive_export_name(
        title="SEO & GEO Outlook 2026",
        document_id="123e4567-e89b-12d3-a456-426614174000",
    )
    assert name.startswith("SEO & GEO Outlook 2026")
    assert name.endswith("[MDA-123e4567]")
    assert "/" not in name
    assert "\\" not in name


def test_google_drive_export_uses_readable_drive_name(tmp_path, monkeypatch):
    app = create_app(
        db_path=tmp_path / "manager_agent_api.db",
        exports_dir=tmp_path / "exports",
    )

    captured: dict[str, str] = {}

    class FakeDriveClient:
        def upload_docx_as_google_doc(self, local_docx_path: Path, *, document_name: str = "") -> dict:
            captured["document_name"] = document_name
            assert local_docx_path.exists()
            return {"webViewLink": "https://docs.google.com/document/d/doc-123/edit"}

    monkeypatch.setattr(
        api_module,
        "_build_google_drive_client_from_config",
        lambda *, service, config: FakeDriveClient(),
    )

    with TestClient(app) as client:
        _seed_oauth_drive_integration(client)

        created = client.post(
            "/documents",
            json=_create_payload("SEO & GEO Outlook 2026"),
        )
        document_id = created.json()["id"]

        exported = client.post(f"/documents/{document_id}/export/drive")
        assert exported.status_code == 200
        assert exported.json()["status"] == "SUCCESS"
        assert exported.json()["external_url"].startswith("https://docs.google.com/document/d/")

        expected_name = build_drive_export_name(
            title="SEO & GEO Outlook 2026",
            document_id=document_id,
        )
        assert captured["document_name"] == expected_name


def test_docx_export_removes_inline_markdown_markers_and_keeps_text(tmp_path):
    exported = export_markdown_like_to_docx(
        title="SEO & GEO Outlook",
        content=(
            "## Strategic Context\n"
            "**Market Shift**: Customer journeys move to AI answers.\n"
            "- **CPC Inflation**: 20.5% increase in 2025.\n"
        ),
        output_dir=tmp_path / "exports",
        document_id="doc-1",
    )
    assert exported.exists()

    docx = DocxDocument(str(exported))
    all_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    assert "**" not in all_text
    assert "__" not in all_text
    assert "Market Shift" in all_text
    assert "CPC Inflation" in all_text


def test_docx_export_does_not_duplicate_title_when_h1_matches(tmp_path):
    exported = export_markdown_like_to_docx(
        title="SEO & GEO Outlook",
        content="# SEO & GEO Outlook\n\n## Executive Summary\n- Point",
        output_dir=tmp_path / "exports",
        document_id="doc-2",
    )
    docx = DocxDocument(str(exported))
    all_text = [paragraph.text.strip() for paragraph in docx.paragraphs if paragraph.text.strip()]
    assert all_text.count("SEO & GEO Outlook") == 1
