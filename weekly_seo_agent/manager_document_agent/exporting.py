from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import re

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def export_markdown_like_to_docx(*, title: str, content: str, output_dir: Path, document_id: str) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"{timestamp}_{document_id}.docx"
    target = output_dir / filename

    docx = DocxDocument()
    _configure_doc_styles(docx)
    content_lines = content.splitlines()
    normalized_title = _normalize_for_heading_compare(title)
    first_h1 = _extract_first_h1(content_lines)
    if normalized_title and normalized_title != _normalize_for_heading_compare(first_h1):
        docx.add_heading(title.strip(), level=0)

    line_index = 0
    while line_index < len(content_lines):
        raw_line = content_lines[line_index]
        line = raw_line.rstrip()
        if not line:
            line_index += 1
            continue
        if _is_table_header_line(line):
            table_block, next_index = _consume_markdown_table(content_lines, line_index)
            if table_block:
                _add_markdown_table(docx, table_block)
                line_index = next_index
                continue
        if line.startswith("### "):
            _add_heading_with_inline_markdown(docx, line[4:].strip(), level=3)
            line_index += 1
            continue
        if line.startswith("## "):
            _add_heading_with_inline_markdown(docx, line[3:].strip(), level=2)
            line_index += 1
            continue
        if line.startswith("# "):
            _add_heading_with_inline_markdown(docx, line[2:].strip(), level=1)
            line_index += 1
            continue
        list_match = _match_markdown_list_item(raw_line)
        if list_match is not None:
            style_name, item_text = list_match
            paragraph = _add_list_paragraph(docx, style_name=style_name)
            _append_inline_markdown_runs(paragraph, item_text)
            line_index += 1
            continue
        paragraph = docx.add_paragraph()
        _append_inline_markdown_runs(paragraph, line)
        paragraph.paragraph_format.space_after = Pt(6)
        line_index += 1

    docx.save(str(target))
    return target


def build_drive_export_name(*, title: str, document_id: str) -> str:
    normalized_title = _normalize_drive_title(title)
    safe_id = re.sub(r"[^a-zA-Z0-9]+", "", (document_id or "").strip())[:8] or "DOC"
    return f"{normalized_title} [MDA-{safe_id}]"


def _normalize_drive_title(title: str) -> str:
    base = re.sub(r"\s+", " ", (title or "").strip())
    if not base:
        base = "Manager Document"
    base = re.sub(r"[\\/:*?\"<>|]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip(" .-_")
    if not base:
        base = "Manager Document"
    if len(base) > 120:
        base = base[:120].rstrip()
    return base


def _extract_first_h1(lines: list[str]) -> str:
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            return line[2:].strip()
        return ""
    return ""


def _normalize_for_heading_compare(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip()).casefold()


def _add_heading_with_inline_markdown(docx: DocxDocument, text: str, *, level: int) -> None:
    paragraph = docx.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    _append_inline_markdown_runs(paragraph, text)


def _append_inline_markdown_runs(paragraph: Paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*[^*\n]+\*\*|__[^_\n]+__|\*[^*\n]+\*|_[^_\n]+_|`[^`\n]+`)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(_strip_inline_markdown_markers(text[cursor:match.start()]))
        token = match.group(0)
        content = token[1:-1]
        run = paragraph.add_run(content)
        if token.startswith("**") or token.startswith("__"):
            run.bold = True
            run.text = token[2:-2]
        elif token.startswith("*") or token.startswith("_"):
            run.italic = True
            run.text = token[1:-1]
        elif token.startswith("`"):
            run.text = token[1:-1]
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(_strip_inline_markdown_markers(text[cursor:]))


def _strip_inline_markdown_markers(value: str) -> str:
    cleaned = value.replace("**", "").replace("__", "")
    cleaned = cleaned.replace("`", "")
    return cleaned


def _is_table_header_line(line: str) -> bool:
    return "|" in line and line.count("|") >= 2


def _is_table_separator_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return False
    parts = [part.strip() for part in stripped.split("|")]
    cells = [cell for cell in parts if cell]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell) is not None for cell in cells)


def _parse_table_row(line: str) -> list[str]:
    parts = [part.strip() for part in line.strip().split("|")]
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [_strip_inline_markdown_markers(cell) for cell in parts]


def _consume_markdown_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    if start_index + 1 >= len(lines):
        return [], start_index + 1
    header_line = lines[start_index].rstrip()
    separator_line = lines[start_index + 1].rstrip()
    if not _is_table_separator_line(separator_line):
        return [], start_index + 1
    rows: list[list[str]] = []
    header = _parse_table_row(header_line)
    if not header:
        return [], start_index + 1
    rows.append(header)

    index = start_index + 2
    while index < len(lines):
        current = lines[index].rstrip()
        if not current.strip():
            break
        if "|" not in current:
            break
        parsed = _parse_table_row(current)
        if not parsed:
            break
        rows.append(parsed)
        index += 1
    return rows, index


def _add_markdown_table(docx: DocxDocument, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    if max_cols == 0:
        return
    table = docx.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"
    for col_index, cell_value in enumerate(rows[0]):
        table.rows[0].cells[col_index].text = cell_value

    for row in rows[1:]:
        table_row = table.add_row()
        for col_index in range(max_cols):
            value = row[col_index] if col_index < len(row) else ""
            table_row.cells[col_index].text = value


def _configure_doc_styles(docx: DocxDocument) -> None:
    normal_style = docx.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    normal_paragraph = normal_style.paragraph_format
    normal_paragraph.space_before = Pt(0)
    normal_paragraph.space_after = Pt(6)
    normal_paragraph.line_spacing = 1.15


def _match_markdown_list_item(raw_line: str) -> tuple[str, str] | None:
    match = re.match(r"^(?P<indent>\s*)(?P<marker>[-*]|\d+[.)])\s+(?P<text>.+?)\s*$", raw_line)
    if match is None:
        return None
    indent = len(match.group("indent").replace("\t", "    "))
    level = min(2, indent // 2)
    marker = match.group("marker")
    text = match.group("text").strip()
    if not text:
        return None
    is_numbered = bool(re.match(r"^\d+[.)]$", marker))
    base = "List Number" if is_numbered else "List Bullet"
    style_name = base if level == 0 else f"{base} {level + 1}"
    return style_name, text


def _add_list_paragraph(docx: DocxDocument, *, style_name: str) -> Paragraph:
    try:
        paragraph = docx.add_paragraph(style=style_name)
    except KeyError:
        fallback = "List Number" if style_name.startswith("List Number") else "List Bullet"
        paragraph = docx.add_paragraph(style=fallback)
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph
