from __future__ import annotations

import re
import unicodedata
from datetime import date, timedelta
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, RGBColor

from weekly_seo_agent.weekly_reporting_agent.models import (
    AnalysisResult,
    DateWindow,
    ExternalSignal,
    Finding,
    KeyDelta,
    MetricSummary,
)


BOLD_MARKDOWN_RE = re.compile(r"\*\*(.+?)\*\*")
SIGNED_VALUE_RE = re.compile(
    r"(?<!\w)([+-](?:\d[\d\s.,]*)(?:k|K|m|M|b|B)?%?(?:\s*pp)?)"
)

DARK_GREEN = RGBColor(0x1B, 0x5E, 0x20)
DARK_RED = RGBColor(0x8B, 0x00, 0x00)

CAMPAIGN_EVENT_TOKENS = (
    "black week",
    "black friday",
    "cyber monday",
    "smart week",
    "allegro days",
    "megaraty",
    "prime day",
    "sale",
    "promocj",
    "wyprzedaz",
    "deal",
    "rabat",
    "kupon",
)

CAMPAIGN_QUERY_TOKENS = (
    "black week",
    "black friday",
    "cyber monday",
    "smart week",
    "allegro days",
    "megaraty",
    "allegro smart",
    "promocj",
    "wyprzedaz",
    "rabat",
    "kupon",
    "prime day",
    "sale",
)

COMPETITOR_NAME_TOKENS = (
    ("temu", "Temu"),
    ("amazon", "Amazon"),
    ("aliexpress", "AliExpress"),
    ("shein", "SHEIN"),
    ("ebay", "eBay"),
    ("ceneo", "Ceneo"),
    ("olx", "OLX"),
    ("empik", "Empik"),
    ("media expert", "Media Expert"),
    ("rtv euro agd", "RTV Euro AGD"),
    ("x-kom", "x-kom"),
    ("morele", "Morele"),
)

BRAND_QUERY_TOKENS = (
    "allegro",
    "аллегро",
    "алегро",
)

NOISY_QUERY_EXACT_TOKENS = {
    "all",
    "alegro",
    "allegor",
    "allegroo",
    "allegto",
    "algro",
    "algero",
    "ałegro",
}

QUERY_CLUSTER_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "Brand & navigation",
        (
            "allegro",
            "allegro.pl",
            "allegro pl",
            "allegro moje",
            "logowanie",
            "strona glowna",
            "kontakt",
            "sledzenie paczki",
        ),
    ),
    (
        "WOŚP & charity events",
        ("wosp", "wieniawa", "nocowanka", "licytac", "musial"),
    ),
    (
        "Winter & weather demand",
        (
            "sanki",
            "fajerwerk",
            "petard",
            "odsniez",
            "snieg",
            "raczki",
            "sol drog",
            "swider do lod",
            "lopata do sniegu",
            "pellet",
        ),
    ),
    (
        "Valentine's demand",
        ("walentyn", "prezent na walentynki", "dla niej", "dla niego"),
    ),
    (
        "Automotive intent",
        ("vin", "motoryz", "samochod", "golf", "opon", "akumulator"),
    ),
    (
        "Home & garden spring",
        ("glebogryz", "wertykulator", "traw", "nasion", "sekator", "kosiark", "ogrod"),
    ),
)

EXECUTIVE_THEME_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "Brand & navigation",
        (
            "allegro",
            "allegro.pl",
            "allegro pl",
            "allegro moje",
            "logowanie",
            "strona glowna",
            "kontakt",
            "sledzenie paczki",
        ),
    ),
    (
        "Heating & fireplace demand",
        ("pellet", "komink", "kominek", "drewno opal", "brykiet", "rozpal"),
    ),
    (
        "Winter sports & anti-slip",
        ("sanki", "raczki", "lyzwy", "narty", "snowboard", "lod", "snieg"),
    ),
    (
        "Gardening & spring prep",
        ("glebogryz", "wertykulator", "sekator", "nasion", "traw", "ogrod", "kosiark"),
    ),
    (
        "Spring fashion",
        ("wiosenn", "kurtka", "sukien", "bluz", "trencz", "mokasyn"),
    ),
    (
        "Summer tires & automotive season",
        ("opony letnie", "opona letnia", "opony", "felgi", "letnie"),
    ),
    (
        "WOŚP & charity events",
        ("wosp", "wieniawa", "nocowanka", "licytac", "musial"),
    ),
    (
        "Valentine's gifts",
        ("walentyn", "prezent dla niej", "prezent dla niego"),
    ),
)

TREND_EVENT_DRIVER_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("Valentine's Day", ("walentyn", "valentin", "prezent dla niej", "prezent dla niego")),
    ("Easter", ("wielkanoc", "easter", "pisank", "zajac", "swieconk")),
    ("Spring gardening", ("glebogryz", "wertykulator", "traw", "nasion", "sekator", "kosiark", "lampy solar")),
    ("Winter seasonality", ("sanki", "fajerwerk", "petard", "odsniez", "snieg", "lod", "raczk", "sol drog", "lopata", "pellet")),
    ("Auto demand", ("vin", "opon", "golf", "samochod", "motoryz")),
    ("WOŚP / charity events", ("wosp", "licytac", "wieniawa", "nocowanka")),
    ("Campaign periods", CAMPAIGN_EVENT_TOKENS),
)

TREND_SEASONAL_CALENDAR_TOKENS: tuple[str, ...] = (
    "walentyn",
    "valentin",
    "wielkanoc",
    "easter",
    "pisank",
    "zajac",
    "swieconk",
    "winter",
    "sanki",
    "snieg",
    "lod",
    "fajerwerk",
    "petard",
    "pellet",
    "spring",
    "wiosen",
    "glebogryz",
    "wertykulator",
    "kosiark",
    "ogrod",
    "opony letnie",
    "opona letnia",
)

TREND_EVENT_CAMPAIGN_TOKENS: tuple[str, ...] = CAMPAIGN_EVENT_TOKENS + (
    "wosp",
    "charity",
    "licytac",
    "event",
    "campaign",
)

EXTERNAL_SIGNAL_SOURCE_QUALITY: tuple[tuple[str, str], ...] = (
    ("Google Search Status", "high"),
    ("Google Search Central Blog", "high"),
    ("Public Holidays", "high"),
    ("NBP FX", "high"),
    ("IMGW warnings", "high"),
    ("SEO Update Analysis", "medium"),
    ("Campaign tracker", "medium"),
    ("News RSS", "medium"),
    ("News HTML", "low"),
    ("SEO Team Presentations", "medium"),
    ("Historical SEO Reports", "medium"),
    ("SEO Status Log", "medium"),
    ("Product Trends", "high"),
)

NEWS_GMV_CATEGORY_RULES: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("Home & heating", ("pellet", "kominek", "piec", "ogrzew", "opa", "brykiet")),
    ("Winter & weather demand", ("sanki", "śnieg", "snieg", "fajerwerk", "petard", "odśnie", "odsnie", "lód", "lod", "raczk")),
    ("Automotive", ("opon", "vin", "samochod", "motoryz", "części", "czesci")),
    ("Electronics", ("iphone", "samsung", "realme", "oppo", "playstation", "ps5")),
    ("Campaigns & promotions", ("allegro days", "smart week", "black friday", "cyber monday", "sale", "promocj", "rabat", "kupon")),
    ("Marketplace competition", ("temu", "amazon", "shein", "ceneo", "olx", "aliexpress")),
    ("Regulatory & tax", ("vat", "ksef", "e-faktur", "podatek", "regulation", "compliance")),
    ("Logistics & payments", ("kurier", "dostaw", "warehouse", "logistics", "inpost", "dpd", "dhl", "payu", "przelewy24", "blik", "visa", "mastercard", "awaria")),
)

# GA4 is intentionally excluded from decision logic until parity checks with GSC are stable.
USE_GA4_IN_REPORT = False
# Keep report short for decision-makers: no appendix in final output.
INCLUDE_APPENDIX_IN_REPORT = False

CONTINUITY_THEME_TOKENS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("Seasonality", ("season", "winter", "ferie", "weather", "temperature")),
    ("Events/Campaigns", ("campaign", "event", "wosp", "walent", "allegro days", "smart week")),
    ("Algorithm", ("algorithm", "update", "discover", "google")),
    ("Technical/UX", ("cwv", "pagespeed", "ux", "device", "template", "page name")),
    ("Non-brand trends", ("non-brand", "trend", "product", "query cluster")),
)

SUPPORTING_CONTEXT_CATEGORY_TOKENS = (
    "algorithm",
    "algorithm/serp context",
    "serp behavior context",
    "serp features",
)


def _pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def _signed_pct(value: float) -> str:
    return f"{value * 100:+.2f}%"


def _fmt_int(value: float | int) -> str:
    try:
        rounded = int(round(float(value)))
    except (TypeError, ValueError):
        return "0"
    return f"{rounded:,}".replace(",", " ")


def _fmt_signed_int(value: float | int) -> str:
    try:
        raw = float(value)
    except (TypeError, ValueError):
        raw = 0.0
    sign = "+" if raw >= 0 else "-"
    return f"{sign}{_fmt_int(abs(raw))}"


def _fmt_compact(value: float | int) -> str:
    try:
        raw = float(value)
    except (TypeError, ValueError):
        raw = 0.0
    abs_raw = abs(raw)
    suffix = ""
    scale = 1.0
    if abs_raw >= 1_000_000_000:
        suffix = "B"
        scale = 1_000_000_000.0
    elif abs_raw >= 1_000_000:
        suffix = "M"
        scale = 1_000_000.0
    elif abs_raw >= 1_000:
        suffix = "k"
        scale = 1_000.0
    if suffix:
        value_scaled = raw / scale
        text = f"{value_scaled:.1f}"
        if text.endswith(".0"):
            text = text[:-2]
        return f"{text}{suffix}"
    return _fmt_int(raw)


def _fmt_signed_compact(value: float | int) -> str:
    try:
        raw = float(value)
    except (TypeError, ValueError):
        raw = 0.0
    sign = "+" if raw >= 0 else "-"
    return f"{sign}{_fmt_compact(abs(raw))}"


def _ga4_metric_available(ga4_context: dict[str, object], metric: str) -> bool:
    availability = ga4_context.get("metric_availability", {})
    if not isinstance(availability, dict):
        return True
    return bool(availability.get(metric, True))


def _ga4_num(summary: dict[str, object], metric: str) -> float | None:
    if not isinstance(summary, dict):
        return None
    value = summary.get(metric)
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _ratio_delta(current: float, baseline: float) -> float:
    if baseline == 0:
        return 1.0 if current > 0 else 0.0
    return (current - baseline) / baseline


def _flatten_findings(scope_results: list[tuple[str, AnalysisResult]]) -> list[tuple[str, Finding]]:
    result: list[tuple[str, Finding]] = []
    for scope_name, analysis in scope_results:
        for finding in analysis.findings:
            result.append((scope_name, finding))
    return result


def _normalize_text(value: str) -> str:
    ascii_text = (
        unicodedata.normalize("NFKD", value)
        .encode("ascii", "ignore")
        .decode("ascii")
        .lower()
    )
    return re.sub(r"\s+", " ", ascii_text)


def _contains_any(text: str, keywords: tuple[str, ...]) -> bool:
    normalized = _normalize_text(text)
    return any(keyword in normalized for keyword in keywords)


def _is_noise_or_brand_query(query: str) -> bool:
    normalized = _normalize_text(query).strip()
    if not normalized:
        return True
    if any(token in normalized for token in BRAND_QUERY_TOKENS):
        return True
    if normalized in NOISY_QUERY_EXACT_TOKENS:
        return True
    compact = normalized.replace(" ", "")
    if compact in NOISY_QUERY_EXACT_TOKENS:
        return True
    if re.fullmatch(r"[a-z]{1,3}", normalized):
        return True
    return False


def _find_scope(scope_results: list[tuple[str, AnalysisResult]], scope_name: str) -> AnalysisResult | None:
    for name, analysis in scope_results:
        if name == scope_name:
            return analysis
    return None


def _campaign_query_rows(query_scope: AnalysisResult | None) -> list[KeyDelta]:
    if query_scope is None:
        return []
    rows = query_scope.top_winners + query_scope.top_losers
    out: list[KeyDelta] = []
    for row in rows:
        text = _normalize_text(row.key)
        if any(token in text for token in CAMPAIGN_QUERY_TOKENS):
            out.append(row)
            continue
        if any(token in text for token, _ in COMPETITOR_NAME_TOKENS) and any(
            token in text for token in ("promocj", "sale", "black", "cyber", "week")
        ):
            out.append(row)
    out.sort(key=lambda item: abs(item.click_delta_vs_previous), reverse=True)
    return out[:8]


def _campaign_event_context(
    external_signals: list[ExternalSignal],
    query_scope: AnalysisResult | None,
) -> dict[str, object]:
    allegro_events: list[ExternalSignal] = []
    competitor_events: list[tuple[ExternalSignal, str]] = []
    seen_allegro: set[str] = set()
    seen_comp: set[str] = set()

    for signal in external_signals:
        text = _normalize_text(f"{signal.title} {signal.details} {signal.source}")
        source_text = _normalize_text(signal.source)
        is_campaign_tracker_source = "campaign tracker" in source_text
        has_campaign_token = any(token in text for token in CAMPAIGN_EVENT_TOKENS)
        has_competitor_token = any(token in text for token, _ in COMPETITOR_NAME_TOKENS)
        has_promo_context = any(
            token in text
            for token in ("promocj", "promo", "sale", "deal", "discount", "black", "cyber", "week")
        )
        if not (has_campaign_token or is_campaign_tracker_source or (has_competitor_token and has_promo_context)):
            continue

        if "allegro" in text or "smart week" in text or "allegro days" in text or "megaraty" in text:
            key = re.sub(
                r"\s*[-|:]\s*[^-|:]{1,80}$",
                "",
                _normalize_text(signal.title),
            ).strip() or _normalize_text(signal.title)
            if key not in seen_allegro:
                seen_allegro.add(key)
                allegro_events.append(signal)

        for token, label in COMPETITOR_NAME_TOKENS:
            if token in text:
                key = (
                    re.sub(
                        r"\s*[-|:]\s*[^-|:]{1,80}$",
                        "",
                        _normalize_text(signal.title),
                    ).strip()
                    + "::"
                    + label.lower()
                )
                if key not in seen_comp:
                    seen_comp.add(key)
                    competitor_events.append((signal, label))
                break

    query_events = _campaign_query_rows(query_scope)
    return {
        "allegro_events": sorted(allegro_events, key=lambda item: item.day, reverse=True)[:10],
        "competitor_events": sorted(
            competitor_events, key=lambda item: item[0].day, reverse=True
        )[:10],
        "query_events": query_events,
    }


def _parse_iso_date(value: object) -> date | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return date.fromisoformat(text[:10])
    except ValueError:
        return None


def _shorten(value: object, limit: int = 120) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 3)].rstrip() + "..."


def _extract_country_hint(text: str, fallback_country: str) -> str:
    normalized = str(text or "").upper()
    token = re.search(r"\b(PL|CZ|SK|HU)\b", normalized)
    if token:
        return token.group(1)
    return (fallback_country or "PL").strip().upper() or "PL"


def _canonical_timeline_event(text: str) -> tuple[str, str]:
    raw = str(text or "").strip()
    if not raw:
        return "", ""

    trimmed = re.sub(r"\s*[-|]\s*[^-|]{1,80}$", "", raw).strip()
    normalized = _normalize_text(trimmed)
    normalized = re.sub(r"\b(via|source|zrodlo)\b.*$", "", normalized).strip()

    known_rules = (
        ("allegro days", "Allegro Days"),
        ("megaraty", "MegaRaty"),
        ("smart week", "Smart Week"),
        ("black friday", "Black Friday"),
        ("cyber monday", "Cyber Monday"),
        ("black week", "Black Week"),
        ("prime day", "Prime Day"),
        ("wosp", "WOSP"),
    )
    for token, label in known_rules:
        if token in normalized:
            return token, label

    tokens = [
        row
        for row in normalized.split()
        if row not in {"planned", "campaign", "event", "update", "breaking", "news"}
    ]
    canonical = " ".join(tokens[:12]).strip() or normalized
    display = " ".join(canonical.split()[:8]).strip() or trimmed
    return canonical, display


def _timeline_impact_rank(value: str) -> int:
    lowered = _normalize_text(value)
    if "high" in lowered:
        return 3
    if "medium" in lowered:
        return 2
    if "low" in lowered:
        return 1
    return 0


def _marketplace_timeline_rows(
    *,
    additional_context: dict[str, object] | None,
    campaign_context: dict[str, object] | None,
    report_country_code: str,
    max_rows: int = 40,
) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    base_country = (report_country_code or "PL").strip().upper() or "PL"

    market_events = (additional_context or {}).get("market_event_calendar", {})
    if isinstance(market_events, dict):
        market_country = str(market_events.get("country_code", "")).strip().upper() or base_country
        market_rows = market_events.get("events", [])
        if isinstance(market_rows, list):
            for row in market_rows:
                if not isinstance(row, dict):
                    continue
                day = _parse_iso_date(row.get("date"))
                if not isinstance(day, date):
                    continue
                conf = int(row.get("confidence", 0) or 0)
                impact = (
                    f"{str(row.get('impact_level', 'LOW')).strip().upper()}/"
                    f"{str(row.get('impact_direction', 'Mixed')).strip()} ({conf}/100)"
                )
                rows.append(
                    {
                        "day": day,
                        "country": market_country,
                        "track": "Market event",
                        "event": _shorten(row.get("title", ""), 130) or "Market signal",
                        "source": _shorten(row.get("source", "GDELT DOC API"), 50),
                        "impact": impact,
                    }
                )

    if isinstance(campaign_context, dict):
        allegro_events = campaign_context.get("allegro_events", [])
        if isinstance(allegro_events, list):
            for signal in allegro_events:
                if not isinstance(signal, ExternalSignal):
                    continue
                country = _extract_country_hint(
                    f"{signal.source} {signal.title} {signal.details}",
                    base_country,
                )
                rows.append(
                    {
                        "day": signal.day,
                        "country": country,
                        "track": "Allegro campaign",
                        "event": _shorten(signal.title, 130),
                        "source": _shorten(signal.source, 50),
                        "impact": str(signal.severity or "info").strip().lower(),
                    }
                )

        competitor_events = campaign_context.get("competitor_events", [])
        if isinstance(competitor_events, list):
            for entry in competitor_events:
                if (
                    not isinstance(entry, tuple)
                    or len(entry) != 2
                    or not isinstance(entry[0], ExternalSignal)
                ):
                    continue
                signal, competitor = entry
                country = _extract_country_hint(
                    f"{signal.source} {signal.title} {signal.details}",
                    base_country,
                )
                rows.append(
                    {
                        "day": signal.day,
                        "country": country,
                        "track": f"Competitor promo ({competitor})",
                        "event": _shorten(signal.title, 130),
                        "source": _shorten(signal.source, 50),
                        "impact": str(signal.severity or "info").strip().lower(),
                    }
                )

    promo_radar = (additional_context or {}).get("competitor_promo_radar", {})
    if isinstance(promo_radar, dict):
        promo_rows = promo_radar.get("rows", [])
        if isinstance(promo_rows, list):
            for row in promo_rows:
                if not isinstance(row, dict):
                    continue
                day = _parse_iso_date(row.get("date"))
                if not isinstance(day, date):
                    continue
                title = str(row.get("title", "")).strip()
                competitor = "market"
                title_norm = _normalize_text(title)
                for token, label in COMPETITOR_NAME_TOKENS:
                    if token in title_norm:
                        competitor = label
                        break
                country = _extract_country_hint(
                    f"{row.get('source', '')} {title}",
                    base_country,
                )
                rows.append(
                    {
                        "day": day,
                        "country": country,
                        "track": f"Promo radar ({competitor})",
                        "event": _shorten(title, 130),
                        "source": _shorten(row.get("source", "Google News RSS query"), 50),
                        "impact": "medium",
                    }
                )

    merged: dict[tuple[str, str, str], dict[str, object]] = {}
    for row in sorted(
        rows,
        key=lambda item: (
            item.get("day") if isinstance(item.get("day"), date) else date.min,
            str(item.get("track", "")),
        ),
    ):
        day = row.get("day")
        if not isinstance(day, date):
            continue

        country = str(row.get("country", "")).strip().upper() or base_country
        track = str(row.get("track", "")).strip()
        source = str(row.get("source", "")).strip()
        impact = str(row.get("impact", "")).strip()
        canonical_key, canonical_label = _canonical_timeline_event(str(row.get("event", "")))
        canonical_key = canonical_key or _normalize_text(str(row.get("event", ""))).strip()
        if not canonical_key:
            continue
        key = (day.isoformat(), _normalize_text(country), canonical_key)
        score = _timeline_impact_rank(impact)

        existing = merged.get(key)
        if existing is None:
            merged[key] = {
                "day": day,
                "country": country,
                "track": track,
                "event": _shorten(canonical_label or str(row.get("event", "")), 130) or "Market signal",
                "source": _shorten(source, 50),
                "impact": impact or "info",
                "_impact_score": score,
                "_tracks": [track] if track else [],
                "_sources": [source] if source else [],
                "canonical_event": canonical_key,
            }
            continue

        tracks = existing.get("_tracks", [])
        if isinstance(tracks, list) and track and track not in tracks:
            tracks.append(track)
        sources = existing.get("_sources", [])
        if isinstance(sources, list) and source and source not in sources:
            sources.append(source)

        current_score = int(existing.get("_impact_score", 0) or 0)
        if score > current_score:
            existing["impact"] = impact or existing.get("impact", "info")
            existing["_impact_score"] = score

    deduped: list[dict[str, object]] = []
    for row in sorted(
        merged.values(),
        key=lambda item: (
            item.get("day") if isinstance(item.get("day"), date) else date.min,
            str(item.get("country", "")),
            str(item.get("event", "")),
        ),
    ):
        tracks = row.pop("_tracks", [])
        sources = row.pop("_sources", [])
        row.pop("_impact_score", None)
        if isinstance(tracks, list) and tracks:
            dedup_tracks = sorted({str(item).strip() for item in tracks if str(item).strip()})
            if dedup_tracks:
                row["track"] = " + ".join(dedup_tracks[:2]) + (" + more" if len(dedup_tracks) > 2 else "")
        if isinstance(sources, list) and sources:
            dedup_sources = sorted({str(item).strip() for item in sources if str(item).strip()})
            if dedup_sources:
                row["source"] = _shorten("; ".join(dedup_sources[:2]), 50)
        deduped.append(row)
    return deduped[: max(1, max_rows)]


def _priority_action_for_hypothesis(row: dict[str, object]) -> tuple[str, str]:
    category = _normalize_text(str(row.get("category", "")))
    if "campaign" in category:
        return (
            "Validate campaign timing and paid-share overlap for the top impacted clusters.",
            "SEO + Commercial | 3d",
        )
    if "algorithm" in category:
        return (
            "Run before/after checks on affected page/query clusters since the update date.",
            "SEO | 48h",
        )
    if "technical" in category or "ux" in category:
        return (
            "Prioritize technical audit for the weakest device/page segments.",
            "SEO + Web Performance | 5d",
        )
    if "data quality" in category:
        return (
            "Fix missing data source and rerun confidence-sensitive hypotheses.",
            "SEO Ops | 24h",
        )
    if "seasonality" in category or "events" in category:
        return (
            "Align content and merchandising timing with demand rotation drivers.",
            "SEO + Merchandising | this week",
        )
    return ("Validate the hypothesis with the next weekly window before escalation.", "SEO Team | next run")


def _top_priority_actions(
    hypotheses: list[dict[str, object]],
    limit: int = 3,
) -> list[str]:
    out: list[str] = []
    for row in hypotheses[: max(1, limit)]:
        action, owner_eta = _priority_action_for_hypothesis(row)
        category = str(row.get("category", "Unknown")).strip() or "Unknown"
        out.append(f"{category}: {action} ({owner_eta})")
    return out


def _direction_marker(
    delta: float,
    *,
    better_when_lower: bool = False,
    tolerance: float = 0.0,
) -> str:
    if abs(float(delta)) <= float(tolerance):
        return "FLAT"
    is_positive = float(delta) > 0.0
    is_good = (not better_when_lower and is_positive) or (better_when_lower and not is_positive)
    return "UP" if is_good else "DOWN"


def _format_marker_with_value(marker: str, value_text: str) -> str:
    return f"{marker} {value_text}".strip()


def _kpi_snapshot_table_lines(
    *,
    totals: dict[str, MetricSummary],
    additional_context: dict[str, object] | None,
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    yoy = totals["yoy_52w"]
    long_kpi = ((additional_context or {}).get("long_window_context", {}) or {}).get("kpi", {})
    if not isinstance(long_kpi, dict):
        long_kpi = {}

    clicks_wow_pct = _ratio_delta(current.clicks, previous.clicks) * 100.0
    clicks_yoy_pct = _ratio_delta(current.clicks, yoy.clicks) * 100.0
    clicks_mom_pct = long_kpi.get("clicks_delta_pct_vs_previous")
    clicks_mom_text = (
        _format_marker_with_value(
            _direction_marker(float(clicks_mom_pct), tolerance=0.15),
            f"{float(clicks_mom_pct):+.2f}%",
        )
        if isinstance(clicks_mom_pct, (int, float))
        else "n/a"
    )

    impressions_wow_pct = _ratio_delta(current.impressions, previous.impressions) * 100.0
    impressions_yoy_pct = _ratio_delta(current.impressions, yoy.impressions) * 100.0
    impressions_mom_pct = long_kpi.get("impressions_delta_pct_vs_previous")
    impressions_mom_text = (
        _format_marker_with_value(
            _direction_marker(float(impressions_mom_pct), tolerance=0.15),
            f"{float(impressions_mom_pct):+.2f}%",
        )
        if isinstance(impressions_mom_pct, (int, float))
        else "n/a"
    )

    ctr_wow_pp = (current.ctr - previous.ctr) * 100.0
    ctr_yoy_pp = (current.ctr - yoy.ctr) * 100.0
    ctr_mom_pp = long_kpi.get("ctr_delta_pp_vs_previous")
    ctr_mom_text = (
        _format_marker_with_value(
            _direction_marker(float(ctr_mom_pp), tolerance=0.02),
            f"{float(ctr_mom_pp):+.2f} pp",
        )
        if isinstance(ctr_mom_pp, (int, float))
        else "n/a"
    )

    position_wow = current.position - previous.position
    position_yoy = current.position - yoy.position
    position_mom = long_kpi.get("position_delta_vs_previous")
    position_mom_text = (
        _format_marker_with_value(
            _direction_marker(float(position_mom), better_when_lower=True, tolerance=0.02),
            f"{float(position_mom):+.2f}",
        )
        if isinstance(position_mom, (int, float))
        else "n/a"
    )

    lines = [
        "## KPI snapshot (WoW/MoM/YoY)",
        "| KPI | WoW | MoM | YoY |",
        "|---|---|---|---|",
        "| Clicks | "
        + _format_marker_with_value(_direction_marker(clicks_wow_pct, tolerance=0.15), f"{clicks_wow_pct:+.2f}%")
        + " | "
        + clicks_mom_text
        + " | "
        + _format_marker_with_value(_direction_marker(clicks_yoy_pct, tolerance=0.15), f"{clicks_yoy_pct:+.2f}%")
        + " |",
        "| Impressions | "
        + _format_marker_with_value(_direction_marker(impressions_wow_pct, tolerance=0.15), f"{impressions_wow_pct:+.2f}%")
        + " | "
        + impressions_mom_text
        + " | "
        + _format_marker_with_value(_direction_marker(impressions_yoy_pct, tolerance=0.15), f"{impressions_yoy_pct:+.2f}%")
        + " |",
        "| CTR | "
        + _format_marker_with_value(_direction_marker(ctr_wow_pp, tolerance=0.02), f"{ctr_wow_pp:+.2f} pp")
        + " | "
        + ctr_mom_text
        + " | "
        + _format_marker_with_value(_direction_marker(ctr_yoy_pp, tolerance=0.02), f"{ctr_yoy_pp:+.2f} pp")
        + " |",
        "| Avg position | "
        + _format_marker_with_value(_direction_marker(position_wow, better_when_lower=True, tolerance=0.02), f"{position_wow:+.2f}")
        + " | "
        + position_mom_text
        + " | "
        + _format_marker_with_value(_direction_marker(position_yoy, better_when_lower=True, tolerance=0.02), f"{position_yoy:+.2f}")
        + " |",
    ]
    return lines


def _driver_scoreboard_lines(hypotheses: list[dict[str, object]], limit: int = 4) -> list[str]:
    rows = _top_hypotheses_for_actions(hypotheses, limit=limit)
    lines = [
        "## Driver scoreboard",
        "| Driver | Impact | Confidence | Action owner |",
        "|---|---|---|---|",
    ]
    if not rows:
        lines.append("| (no prioritized drivers) | - | - | SEO Team |")
        return lines
    for row in rows[: max(1, limit)]:
        category = str(row.get("category", "Unknown")).strip() or "Unknown"
        impact_score = int(row.get("impact_score", 0) or 0)
        if impact_score >= 4:
            impact_label = "High"
        elif impact_score == 3:
            impact_label = "Medium"
        else:
            impact_label = "Low"
        confidence = int(row.get("confidence", 0) or 0)
        owner = str(row.get("owner", "")).strip() or "SEO Team"
        lines.append(
            f"| {category} | {impact_label} ({impact_score}/5) | {_confidence_bucket(confidence)} ({confidence}/100) | {owner} |"
        )
    return lines


def _meeting_ready_talking_points(
    *,
    totals: dict[str, MetricSummary],
    hypotheses: list[dict[str, object]],
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    wow_clicks_pct = _ratio_delta(current.clicks, previous.clicks) * 100.0
    top_rows = _top_hypotheses_for_actions(hypotheses, limit=2)
    top_driver = str(top_rows[0].get("category", "Demand timing")).strip() if top_rows else "Demand timing"
    second_driver = str(top_rows[1].get("category", "Campaign timing")).strip() if len(top_rows) > 1 else "Campaign timing"
    primary_action = _top_priority_actions(hypotheses, limit=1)
    action_text = primary_action[0] if primary_action else "SEO Team: Validate top driver on next run (SEO Team | next run)"
    direction_text = "up" if wow_clicks_pct >= 0 else "down"
    return [
        "## Meeting-ready talking points",
        f"- Start with: organic clicks are {direction_text} WoW ({wow_clicks_pct:+.2f}%), and this week is mainly a demand/mix story.",
        f"- Main business risk: if `{top_driver}` persists without action, short-term traffic allocation can underperform in priority categories.",
        f"- Main opportunity: `{second_driver}` is actionable now and can recover traffic share in the next cycle.",
        "- Decision to align on: keep technical SEO escalation as step two unless efficiency weakens.",
        f"- Confirm next step and owner: {action_text}.",
    ]


def _context_snapshot_lines(
    *,
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    additional_context: dict[str, object] | None,
) -> list[str]:
    lines: list[str] = ["## Context snapshot"]
    query_scope = _find_scope(scope_results, "query")
    campaign_context = _campaign_event_context(
        external_signals=external_signals,
        query_scope=query_scope,
    )
    country_hint = str((additional_context or {}).get("country_code", "")).strip().upper() or "PL"
    timeline_rows = _marketplace_timeline_rows(
        additional_context=additional_context,
        campaign_context=campaign_context,
        report_country_code=country_hint,
        max_rows=20,
    )
    if timeline_rows:
        first_day = timeline_rows[0].get("day")
        last_day = timeline_rows[-1].get("day")
        if isinstance(first_day, date) and isinstance(last_day, date):
            lines.append(
                "- **Marketplace timeline (market events + promo calendar)**: "
                f"{_fmt_int(len(timeline_rows))} events on one axis "
                f"({first_day.isoformat()} to {last_day.isoformat()}) for {country_hint}."
            )

    serp_split_line = _serp_appearance_summary_text(additional_context)
    if serp_split_line:
        lines.append("- **SERP appearance mix (WoW/MoM/YoY)**: " + serp_split_line.replace("SERP appearance split (GSC searchAppearance): ", ""))
    daily_serp_line = _daily_serp_feature_shift_line(additional_context)
    if daily_serp_line:
        lines.append("- **Daily SERP feature shifts**: " + daily_serp_line.replace("Daily SERP feature-share shifts: ", ""))
    daily_anomaly_line = _daily_anomaly_detector_line(additional_context)
    if daily_anomaly_line:
        lines.append("- **Daily KPI anomalies**: " + daily_anomaly_line.replace("Daily KPI anomaly detector: ", ""))

    updates_timeline_line = _google_updates_timeline_text(additional_context)
    if updates_timeline_line:
        lines.append("- **Google updates timeline (13 months)**: " + updates_timeline_line.replace("Google update timeline (13M): ", ""))
        lines.append(updates_timeline_line)

    case_study_line = _serp_case_study_text(additional_context)
    if case_study_line:
        lines.append("- **External case-study context (13 months)**: " + case_study_line.replace("SERP case-study scanner (13M): ", ""))
        lines.append(case_study_line)

    daily_story = _build_daily_gsc_storyline(
        additional_context=additional_context,
        external_signals=external_signals,
        weather_summary=weather_summary,
        top_n=3,
    )
    daily_executive_line = str(daily_story.get("executive_line", "")).strip()
    if daily_executive_line:
        lines.append(daily_executive_line)

    if len(lines) == 1:
        lines.append("- Context snapshot has no additional signals in this run.")
    return lines


def _dedupe_report_lines(lines: list[str]) -> list[str]:
    out: list[str] = []
    seen_text: set[str] = set()
    previous_blank = False
    for raw in lines:
        line = str(raw)
        stripped = line.strip()
        if not stripped:
            if previous_blank:
                continue
            out.append("")
            previous_blank = True
            continue
        previous_blank = False
        if stripped.startswith("#") or stripped.startswith("|"):
            out.append(line)
            continue
        norm = _normalize_text(stripped)
        if norm in seen_text:
            continue
        seen_text.add(norm)
        out.append(line)
    return out


def _enforce_section_line_limits(lines: list[str]) -> list[str]:
    limits = {
        "Executive summary": 6,
        "What is happening and why": 40,
        "KPI snapshot (WoW/MoM/YoY)": 8,
        "Driver scoreboard": 8,
        "Meeting-ready talking points": 6,
        "Context snapshot": 9,
        "Evidence coverage check": 12,
        "Hypothesis protocol": 10,
        "Validation plan (next week)": 4,
        "Counterfactual checks": 4,
        "Causality guardrail": 3,
        "Escalation rule": 3,
        "Evidence ledger": 14,
    }
    out: list[str] = []
    idx = 0
    while idx < len(lines):
        line = str(lines[idx])
        out.append(line)
        stripped = line.strip()
        if not stripped.startswith("## "):
            idx += 1
            continue
        title = stripped[3:].strip()
        max_body = limits.get(title)
        if max_body is None:
            idx += 1
            continue
        body: list[str] = []
        look = idx + 1
        while look < len(lines) and not str(lines[look]).strip().startswith("## "):
            body.append(str(lines[look]))
            look += 1
        if len(body) > max_body:
            trimmed = body[:max_body]
            # If section has table, keep header/separator and first rows.
            if any(str(item).strip().startswith("|") for item in body):
                table_header: list[str] = []
                table_rows: list[str] = []
                for row in body:
                    s = str(row).strip()
                    if s.startswith("|"):
                        if len(table_header) < 2:
                            table_header.append(row)
                        else:
                            table_rows.append(row)
                non_table = [row for row in body if not str(row).strip().startswith("|")]
                merged = non_table[: max(0, max_body - min(4, len(table_header) + 2))]
                merged.extend(table_header[:2])
                merged.extend(table_rows[:2])
                trimmed = merged[:max_body]
            trimmed.append("- (Auto-compressed for readability.)")
            out.extend(trimmed)
        else:
            out.extend(body)
        idx = look
    return out


def _cluster_label_for_query(query: str) -> str:
    normalized = _normalize_text(query)
    for label, tokens in QUERY_CLUSTER_RULES:
        if any(token in normalized for token in tokens):
            return label
    return "Other non-brand demand"


def _executive_theme_for_query(query: str) -> str:
    normalized = _normalize_text(query)
    for label, tokens in EXECUTIVE_THEME_RULES:
        if any(token in normalized for token in tokens):
            return label
    return "Other demand themes"


def _executive_yoy_theme_summary(
    query_scope: AnalysisResult | None,
    limit: int = 3,
) -> tuple[str, str]:
    if query_scope is None:
        return "", ""

    by_query: dict[str, KeyDelta] = {}
    for row in query_scope.top_winners + query_scope.top_losers:
        key = _normalize_text(row.key)
        if not key:
            continue
        if _is_noise_or_brand_query(row.key):
            continue
        current = by_query.get(key)
        if current is None or abs(row.click_delta_vs_yoy) > abs(current.click_delta_vs_yoy):
            by_query[key] = row

    buckets: dict[str, dict[str, object]] = {}
    for row in by_query.values():
        normalized_key = _normalize_text(row.key)
        if any(token in normalized_key for token in BRAND_QUERY_TOKENS):
            continue
        if abs(row.click_delta_vs_yoy) < 100:
            continue
        label = _executive_theme_for_query(row.key)
        payload = buckets.setdefault(
            label,
            {"delta": 0.0, "examples": []},
        )
        payload["delta"] = float(payload.get("delta", 0.0)) + float(row.click_delta_vs_yoy)
        examples = payload.get("examples")
        if isinstance(examples, list) and len(examples) < 3:
            examples.append(str(row.key).strip())

    if not buckets:
        return "", ""

    growth = sorted(
        [
            (label, float(payload.get("delta", 0.0)), payload)
            for label, payload in buckets.items()
            if float(payload.get("delta", 0.0)) > 0
        ],
        key=lambda item: item[1],
        reverse=True,
    )[:limit]
    declines = sorted(
        [
            (label, float(payload.get("delta", 0.0)), payload)
            for label, payload in buckets.items()
            if float(payload.get("delta", 0.0)) < 0
        ],
        key=lambda item: item[1],
    )[:limit]

    if len(growth) > 1:
        growth = [row for row in growth if row[0] != "Other demand themes"][:limit]
    if len(declines) > 1:
        declines = [row for row in declines if row[0] != "Other demand themes"][:limit]

    def _render(label: str, delta: float, payload: dict[str, object]) -> str:
        rendered_label = label
        if label == "Other demand themes":
            examples = payload.get("examples")
            if isinstance(examples, list) and examples:
                rendered_examples = ", ".join(
                    f"`{str(item).strip()}`"
                    for item in examples[:2]
                    if str(item).strip()
                )
                if rendered_examples:
                    rendered_label = f"Other demand themes (e.g. {rendered_examples})"
        return f"{rendered_label} ({_fmt_signed_int(delta)} YoY)"

    growth_line = ""
    if growth:
        growth_line = "; ".join(
            _render(label, delta, payload)
            for label, delta, payload in growth
        )
    decline_line = ""
    if declines:
        decline_line = "; ".join(
            _render(label, delta, payload)
            for label, delta, payload in declines
        )
    return growth_line, decline_line


def _query_cluster_rows(query_scope: AnalysisResult | None) -> list[dict[str, object]]:
    if query_scope is None:
        return []

    by_query: dict[str, KeyDelta] = {}
    for row in query_scope.top_winners + query_scope.top_losers:
        key = _normalize_text(row.key)
        if not key:
            continue
        if _is_noise_or_brand_query(row.key):
            continue
        previous = by_query.get(key)
        if previous is None or row.current_clicks > previous.current_clicks:
            by_query[key] = row

    buckets: dict[str, dict[str, object]] = {}
    for row in by_query.values():
        label = _cluster_label_for_query(row.key)
        payload = buckets.setdefault(
            label,
            {
                "cluster": label,
                "rows": 0,
                "current_clicks": 0.0,
                "delta_vs_previous": 0.0,
                "delta_vs_yoy": 0.0,
                "queries": [],
            },
        )
        payload["rows"] = int(payload.get("rows", 0)) + 1
        payload["current_clicks"] = float(payload.get("current_clicks", 0.0)) + float(
            row.current_clicks
        )
        payload["delta_vs_previous"] = float(payload.get("delta_vs_previous", 0.0)) + float(
            row.click_delta_vs_previous
        )
        payload["delta_vs_yoy"] = float(payload.get("delta_vs_yoy", 0.0)) + float(
            row.click_delta_vs_yoy
        )
        payload_queries = payload.get("queries")
        if isinstance(payload_queries, list):
            payload_queries.append((float(row.current_clicks), row.key))

    out = list(buckets.values())
    for row in out:
        samples = row.get("queries")
        if isinstance(samples, list):
            ordered = sorted(
                [item for item in samples if isinstance(item, tuple) and len(item) == 2],
                key=lambda item: float(item[0]),
                reverse=True,
            )
            row["samples"] = [str(item[1]) for item in ordered[:3]]
        else:
            row["samples"] = []
        row.pop("queries", None)

    out.sort(
        key=lambda row: (
            float(row.get("current_clicks", 0.0)),
            abs(float(row.get("delta_vs_previous", 0.0))),
        ),
        reverse=True,
    )
    return out[:8]


def _query_cluster_summary_line(query_scope: AnalysisResult | None) -> str:
    rows = _query_cluster_rows(query_scope)
    if not rows:
        return ""
    gains = sorted(
        [row for row in rows if float(row.get("delta_vs_previous", 0.0)) > 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
        reverse=True,
    )[:3]
    losses = sorted(
        [row for row in rows if float(row.get("delta_vs_previous", 0.0)) < 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
    )[:3]
    largest = sorted(
        rows,
        key=lambda row: float(row.get("current_clicks", 0.0)),
        reverse=True,
    )[:3]

    parts: list[str] = []
    if gains:
        gains_text = "; ".join(
            f"{row.get('cluster', '')} ({_fmt_signed_int(row.get('delta_vs_previous', 0.0))} vs prev)"
            for row in gains
            if str(row.get("cluster", "")).strip()
        )
        if gains_text:
            parts.append(f"gains: {gains_text}")
    if losses:
        losses_text = "; ".join(
            f"{row.get('cluster', '')} ({_fmt_signed_int(row.get('delta_vs_previous', 0.0))} vs prev)"
            for row in losses
            if str(row.get("cluster", "")).strip()
        )
        if losses_text:
            parts.append(f"losses: {losses_text}")
    if largest:
        largest_text = "; ".join(
            f"{row.get('cluster', '')} ({_fmt_int(row.get('current_clicks', 0.0))} clicks)"
            for row in largest
            if str(row.get("cluster", "")).strip()
        )
        if largest_text:
            parts.append(f"largest demand clusters: {largest_text}")
    if not parts:
        return ""
    return "Query-cluster split by direction: " + " | ".join(parts) + "."


def _query_cluster_contribution_rows(
    query_scope: AnalysisResult | None,
    total_click_delta_yoy: float,
    limit: int = 8,
) -> list[dict[str, object]]:
    rows = _query_cluster_rows(query_scope)
    if not rows:
        return []

    tracked_total = sum(float(row.get("delta_vs_yoy", 0.0)) for row in rows)
    if abs(total_click_delta_yoy) >= 1.0:
        denominator = abs(total_click_delta_yoy)
        basis = "total_yoy"
    elif abs(tracked_total) >= 1.0:
        denominator = abs(tracked_total)
        basis = "tracked_yoy"
    else:
        denominator = 1.0
        basis = "tracked_yoy"

    out: list[dict[str, object]] = []
    for row in rows:
        delta_yoy = float(row.get("delta_vs_yoy", 0.0))
        out.append(
            {
                "cluster": str(row.get("cluster", "")).strip(),
                "current_clicks": float(row.get("current_clicks", 0.0)),
                "delta_vs_yoy": delta_yoy,
                "contribution_pct": (delta_yoy / denominator) * 100.0,
                "basis": basis,
                "rows": int(row.get("rows", 0)),
                "samples": row.get("samples", []),
            }
        )
    out.sort(key=lambda item: abs(float(item.get("delta_vs_yoy", 0.0))), reverse=True)
    return out[: max(1, limit)]


def _normalized_page_key(raw_url: str) -> str:
    raw = str(raw_url or "").strip()
    if not raw:
        return ""
    parsed = urlparse(raw)
    host = (parsed.netloc or "").strip().lower()
    if host.startswith("www."):
        host = host[4:]
    path = (parsed.path or "/").strip()
    if not path.startswith("/"):
        path = "/" + path
    if path != "/":
        path = path.rstrip("/")
    query = f"?{parsed.query}" if parsed.query else ""
    return f"{host}{path}{query}"


def _pick_source_quality(source: str) -> str:
    lowered = _normalize_text(source)
    for token, label in EXTERNAL_SIGNAL_SOURCE_QUALITY:
        if _normalize_text(token) in lowered:
            return label
    return "low"


def _signal_quality_rows(
    external_signals: list[ExternalSignal],
    limit: int = 12,
) -> list[dict[str, object]]:
    deduped: dict[str, dict[str, object]] = {}
    severity_rank = {"info": 1, "medium": 2, "high": 3}
    quality_rank = {"low": 1, "medium": 2, "high": 3}

    for signal in sorted(external_signals, key=lambda row: row.day, reverse=True):
        title_key = _normalize_text(re.sub(r"\s*-\s*[^-]{1,50}$", "", signal.title)).strip()
        if not title_key:
            title_key = _normalize_text(signal.title)
        if not title_key:
            continue

        entry = deduped.get(title_key)
        source_quality = _pick_source_quality(signal.source)
        if entry is None:
            deduped[title_key] = {
                "title": signal.title,
                "latest_day": signal.day,
                "severity": signal.severity.lower(),
                "sources": {signal.source},
                "mentions": 1,
                "quality": source_quality,
            }
            continue

        entry["mentions"] = int(entry.get("mentions", 0)) + 1
        sources = entry.get("sources")
        if isinstance(sources, set):
            sources.add(signal.source)
        if signal.day > entry.get("latest_day"):  # type: ignore[arg-type]
            entry["latest_day"] = signal.day
            entry["title"] = signal.title
        if severity_rank.get(signal.severity.lower(), 1) > severity_rank.get(
            str(entry.get("severity", "info")).lower(), 1
        ):
            entry["severity"] = signal.severity.lower()
        if quality_rank.get(source_quality, 1) > quality_rank.get(
            str(entry.get("quality", "low")).lower(), 1
        ):
            entry["quality"] = source_quality

    rows: list[dict[str, object]] = []
    for payload in deduped.values():
        sources = payload.get("sources")
        source_count = len(sources) if isinstance(sources, set) else 0
        quality = str(payload.get("quality", "low")).lower()
        if source_count >= 3 and quality == "medium":
            quality = "high"
        rows.append(
            {
                "date": payload.get("latest_day"),
                "title": str(payload.get("title", "")).strip(),
                "severity": str(payload.get("severity", "info")).lower(),
                "quality": quality,
                "source_count": source_count,
                "mentions": int(payload.get("mentions", 0)),
            }
        )
    rows.sort(
        key=lambda item: (
            quality_rank.get(str(item.get("quality", "low")), 1),
            severity_rank.get(str(item.get("severity", "info")), 1),
            int(item.get("mentions", 0)),
        ),
        reverse=True,
    )
    return rows[: max(1, limit)]


def _news_category_impact_rows(
    external_signals: list[ExternalSignal],
    limit: int = 8,
) -> list[dict[str, object]]:
    category_map: dict[str, dict[str, object]] = {}
    severity_weight = {"high": 3.0, "medium": 2.0, "info": 1.0}
    for signal in external_signals:
        source_norm = _normalize_text(signal.source)
        if "news" not in source_norm and "campaign tracker" not in source_norm:
            continue
        blob = _normalize_text(f"{signal.title} {signal.details} {signal.source}")
        category = "Other market events"
        for label, tokens in NEWS_GMV_CATEGORY_RULES:
            if any(_normalize_text(token) in blob for token in tokens):
                category = label
                break
        row = category_map.setdefault(
            category,
            {
                "category": category,
                "signals": 0,
                "weighted_impact": 0.0,
                "latest_day": signal.day,
                "examples": [],
            },
        )
        row["signals"] = int(row.get("signals", 0)) + 1
        row["weighted_impact"] = float(row.get("weighted_impact", 0.0)) + severity_weight.get(
            signal.severity.lower(), 1.0
        )
        if signal.day > row.get("latest_day"):  # type: ignore[arg-type]
            row["latest_day"] = signal.day
        examples = row.get("examples", [])
        if isinstance(examples, list) and signal.title and signal.title not in examples:
            examples.append(signal.title)
            row["examples"] = examples[:3]

    rows = list(category_map.values())
    rows.sort(
        key=lambda item: (
            float(item.get("weighted_impact", 0.0)),
            int(item.get("signals", 0)),
        ),
        reverse=True,
    )
    return rows[: max(1, limit)]


def _trend_page_coverage_rows(
    scope_results: list[tuple[str, AnalysisResult]],
    additional_context: dict[str, object] | None,
    limit: int = 8,
) -> list[dict[str, object]]:
    product_trends = (additional_context or {}).get("product_trends", {})
    if not isinstance(product_trends, dict):
        return []

    current_rows = product_trends.get("current_non_brand", [])
    upcoming_rows = product_trends.get("upcoming_31d", [])
    if not isinstance(current_rows, list):
        current_rows = []
    if not isinstance(upcoming_rows, list):
        upcoming_rows = []

    candidates: list[dict[str, object]] = []
    seen_trends: set[str] = set()
    for row in current_rows + upcoming_rows:
        if not isinstance(row, dict):
            continue
        trend = str(row.get("trend", "")).strip()
        if not trend:
            continue
        trend_key = _normalize_text(trend)
        if trend_key in seen_trends:
            continue
        seen_trends.add(trend_key)
        candidates.append(row)
        if len(candidates) >= max(5, limit):
            break

    page_scope = _find_scope(scope_results, "page")
    page_index: dict[str, KeyDelta] = {}
    if page_scope is not None:
        for row in page_scope.top_winners + page_scope.top_losers:
            key = _normalized_page_key(row.key)
            if not key:
                continue
            existing = page_index.get(key)
            if existing is None or row.current_clicks > existing.current_clicks:
                page_index[key] = row

    out: list[dict[str, object]] = []
    for row in candidates:
        trend = str(row.get("trend", "")).strip()
        page = str(row.get("page", "")).strip()
        page_key = _normalized_page_key(page)
        page_perf = page_index.get(page_key) if page_key else None
        median_pos = _safe_float(row, "median_position")

        status = "Coverage present"
        gap_reason = "Landing page exists and has measurable demand."
        if not page:
            status = "Coverage gap"
            gap_reason = "Trend row has no mapped landing page URL."
        elif page_perf is None:
            status = "Coverage gap"
            gap_reason = "Landing page not visible among tracked top page movers."
        else:
            if page_perf.click_delta_vs_previous < -1000:
                status = "At-risk coverage"
                gap_reason = "Mapped landing page is losing clicks WoW."
            elif page_perf.current_position > 6.0:
                status = "Optimization gap"
                gap_reason = "Mapped landing page has ranking headroom (position > 6)."
            elif page_perf.current_ctr < 0.02 and page_perf.current_clicks > 500:
                status = "Optimization gap"
                gap_reason = "Mapped landing page has low CTR for current demand."

        out.append(
            {
                "trend": trend,
                "page": page or "-",
                "sheet_value": _safe_float(row, "value"),
                "sheet_date": str(row.get("date", "")).strip(),
                "median_position": median_pos,
                "status": status,
                "gap_reason": gap_reason,
                "page_delta_vs_prev": float(page_perf.click_delta_vs_previous) if page_perf else 0.0,
                "page_delta_vs_yoy": float(page_perf.click_delta_vs_yoy) if page_perf else 0.0,
                "page_ctr": float(page_perf.current_ctr) if page_perf else 0.0,
                "page_position": float(page_perf.current_position) if page_perf else 0.0,
            }
        )
    return out[: max(1, limit)]


def _hypothesis_continuity_rows(
    hypotheses: list[dict[str, object]],
    additional_context: dict[str, object] | None,
    limit: int = 6,
) -> list[dict[str, str]]:
    historical = (additional_context or {}).get("historical_reports", {})
    if not isinstance(historical, dict) or not historical.get("enabled"):
        return []

    notes: list[str] = []
    recent_reports = historical.get("recent_reports", [])
    if isinstance(recent_reports, list):
        notes.extend(_collect_context_notes(recent_reports))
    yoy_report = historical.get("yoy_report", {})
    if isinstance(yoy_report, dict):
        notes.extend(_collect_context_notes([yoy_report], fallback_key="excerpt"))

    notes_blob = _normalize_text(" ".join(notes))
    out: list[dict[str, str]] = []
    for row in hypotheses[: max(1, limit)]:
        category = str(row.get("category", "")).strip() or "Unclassified"
        thesis = str(row.get("thesis", "")).strip()
        category_key = _normalize_text(category)
        matched_tokens: tuple[str, ...] = ()
        for label, tokens in CONTINUITY_THEME_TOKENS:
            if _normalize_text(label) in category_key:
                matched_tokens = tokens
                break
        if not matched_tokens:
            matched_tokens = tuple(
                token for token in re.findall(r"[a-z]{4,}", _normalize_text(thesis))[:4]
            )

        if not notes:
            status = "No prior-report context"
            evidence = "No earlier weekly reports with extractable highlights."
        elif matched_tokens and any(token in notes_blob for token in matched_tokens):
            status = "Recurring vs prior report"
            hit = next((token for token in matched_tokens if token in notes_blob), "")
            evidence = f"Theme overlap found in prior reports (token: `{hit}`)."
        else:
            status = "New this week"
            evidence = "No direct overlap with extracted hypotheses/themes from prior reports."
        out.append(
            {
                "category": category,
                "status": status,
                "evidence": evidence,
            }
        )
    return out


def _follow_up_flags(
    totals: dict[str, MetricSummary],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
) -> list[str]:
    flags: list[str] = []
    current = totals["current_28d"]
    yoy = totals["yoy_52w"]

    brand_rows = (segment_diagnostics or {}).get("brand_non_brand") or []
    non_brand_row = next(
        (row for row in brand_rows if str(row.get("segment", "")).strip() == "non_brand"),
        None,
    )
    if isinstance(non_brand_row, dict):
        yoy_drop = float(non_brand_row.get("delta_pct_vs_yoy", 0.0))
        ctr_drop_pp = (current.ctr - yoy.ctr) * 100.0
        if yoy_drop <= -0.15 and ctr_drop_pp <= -0.30:
            flags.append(
                "Non-brand pressure flag: non-brand clicks are down "
                f"{yoy_drop * 100:.2f}% YoY and site CTR is down {ctr_drop_pp:.2f} pp YoY. Further analysis recommended."
            )

    page_rows = (segment_diagnostics or {}).get("page_template") or []
    page_loser = _top_negative_segment(page_rows, min_clicks=3000.0)
    if isinstance(page_loser, dict):
        pct_drop = float(page_loser.get("delta_pct_vs_previous", 0.0)) * 100.0
        if pct_drop <= -20.0:
            flags.append(
                "Page Name concentration flag: "
                f"`{page_loser.get('segment', '')}` is down {pct_drop:.2f}% WoW. Run section-level diagnosis."
            )

    seo_update_rows = [
        row
        for row in external_signals
        if "seo update analysis" in _normalize_text(row.source)
        and "risk" in _normalize_text(row.details)
    ]
    if seo_update_rows:
        latest = max(seo_update_rows, key=lambda row: row.day)
        flags.append(
            "Algorithm-volatility flag: external SEO media indicate elevated risk "
            f"({latest.day.isoformat()}: `{latest.title}`). Track affected query/page clusters."
        )

    forecast_avg = float(weather_summary.get("forecast_avg_temp_c", 0.0))
    current_avg = float(weather_summary.get("avg_temp_current_c", 0.0))
    if abs(forecast_avg - current_avg) >= 3.0:
        flags.append(
            "Weather-shift flag: next-week average temperature differs by "
            f"{forecast_avg - current_avg:+.1f}C vs current week. Monitor season-sensitive demand timing."
        )
    return flags[:5]


def _fallback_non_brand_yoy_from_gsc(
    query_scope: AnalysisResult | None,
    top_rows: int,
) -> list[dict[str, object]]:
    if query_scope is None:
        return []
    rows = query_scope.top_winners + query_scope.top_losers
    deduped: dict[str, dict[str, object]] = {}
    for row in rows:
        normalized_key = _normalize_text(row.key)
        if not normalized_key:
            continue
        if _is_noise_or_brand_query(row.key):
            continue
        if row.current_clicks <= 0 and row.yoy_clicks <= 0:
            continue
        delta = row.current_clicks - row.yoy_clicks
        if abs(delta) < 100:
            continue
        key = normalized_key
        payload = {
            "trend": row.key,
            "current_value": float(row.current_clicks),
            "previous_value": float(row.yoy_clicks),
            "delta_value": float(delta),
            "delta_pct": float(row.click_delta_pct_vs_yoy * 100.0),
            "sheet": "GSC query fallback",
        }
        current = deduped.get(key)
        if current is None or abs(float(payload["delta_value"])) > abs(float(current["delta_value"])):
            deduped[key] = payload

    ranked = sorted(
        deduped.values(),
        key=lambda item: abs(float(item.get("delta_value", 0.0))),
        reverse=True,
    )
    return ranked[: max(1, top_rows)]


def _sum_losses(rows: list[KeyDelta], keywords: tuple[str, ...]) -> float:
    total = 0.0
    for row in rows:
        if row.click_delta_vs_previous < 0 and _contains_any(row.key, keywords):
            total += -row.click_delta_vs_previous
    return total


def _sum_gains(rows: list[KeyDelta], keywords: tuple[str, ...]) -> float:
    total = 0.0
    for row in rows:
        if row.click_delta_vs_previous > 0 and _contains_any(row.key, keywords):
            total += row.click_delta_vs_previous
    return total


def _top_negative_segment(
    rows: list[dict[str, float | str]] | None,
    min_clicks: float = 1000.0,
) -> dict[str, float | str] | None:
    if not rows:
        return None
    candidates = [
        row
        for row in rows
        if float(row.get("current_clicks", 0.0)) >= min_clicks
        and float(row.get("delta_vs_previous", 0.0)) < 0
    ]
    if not candidates:
        return None
    return sorted(candidates, key=lambda row: float(row.get("delta_vs_previous", 0.0)))[0]


def _safe_float(row: dict[str, object], key: str, default: float = 0.0) -> float:
    try:
        return float(row.get(key, default))
    except (TypeError, ValueError):
        return default


def _position_delta_label(delta: float) -> str:
    if delta > 0:
        return f"worsened by {abs(delta):.2f}"
    if delta < 0:
        return f"improved by {abs(delta):.2f}"
    return "unchanged (0.00)"


def _sanitize_data_gap_message(message: object) -> str:
    text = str(message or "").strip()
    if not text:
        return ""
    text = re.sub(r"https?://\S+", "[url redacted]", text)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > 220:
        text = text[:217].rstrip() + "..."
    return text


def _format_count_yoy_change(current: float, yoy: float) -> str:
    current_n = int(round(float(current)))
    yoy_n = int(round(float(yoy)))
    if yoy_n <= 0:
        if current_n <= 0:
            return f"{current_n} vs {yoy_n} (no activity in both periods)."
        return f"{current_n} vs {yoy_n} (new activity vs no YoY baseline; % change n/a)."
    delta_pct = ((current_n - yoy_n) / yoy_n) * 100.0
    return f"{current_n} vs {yoy_n} ({delta_pct:+.1f}%)."


def _confidence_bucket(value: int | float | None) -> str:
    try:
        score = int(float(value or 0))
    except (TypeError, ValueError):
        score = 0
    if score >= 75:
        return "High"
    if score >= 60:
        return "Medium"
    return "Low"


def _hypothesis_impact_score(row: dict[str, object]) -> int:
    category = _normalize_text(str(row.get("category", "")))
    if any(token in category for token in ("technical", "seo visibility", "algorithm", "page name", "demand mix")):
        return 5
    if any(token in category for token in ("campaign", "seasonality", "events", "non-brand", "competitive")):
        return 4
    if any(token in category for token in ("macro", "weather", "continuity", "internal")):
        return 3
    if any(token in category for token in ("data quality",)):
        return 2
    return 3


def _hypothesis_controllability_score(row: dict[str, object]) -> int:
    owner = _normalize_text(str(row.get("owner", "")))
    category = _normalize_text(str(row.get("category", "")))
    if any(token in owner for token in ("seo ops", "web performance", "seo + product")):
        return 5
    if any(token in category for token in ("technical", "page name", "execution continuity", "internal initiatives")):
        return 4
    if any(token in category for token in ("campaign", "demand mix", "non-brand", "events", "seasonality")):
        return 3
    if any(token in category for token in ("macro", "weather", "algorithm", "serp behavior")):
        return 2
    return 3


def _apply_driver_priority_model(hypotheses: list[dict[str, object]]) -> list[dict[str, object]]:
    for idx, row in enumerate(hypotheses):
        if not isinstance(row, dict):
            continue
        confidence = int(row.get("confidence", 0) or 0)
        impact = _hypothesis_impact_score(row)
        controllability = _hypothesis_controllability_score(row)
        score = int(round((impact / 5.0) * (controllability / 5.0) * (max(0, min(100, confidence)) / 100.0) * 100.0))
        score = max(1, min(100, score))
        row["impact_score"] = impact
        row["controllability_score"] = controllability
        row["driver_priority_score"] = score
        row["priority_rank_hint"] = idx + 1
    hypotheses.sort(
        key=lambda row: (
            1 if bool(row.get("supporting_context_only")) else 0,
            -int(row.get("driver_priority_score", 0) or 0),
            -int(row.get("confidence", 0) or 0),
        )
    )
    return hypotheses


def _top_hypotheses_for_actions(hypotheses: list[dict[str, object]], limit: int = 3) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for row in hypotheses:
        if not isinstance(row, dict):
            continue
        if bool(row.get("supporting_context_only")):
            continue
        rows.append(row)
    if not rows:
        rows = [row for row in hypotheses if isinstance(row, dict)]
    rows.sort(
        key=lambda row: (
            -int(row.get("driver_priority_score", 0) or 0),
            -int(row.get("confidence", 0) or 0),
        )
    )
    return rows[: max(1, limit)]


def _uncertainty_action_template(confidence: int) -> str:
    if confidence >= 80:
        return "Likely driver; execute action now and verify with next-week checkpoint."
    if confidence >= 65:
        return "Plausible driver; validate before major escalation."
    return "Low-certainty signal; monitor and collect more evidence before action."


def _next_week_validation_plan_lines(hypotheses: list[dict[str, object]], limit: int = 3) -> list[str]:
    rows = _top_hypotheses_for_actions(hypotheses, limit=limit)
    lines: list[str] = []
    for row in rows:
        category = str(row.get("category", "Unknown")).strip() or "Unknown"
        metric = str(row.get("validation_metric", "")).strip() or "segment-level KPI check"
        date_label = str(row.get("validation_date", "")).strip() or (date.today() + timedelta(days=7)).isoformat()
        owner = str(row.get("owner", "")).strip() or "SEO Team"
        priority = int(row.get("driver_priority_score", 0) or 0)
        confidence = int(row.get("confidence", 0) or 0)
        lines.append(
            f"- `{category}` [{priority}/100 priority; {_confidence_bucket(confidence)} confidence]: "
            f"validate `{metric}` by {date_label} (owner: {owner})."
        )
    return lines[: max(1, limit)]


def _counterfactual_check_lines(hypotheses: list[dict[str, object]], limit: int = 3) -> list[str]:
    rows = _top_hypotheses_for_actions(hypotheses, limit=limit)
    out: list[str] = []
    for row in rows:
        category = str(row.get("category", "Unknown")).strip() or "Unknown"
        falsifier = str(row.get("falsifier", "")).strip() or "No falsifier provided."
        out.append(f"- `{category}`: {falsifier}")
    return out[: max(1, limit)]


def _causality_guardrail_summary(hypotheses: list[dict[str, object]]) -> str:
    if not hypotheses:
        return ""
    correlation_only = [
        row for row in hypotheses if isinstance(row, dict) and str(row.get("causality_level", "")).strip() == "correlation-only"
    ]
    if not correlation_only:
        return (
            "Causality guardrail: top hypotheses are triangulated across primary KPI evidence and context signals; "
            "continue validation before hard root-cause lock."
        )
    categories = ", ".join(
        f"`{str(row.get('category', '')).strip()}`"
        for row in correlation_only[:3]
        if str(row.get("category", "")).strip()
    )
    return (
        "Causality guardrail: correlation-only context cannot be treated as standalone root cause "
        "or trigger technical escalation. "
        + (f"Flagged: {categories}." if categories else "")
    ).strip()


def _contradiction_reconciliation_lines(
    *,
    totals: dict[str, MetricSummary],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    contradictions: list[str] = []

    clicks_delta = float(current.clicks - previous.clicks)
    impressions_delta = float(current.impressions - previous.impressions)
    ctr_delta_pp = (float(current.ctr) - float(previous.ctr)) * 100.0
    position_delta = float(current.position - previous.position)

    template_rows = (segment_diagnostics or {}).get("page_template") or []
    if isinstance(template_rows, list):
        home_row = next(
            (row for row in template_rows if str(row.get("segment", "")).strip().lower() == "home"),
            None,
        )
        if isinstance(home_row, dict):
            home_delta = float(home_row.get("delta_vs_previous", 0.0) or 0.0)
            if clicks_delta > 0 and home_delta < 0:
                contradictions.append(
                    "Total clicks are up while `home` is down; reconcile as routing mix shift "
                    "between templates/SERP features rather than a sitewide demand collapse."
                )

    brand_context = (additional_context or {}).get("google_trends_brand", {})
    brand_summary = brand_context.get("summary", {}) if isinstance(brand_context, dict) else {}
    brand_proxy = _brand_proxy_from_gsc(segment_diagnostics)
    if isinstance(brand_summary, dict) and isinstance(brand_proxy, dict) and brand_context.get("enabled"):
        trends_wow = float(brand_summary.get("delta_pct_vs_previous", 0.0) or 0.0)
        brand_clicks_wow = float(brand_proxy.get("delta_pct_vs_previous", 0.0) or 0.0)
        if trends_wow > 0.2 and brand_clicks_wow < -0.2:
            contradictions.append(
                "Brand interest is up in Google Trends but brand organic clicks are down in GSC; "
                "reconcile as possible paid-overlap/SERP allocation shift before technical SEO diagnosis."
            )

    if clicks_delta < 0 and impressions_delta >= 0 and ctr_delta_pp < 0:
        contradictions.append(
            "Impressions are stable/up while clicks decline with weaker CTR; "
            "reconcile as visibility-to-click efficiency issue (SERP mix/competition), not pure demand drop."
        )
    if clicks_delta < 0 and position_delta <= 0 and ctr_delta_pp < 0:
        contradictions.append(
            "Average position improved/stayed flat while clicks fell; "
            "reconcile via SERP layout and feature-share shift rather than ranking-only explanation."
        )
    return contradictions[:3]


def _technical_seo_escalation_gate(
    *,
    totals: dict[str, MetricSummary],
    hypotheses: list[dict[str, object]],
    contradiction_count: int = 0,
) -> dict[str, object]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    wow_clicks_pct = _ratio_delta(current.clicks, previous.clicks)
    ctr_wow_pp = (float(current.ctr) - float(previous.ctr)) * 100.0
    pos_wow = float(current.position - previous.position)

    technical_like = [
        row
        for row in hypotheses
        if isinstance(row, dict)
        and not bool(row.get("supporting_context_only"))
        and any(
            token in _normalize_text(str(row.get("category", "")))
            for token in ("technical", "seo visibility", "algorithm", "page name")
        )
    ]
    top_technical_conf = max((int(row.get("confidence", 0) or 0) for row in technical_like), default=0)

    severe_efficiency_deterioration = wow_clicks_pct <= -0.08 and ctr_wow_pp <= -0.15 and pos_wow >= 0.15
    moderate_efficiency_deterioration = wow_clicks_pct <= -0.04 and ctr_wow_pp <= -0.08 and pos_wow >= 0.05

    if severe_efficiency_deterioration and top_technical_conf >= 70:
        return {
            "status": "Escalate now",
            "reason": (
                "Material WoW efficiency deterioration with high-confidence technical/algorithm hypothesis."
            ),
            "next_action": "Trigger technical SEO investigation in 24-48h (templates, indexing, internal linking, CWV).",
        }
    if severe_efficiency_deterioration or (moderate_efficiency_deterioration and top_technical_conf >= 65):
        return {
            "status": "Escalate if persists",
            "reason": "Efficiency deterioration is visible but causality is not yet fully locked.",
            "next_action": "Run one strict validation cycle next week; escalate immediately if deterioration persists.",
        }
    if contradiction_count > 0:
        return {
            "status": "Hold escalation",
            "reason": "Mixed directional signals require reconciliation before technical escalation.",
            "next_action": "Prioritize reconciliation checks (routing/SERP mix/paid overlap) in the next run.",
        }
    return {
        "status": "No technical escalation",
        "reason": "Current movement is better explained by demand/timing/allocation context.",
        "next_action": "Continue monitoring with hypothesis validation plan; escalate only if efficiency weakens.",
    }


def _comparability_summary(
    additional_context: dict[str, object] | None,
    *,
    brand_trends_available: bool,
    brand_proxy_available: bool,
) -> str:
    parts: list[str] = []
    coverage = (additional_context or {}).get("gsc_data_coverage", {})
    if isinstance(coverage, dict) and coverage.get("p52w_mode") == "masked_to_days_with_data":
        days_with_data = int(coverage.get("days_with_data", 0) or 0)
        days_total = int(coverage.get("days_total", 0) or 0)
        if days_total > 0:
            parts.append(f"GSC YoY is aligned only to available current-week days ({days_with_data}/{days_total}).")
    if brand_trends_available:
        parts.append("Google Trends is an index (0-100), so use it directionally, not as absolute demand volume.")
    if brand_proxy_available:
        parts.append("GSC totals can differ by chosen dimensions in UI/API breakdowns.")
    return " ".join(parts).strip()


def _source_quality_summary(additional_context: dict[str, object] | None) -> str:
    rows = _source_freshness_rows(additional_context)
    reliability = (additional_context or {}).get("external_source_reliability", {})
    reliability_text = ""
    if isinstance(reliability, dict) and reliability.get("enabled"):
        summary = reliability.get("summary", {})
        if isinstance(summary, dict):
            weighted_score = float(summary.get("weighted_score", 0.0) or 0.0)
            sources_count = int(summary.get("sources_count", 0) or 0)
            tiers = summary.get("tier_counts", {})
            official = 0
            high_quality = 0
            if isinstance(tiers, dict):
                official = int(tiers.get("official", 0) or 0) + int(tiers.get("public-institution", 0) or 0)
                high_quality = int(tiers.get("high-quality media", 0) or 0)
            reliability_text = (
                "Source reliability score "
                f"{weighted_score:.1f}/100 across {sources_count} sources "
                f"(official/public={official}, high-quality media={high_quality})."
            )
    if not rows:
        return reliability_text
    fresh = sum(1 for row in rows if str(row.get("status", "")).strip().lower() == "fresh")
    stale = sum(1 for row in rows if str(row.get("status", "")).strip().lower() == "stale")
    degraded = sum(1 for row in rows if str(row.get("status", "")).strip().lower() == "degraded")
    total = len(rows)
    if stale == 0 and degraded == 0:
        freshness = f"Data quality check: all core sources are fresh ({fresh}/{total})."
    else:
        freshness = (
            "Data quality check: some sources are delayed or degraded "
            f"(fresh={fresh}, stale={stale}, degraded={degraded}; total={total})."
        )
    if reliability_text:
        return freshness + " " + reliability_text
    return freshness


QUALITY_GUARDRAIL_DROP_PREFIXES = (
    "macro context:",
    "macro backdrop (annual):",
    "labor-market backdrop:",
    "competitor promo radar:",
    "broader 28d context:",
    "hypothesis continuity:",
    "seo/geo publication context:",
    "trade-plan yoy hypotheses (channels):",
    "trade-plan yoy hypotheses (campaigns active this week):",
    "trade-plan yoy hypotheses (upcoming campaigns, not active this week):",
    "gsc signals suggest visibility shifted across serp result types",
    "campaign context:",
    "weekly market storyline on one timeline:",
    "platform/regulatory context:",
    "**forward 7d**:",
)


def _word_count_simple(text: str) -> int:
    return len(re.findall(r"\b\w+\b", str(text or ""), flags=re.UNICODE))


def enforce_manager_quality_guardrail(
    report_text: str,
    *,
    max_words: int = 1380,
) -> str:
    text = str(report_text or "").strip()
    if not text:
        return ""

    lines = text.splitlines()
    lowered = text.lower()

    has_confirmed_split = (
        "confirmed vs hypothesis" in lowered
        or "confirmed facts vs plausible drivers vs open questions" in lowered
    )
    if not has_confirmed_split:
        lines.append("")
        lines.append("## Confirmed Facts Vs Plausible Drivers Vs Open Questions")
        lines.append("### Confirmed facts from data")
        lines.append("- Weekly KPI movement is grounded in evidence anchors from this run [E1].")
        lines.append("### Plausible drivers (need validation)")
        lines.append("- Causal interpretation remains provisional and is validated in the next run [E2].")
        lines.append("### Open questions for next run")
        lines.append("- Which hypothesis should be falsified first before escalation? [E3]")

    lowered = "\n".join(lines).lower()
    if "priority actions" not in lowered or "owner | eta" not in lowered:
        lines.append("- **Priority actions (owner | ETA)**: [SEO Team | next run] Validate the top hypothesis against refreshed evidence.")
    marker_tokens = ("falsifier", "validation metric", "validation date")
    if not all(token in lowered for token in marker_tokens):
        marker_line = "- Hypothesis fields: falsifier | validation metric | validation date."
        insert_at = next(
            (
                idx + 1
                for idx, row in enumerate(lines)
                if row.strip().lower() == "## hypothesis protocol"
            ),
            -1,
        )
        if insert_at >= 0:
            lines.insert(insert_at, marker_line)
        else:
            lines.append(marker_line)

    joined_for_refs = "\n".join(lines)
    evidence_ref_count = len(re.findall(r"\[E\d+\]", joined_for_refs))
    if evidence_ref_count < 5:
        lines.append("- Evidence anchors reminder: [E1], [E2], [E3], [E4], [E5].")

    def _drop_candidate(line: str) -> bool:
        row = line.strip().lower()
        return any(row.startswith(prefix) for prefix in QUALITY_GUARDRAIL_DROP_PREFIXES)

    word_count = _word_count_simple("\n".join(lines))
    if word_count > max_words:
        pruned: list[str] = []
        for row in lines:
            if word_count > max_words and _drop_candidate(row):
                word_count -= _word_count_simple(row)
                continue
            pruned.append(row)
        lines = pruned

    word_count = _word_count_simple("\n".join(lines))
    if word_count > max_words:
        compacted: list[str] = []
        in_narrative = False
        keep_tokens = (
            "wow diagnosis",
            "yoy diagnosis",
            "daily trend view",
            "confirmed vs hypothesis",
            "reasoning ledger",
        )
        for row in lines:
            stripped = row.strip()
            lowered_row = stripped.lower()
            if lowered_row == "## what is happening and why":
                in_narrative = True
                compacted.append(row)
                continue
            if stripped.startswith("## ") and lowered_row != "## what is happening and why":
                in_narrative = False
            if (
                word_count > max_words
                and in_narrative
                and stripped.startswith("- ")
                and not any(token in lowered_row for token in keep_tokens)
            ):
                word_count -= _word_count_simple(row)
                continue
            compacted.append(row)
        lines = compacted

    word_count = _word_count_simple("\n".join(lines))
    if word_count > max_words:
        hard_trimmed: list[str] = []
        current_words = 0
        for row in lines:
            row_words = _word_count_simple(row)
            stripped = row.strip()
            if stripped.startswith("## ") or stripped.startswith("### "):
                hard_trimmed.append(row)
                current_words += row_words
                continue
            if (
                stripped.startswith("- Hypothesis fields:")
                or "priority actions (owner | eta)" in stripped.lower()
                or "evidence anchors reminder" in stripped.lower()
            ):
                hard_trimmed.append(row)
                current_words += row_words
                continue
            if current_words + row_words > max_words and stripped and not stripped.startswith("|"):
                continue
            hard_trimmed.append(row)
            current_words += row_words
        lines = hard_trimmed

    final_joined = "\n".join(lines)
    final_evidence_refs = len(re.findall(r"\[E\d+\]", final_joined))
    if final_evidence_refs < 5:
        lines.append("- Evidence anchors reminder: [E1], [E2], [E3], [E4], [E5].")

    return "\n".join(lines).strip() + "\n"


def _parse_iso_day_safe(value: object) -> date | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return date.fromisoformat(text[:10])
    except ValueError:
        return None


def _weather_daily_map(
    weather_summary: dict[str, float],
    key: str,
) -> dict[date, tuple[float, float]]:
    payload = weather_summary.get(key, [])
    if not isinstance(payload, list):
        return {}
    out: dict[date, tuple[float, float]] = {}
    for row in payload:
        if not isinstance(row, dict):
            continue
        day = _parse_iso_day_safe(row.get("date"))
        if day is None:
            continue
        try:
            temp_c = float(row.get("temp_c", 0.0) or 0.0)
            precip_mm = float(row.get("precip_mm", 0.0) or 0.0)
        except (TypeError, ValueError):
            continue
        out[day] = (temp_c, precip_mm)
    return out


def _signal_context_label(signal: ExternalSignal) -> str:
    blob = _normalize_text(f"{signal.source} {signal.title} {signal.details}")
    if "source degraded" in blob:
        return ""
    if (
        "campaign tracker" in blob
        or "trade plan" in blob
        or "planned campaign" in blob
        or any(token in blob for token in CAMPAIGN_EVENT_TOKENS)
    ):
        return f"campaign signal: {signal.title}"
    if any(token in blob for token in ("market events api", "public holidays", "platform+regulatory", "regulatory", "tax", "vat")):
        return f"event signal: {signal.title}"
    if any(token in blob for token in ("weekly seo digest", "news", "search status", "search central blog", "searchenginejournal", "seroundtable")):
        return f"news signal: {signal.title}"
    if "weather" in blob:
        return f"weather signal: {signal.title}"
    return ""


def _trade_plan_tags_for_day(
    additional_context: dict[str, object] | None,
    day: date,
    limit: int = 2,
) -> list[str]:
    trade_plan = (additional_context or {}).get("trade_plan", {})
    if not isinstance(trade_plan, dict) or not trade_plan.get("enabled"):
        return []
    campaign_rows = trade_plan.get("campaign_rows", [])
    if not isinstance(campaign_rows, list):
        return []
    tags: list[str] = []
    for row in campaign_rows:
        if not isinstance(row, dict):
            continue
        campaign = str(row.get("campaign", "")).strip()
        if not campaign:
            continue
        first_day = _parse_iso_day_safe(row.get("first_date"))
        last_day = _parse_iso_day_safe(row.get("last_date"))
        if first_day is None or last_day is None:
            continue
        if first_day <= day <= last_day:
            tags.append(f"trade-plan campaign active: {campaign}")
        if len(tags) >= max(1, limit):
            break
    return tags


def _daily_weather_impact_hint(temp_diff: float, precip_diff: float) -> tuple[str, int] | None:
    if abs(temp_diff) < 2.0 and abs(precip_diff) < 4.0:
        return None

    if temp_diff <= -2.0:
        temp_state = "colder"
    elif temp_diff >= 2.0:
        temp_state = "warmer"
    else:
        temp_state = "temperature-stable"

    if precip_diff >= 4.0:
        precip_state = "wetter"
    elif precip_diff <= -4.0:
        precip_state = "drier"
    else:
        precip_state = "precipitation-stable"

    if temp_state == "temperature-stable" and precip_state == "precipitation-stable":
        return None

    if temp_state != "temperature-stable" and precip_state != "precipitation-stable":
        weather_state = f"{temp_state} and {precip_state}"
    elif temp_state != "temperature-stable":
        weather_state = temp_state
    else:
        weather_state = precip_state

    if temp_diff <= -2.0 and precip_diff >= 4.0:
        likely_effect = "likely lower discretionary outdoor demand and stronger indoor/home demand rotation"
    elif temp_diff >= 2.0 and precip_diff <= -4.0:
        likely_effect = "likely higher discretionary/outdoor demand and softer weather-driven indoor demand"
    elif precip_diff >= 4.0:
        likely_effect = "likely demand reallocation due to heavier precipitation"
    elif precip_diff <= -4.0:
        likely_effect = "likely demand reallocation due to lighter precipitation"
    elif temp_diff <= -2.0:
        likely_effect = "likely colder-weather demand rotation"
    else:
        likely_effect = "likely warmer-weather demand rotation"

    score = int(round(52.0 + abs(temp_diff) * 2.5 + abs(precip_diff) * 1.2))
    if temp_state != "temperature-stable" and precip_state != "precipitation-stable":
        score += 4
    score = max(52, min(82, score))

    label = (
        "weather timing hint: "
        f"{weather_state} vs previous-weekday; {likely_effect} "
        f"(non-causal, {_confidence_bucket(score)} confidence: {score}/100)"
    )
    return label, score


def _daily_context_tags_for_day(
    *,
    day: date,
    external_signals: list[ExternalSignal],
    additional_context: dict[str, object] | None,
    weather_summary: dict[str, float],
    limit: int = 3,
) -> list[str]:
    scored: list[tuple[float, str]] = []
    for signal in external_signals:
        lag_days = abs((signal.day - day).days)
        if lag_days > 1:
            continue
        label = _signal_context_label(signal)
        if not label:
            continue
        severity_bonus = {"high": 2.0, "medium": 1.0, "info": 0.5}.get(signal.severity.lower(), 0.5)
        if label.startswith("campaign signal:"):
            category_bonus = 4.0
        elif label.startswith("event signal:"):
            category_bonus = 3.0
        elif label.startswith("news signal:"):
            category_bonus = 2.0
        else:
            category_bonus = 1.0
        score = (category_bonus * 10.0) + (severity_bonus * 2.0) - float(lag_days)
        scored.append((score, label))

    tags: list[str] = []
    seen: set[str] = set()
    for _, label in sorted(scored, key=lambda item: item[0], reverse=True):
        canonical = _normalize_text(label)
        if not canonical or canonical in seen:
            continue
        seen.add(canonical)
        tags.append(label)
        if len(tags) >= max(1, limit):
            break

    for row in _trade_plan_tags_for_day(additional_context=additional_context, day=day, limit=2):
        canonical = _normalize_text(row)
        if canonical and canonical not in seen:
            seen.add(canonical)
            tags.append(row)
        if len(tags) >= max(1, limit):
            break

    current_weather = _weather_daily_map(weather_summary, "daily_current")
    previous_weather = _weather_daily_map(weather_summary, "daily_previous")
    current_row = current_weather.get(day)
    previous_row = previous_weather.get(day - timedelta(days=7))
    limit_safe = max(1, limit)
    if current_row and previous_row:
        temp_diff = float(current_row[0] - previous_row[0])
        precip_diff = float(current_row[1] - previous_row[1])
        weather_hint = _daily_weather_impact_hint(temp_diff=temp_diff, precip_diff=precip_diff)
        if weather_hint:
            weather_tag, _ = weather_hint
            weather_tag += f"; observed change: temp {temp_diff:+.1f}C, precipitation {precip_diff:+.1f}mm"
            canonical = _normalize_text(weather_tag)
            if canonical and canonical not in seen:
                if len(tags) >= limit_safe:
                    tags = tags[: limit_safe - 1] + [weather_tag]
                else:
                    tags.append(weather_tag)

    return tags[:limit_safe]


def _build_daily_gsc_storyline(
    *,
    additional_context: dict[str, object] | None,
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    top_n: int = 3,
) -> dict[str, object]:
    context = (additional_context or {}).get("gsc_daily_rows", {})
    if not isinstance(context, dict) or not context.get("enabled"):
        return {"enabled": False, "executive_line": "", "narrative_lines": []}

    raw_rows = context.get("rows", [])
    if not isinstance(raw_rows, list) or not raw_rows:
        return {"enabled": False, "executive_line": "", "narrative_lines": []}

    weekly_clicks = float(context.get("weekly_clicks_sum", 0.0) or 0.0)
    material_click_threshold = max(500.0, weekly_clicks * 0.03) if weekly_clicks > 0.0 else 500.0

    candidate_rows: list[dict[str, object]] = []
    for row in raw_rows:
        if not isinstance(row, dict):
            continue
        if not bool(row.get("previous_weekday_has_data")):
            continue
        candidate_rows.append(row)
    if not candidate_rows:
        return {"enabled": False, "executive_line": "", "narrative_lines": []}

    material_rows = [
        row
        for row in candidate_rows
        if abs(float(row.get("delta_clicks_vs_previous_weekday", 0.0) or 0.0)) >= material_click_threshold
        or abs(float(row.get("delta_pct_vs_previous_weekday", 0.0) or 0.0)) >= 6.0
    ]
    ranked = material_rows if material_rows else candidate_rows
    ranked = sorted(
        ranked,
        key=lambda row: (
            abs(float(row.get("delta_clicks_vs_previous_weekday", 0.0) or 0.0)),
            abs(float(row.get("delta_pct_vs_previous_weekday", 0.0) or 0.0)),
        ),
        reverse=True,
    )[: max(1, top_n)]

    detail_snippets: list[str] = []
    narrative_lines: list[str] = [
        "**Daily trend view (GSC by day)**: day-level checks show which specific dates drove the weekly outcome; "
        "context tags (weather/news/campaigns) are directional hints, not proof of causality."
    ]
    context_tags_flat: list[str] = []

    for row in ranked:
        day = _parse_iso_day_safe(row.get("date"))
        if day is None:
            continue
        delta_clicks = float(row.get("delta_clicks_vs_previous_weekday", 0.0) or 0.0)
        delta_pct = float(row.get("delta_pct_vs_previous_weekday", 0.0) or 0.0)
        delta_yoy_pct = float(row.get("delta_pct_vs_yoy_day", 0.0) or 0.0)
        yoy_has_data = bool(row.get("yoy_day_has_data"))
        direction = "up" if delta_clicks > 0 else ("down" if delta_clicks < 0 else "flat")
        base = (
            f"{day.isoformat()} {direction} {_fmt_signed_compact(delta_clicks)} "
            f"({delta_pct:+.1f}% vs previous-weekday"
            + (f", {delta_yoy_pct:+.1f}% vs YoY weekday" if yoy_has_data else "")
            + ")"
        )
        detail_snippets.append(base)

        day_tags = _daily_context_tags_for_day(
            day=day,
            external_signals=external_signals,
            additional_context=additional_context,
            weather_summary=weather_summary,
            limit=3,
        )
        context_tags_flat.extend(day_tags)
        narrative_lines.append(
            "- "
            + base
            + (f". Co-occurring context: {'; '.join(day_tags)}." if day_tags else ".")
        )

    if not detail_snippets:
        return {"enabled": False, "executive_line": "", "narrative_lines": []}

    deduped_context: list[str] = []
    seen_ctx: set[str] = set()
    for row in context_tags_flat:
        canonical = _normalize_text(row)
        if not canonical or canonical in seen_ctx:
            continue
        seen_ctx.add(canonical)
        deduped_context.append(row)
        if len(deduped_context) >= 4:
            break

    executive_line = "- **Daily GSC pulse (day-by-day)**: strongest moves were " + "; ".join(detail_snippets[:2]) + "."
    if deduped_context:
        executive_line += " Co-occurring context: " + "; ".join(deduped_context[:3]) + "."

    days_with_data = int(context.get("days_with_data", 0) or 0)
    days_total = int(context.get("days_total", 0) or 0)
    days_with_prev = int(context.get("days_with_previous_weekday_data", 0) or 0)
    if days_total > 0:
        narrative_lines.append(
            "- Coverage: "
            f"{days_with_data}/{days_total} days had usable day-level GSC data; "
            f"weekday-aligned WoW comparisons were available for {days_with_prev} days."
        )

    return {
        "enabled": True,
        "executive_line": executive_line,
        "narrative_lines": narrative_lines[:6],
    }


def _compact_manager_section(lines: list[str], max_lines: int = 24) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    skip_prefixes = (
        "DuckDuckGo context scan:",
        "Channel pressure signal:",
        "Campaign YoY signal (trade plan):",
        "Trade-plan YoY hypotheses (channels):",
        "Trade-plan YoY hypotheses (campaigns active this week):",
        "Trade-plan YoY hypotheses (upcoming campaigns, not active this week):",
        "SEO/GEO publication context:",
    )

    def _norm(text: str) -> str:
        return re.sub(r"\s+", " ", re.sub(r"[`*#|]", "", text)).strip().lower()

    idx = 0
    while idx < len(lines):
        row = lines[idx]
        stripped = row.strip()
        if stripped.startswith("|"):
            idx += 1
            continue
        if any(stripped.startswith(prefix) for prefix in skip_prefixes):
            idx += 1
            continue
        key = _norm(row)
        if key and key in seen and not stripped.startswith("**"):
            idx += 1
            continue
        if key:
            seen.add(key)
        out.append(row)

        confirmed_split_headers = {
            "**Confirmed vs hypothesis**",
            "**Confirmed facts vs plausible drivers vs open questions**",
        }
        if stripped in confirmed_split_headers:
            look_ahead = idx + 1
            while look_ahead < len(lines):
                candidate = lines[look_ahead]
                cand_strip = candidate.strip()
                if cand_strip.startswith("**") and cand_strip not in confirmed_split_headers:
                    break
                if (
                    cand_strip.startswith("Confirmed facts from data:")
                    or cand_strip.startswith("Working hypotheses:")
                    or cand_strip.startswith("Plausible drivers (need validation):")
                    or cand_strip.startswith("Open questions for next run:")
                    or cand_strip.startswith("- ")
                ):
                    cand_key = _norm(candidate)
                    if cand_key not in seen:
                        seen.add(cand_key)
                        out.append(candidate)
                look_ahead += 1
                if len(out) >= max_lines:
                    break
            idx = look_ahead
            if len(out) >= max_lines:
                break
            continue

        if stripped == "**Reasoning ledger (facts -> hypotheses -> validation)**":
            look_ahead = idx + 1
            while look_ahead < len(lines):
                candidate = lines[look_ahead]
                cand_strip = candidate.strip()
                if cand_strip.startswith("**") and cand_strip != "**Reasoning ledger (facts -> hypotheses -> validation)**":
                    break
                if cand_strip.startswith("- "):
                    cand_key = _norm(candidate)
                    if cand_key not in seen:
                        seen.add(cand_key)
                        out.append(candidate)
                look_ahead += 1
                if len(out) >= max_lines:
                    break
            idx = look_ahead
            if len(out) >= max_lines:
                break
            continue

        idx += 1
        if len(out) >= max_lines:
            break

    return out[:max_lines]


def _source_freshness_rows(additional_context: dict[str, object] | None) -> list[dict[str, object]]:
    payload = (additional_context or {}).get("source_freshness", [])
    if not isinstance(payload, list):
        return []
    rows: list[dict[str, object]] = []
    for row in payload:
        if not isinstance(row, dict):
            continue
        rows.append(
            {
                "source": str(row.get("source", "")).strip(),
                "status": str(row.get("status", "")).strip(),
                "last_day": str(row.get("last_day", "")).strip(),
                "ttl_hours": float(row.get("ttl_hours", 0.0) or 0.0),
                "cache_mode": str(row.get("cache_mode", "")).strip(),
                "note": str(row.get("note", "")).strip(),
            }
        )
    return rows


def _build_evidence_ledger(
    *,
    windows: dict[str, DateWindow],
    external_signals: list[ExternalSignal],
    additional_context: dict[str, object] | None,
    weather_summary: dict[str, float] | None = None,
    limit: int = 12,
) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    current = windows.get("current_28d")
    previous = windows.get("previous_28d")
    if isinstance(current, DateWindow):
        rows.append(
            {
                "id": "E1",
                "source": "GSC weekly window",
                "date": f"{current.start.isoformat()}..{current.end.isoformat()}",
                "note": "Primary KPI baseline window.",
            }
        )
    analysis_start = previous.start if isinstance(previous, DateWindow) else (current.start if isinstance(current, DateWindow) else None)
    analysis_end = current.end if isinstance(current, DateWindow) else None
    forward_end = analysis_end + timedelta(days=14) if isinstance(analysis_end, date) else None
    filtered_signals: list[ExternalSignal] = []
    for signal in external_signals:
        if isinstance(analysis_start, date) and signal.day < analysis_start:
            continue
        if isinstance(forward_end, date) and signal.day > forward_end:
            continue
        filtered_signals.append(signal)
    for signal in filtered_signals[: max(1, limit - len(rows))]:
        note = str(signal.title).strip()[:140]
        if isinstance(analysis_end, date) and signal.day > analysis_end:
            note = f"Forward context: {note}"
        rows.append(
            {
                "id": f"E{len(rows) + 1}",
                "source": str(signal.source).strip(),
                "date": signal.day.isoformat(),
                "note": note,
            }
        )
    if isinstance(additional_context, dict):
        trade_plan = additional_context.get("trade_plan", {})
        if isinstance(trade_plan, dict) and trade_plan.get("enabled"):
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "Trade plan sheet",
                    "date": str(((trade_plan.get("windows", {}) or {}).get("current", {}) or {}).get("start", "")).strip(),
                    "note": str(((trade_plan.get("sheet", {}) or {}).get("tab", "")).strip() or "trade-plan context"),
                }
            )
        weekly_news = additional_context.get("weekly_news_digest", {})
        if isinstance(weekly_news, dict) and weekly_news.get("enabled"):
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "Weekly SEO/GEO digest",
                    "date": f"{weekly_news.get('window_start','')}..{weekly_news.get('window_end','')}",
                    "note": f"rows={len(weekly_news.get('rows', []) if isinstance(weekly_news.get('rows', []), list) else [])}",
                }
            )
        updates_timeline = additional_context.get("google_updates_timeline", {})
        if isinstance(updates_timeline, dict) and updates_timeline.get("enabled"):
            summary = updates_timeline.get("summary", {})
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "Google updates timeline (13M)",
                    "date": f"{updates_timeline.get('scan_start','')}..{updates_timeline.get('scan_end','')}",
                    "note": (
                        f"30d={int((summary or {}).get('count_current_30d', 0) or 0)}; "
                        f"latest={str((summary or {}).get('latest_update_date', '')).strip()}"
                        if isinstance(summary, dict)
                        else "timeline summary"
                    ),
                }
            )
        case_studies = additional_context.get("serp_case_studies", {})
        if isinstance(case_studies, dict) and case_studies.get("enabled"):
            summary = case_studies.get("summary", {})
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "SERP case-study scanner (13M)",
                    "date": f"{case_studies.get('scan_start','')}..{case_studies.get('scan_end','')}",
                    "note": (
                        f"total={int((summary or {}).get('total_count_13m', 0) or 0)}; "
                        f"latest={str((summary or {}).get('latest_case_date', '')).strip()}"
                        if isinstance(summary, dict)
                        else "case-study summary"
                    ),
                }
            )
        market_events = additional_context.get("market_event_calendar", {})
        if isinstance(market_events, dict) and market_events.get("enabled"):
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "Market event calendar",
                    "date": str(market_events.get("country_code", "")).strip() or "-",
                    "note": f"rows={len(market_events.get('events', []) if isinstance(market_events.get('events', []), list) else [])}",
                }
            )
        trends = additional_context.get("product_trends", {})
        if isinstance(trends, dict) and trends.get("enabled"):
            rows.append(
                {
                    "id": f"E{len(rows) + 1}",
                    "source": "Product trends sheets",
                    "date": str(trends.get("horizon_days", 31)),
                    "note": "YoY/current/upcoming non-brand trend context.",
                }
            )
    if isinstance(weather_summary, dict):
        rows.append(
            {
                "id": f"E{len(rows) + 1}",
                "source": "Weather context",
                "date": "-",
                "note": (
                    f"temp_diff={float(weather_summary.get('avg_temp_diff_c', 0.0) or 0.0):+.1f}C; "
                    f"precip_change={float(weather_summary.get('precip_change_pct', 0.0) or 0.0):+.1f}%"
                ),
            }
        )
    return rows[: max(1, limit)]


def _pick_evidence_anchor_id(
    evidence_ledger: list[dict[str, str]],
    *,
    source_tokens: tuple[str, ...],
) -> str:
    for row in evidence_ledger:
        source = _normalize_text(str(row.get("source", "")))
        note = _normalize_text(str(row.get("note", "")))
        blob = f"{source} {note}".strip()
        if any(token in blob for token in source_tokens):
            return str(row.get("id", "")).strip()
    return ""


def _claim_evidence_completeness_rows(
    *,
    executive_lines: list[str],
    narrative_lines: list[str],
    evidence_ledger: list[dict[str, str]],
) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    candidates = [str(line).strip() for line in (executive_lines + narrative_lines) if str(line).strip()]
    for line in candidates:
        normalized = _normalize_text(line)
        if any(token in normalized for token in ("what changed", "wow diagnosis", "yoy diagnosis", "facts observed")):
            rows.append({"claim": line, "anchor_id": "E1", "status": "mapped"})
            continue
        if any(token in normalized for token in ("campaign", "trade-plan", "marketplace timeline", "weekly market storyline")):
            anchor = _pick_evidence_anchor_id(
                evidence_ledger,
                source_tokens=("trade plan", "campaign tracker", "market event", "weekly seo/geo digest"),
            )
            rows.append({"claim": line, "anchor_id": anchor, "status": "mapped" if anchor else "missing"})
            continue
        if any(token in normalized for token in ("weather context", "weather timing hint", "forward 7d")):
            anchor = _pick_evidence_anchor_id(evidence_ledger, source_tokens=("weather",))
            rows.append({"claim": line, "anchor_id": anchor, "status": "mapped" if anchor else "missing"})
            continue
        if any(token in normalized for token in ("google updates timeline", "algorithm context")):
            anchor = _pick_evidence_anchor_id(evidence_ledger, source_tokens=("google updates timeline", "google search status", "google search central"))
            rows.append({"claim": line, "anchor_id": anchor, "status": "mapped" if anchor else "missing"})
            continue
        if any(token in normalized for token in ("case-study", "serp behavior context", "serp appearance")):
            anchor = _pick_evidence_anchor_id(evidence_ledger, source_tokens=("serp case-study scanner", "google updates timeline"))
            rows.append({"claim": line, "anchor_id": anchor, "status": "mapped" if anchor else "missing"})
            continue
        if any(token in normalized for token in ("brand demand", "brand search trend")):
            rows.append({"claim": line, "anchor_id": "E1", "status": "mapped"})
            continue
    # De-duplicate by normalized claim text.
    dedup: dict[str, dict[str, str]] = {}
    for row in rows:
        key = _normalize_text(str(row.get("claim", "")))
        if not key:
            continue
        existing = dedup.get(key)
        if existing is None:
            dedup[key] = row
            continue
        if existing.get("status") != "mapped" and row.get("status") == "mapped":
            dedup[key] = row
    return list(dedup.values())[:12]


def _governance_lines(
    *,
    run_date: date,
    report_country_code: str,
    windows: dict[str, DateWindow],
    additional_context: dict[str, object] | None,
) -> list[str]:
    lines: list[str] = []
    lines.append("## Governance and provenance")
    lines.append(f"- Report date: {run_date.isoformat()} | Country: {report_country_code.strip().upper() or 'PL'}.")
    model_name = str((additional_context or {}).get("gaia_model", "")).strip() if isinstance(additional_context, dict) else ""
    if model_name:
        lines.append(f"- Model version: {model_name}.")
    current = windows.get("current_28d")
    previous = windows.get("previous_28d")
    yoy = windows.get("yoy_52w")
    if isinstance(current, DateWindow) and isinstance(previous, DateWindow) and isinstance(yoy, DateWindow):
        lines.append(
            "- Data windows: "
            f"current {current.start.isoformat()}..{current.end.isoformat()}, "
            f"previous {previous.start.isoformat()}..{previous.end.isoformat()}, "
            f"YoY {yoy.start.isoformat()}..{yoy.end.isoformat()}."
        )
    reviewer = str((additional_context or {}).get("governance_human_reviewer", "")).strip() if isinstance(additional_context, dict) else ""
    lines.append(f"- Human reviewer: {reviewer or 'not assigned'}")
    lines.append("- Known limitations: external signals can be delayed/noisy; causal claims are probabilistic and require next-run validation.")
    return lines


def _impact_attribution_rows(
    *,
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    limit: int = 8,
) -> list[dict[str, object]]:
    total_delta = float(totals["current_28d"].clicks - totals["previous_28d"].clicks)
    denom = abs(total_delta) if abs(total_delta) > 0 else 1.0
    rows: list[dict[str, object]] = []

    query_scope = _find_scope(scope_results, "query")
    for row in _query_cluster_rows(query_scope):
        delta = float(row.get("delta_vs_previous", 0.0))
        if abs(delta) < 500:
            continue
        rows.append(
            {
                "dimension": "Query cluster",
                "name": str(row.get("cluster", "")).strip(),
                "delta": delta,
                "share_of_total": (delta / denom) * 100.0,
            }
        )

    for group_name, key in (
        ("Brand split", "brand_non_brand"),
        ("Device", "device"),
        ("Page Name", "page_template"),
    ):
        for row in (segment_diagnostics or {}).get(key, []) or []:
            if not isinstance(row, dict):
                continue
            delta = float(row.get("delta_vs_previous", 0.0))
            if abs(delta) < 500:
                continue
            rows.append(
                {
                    "dimension": group_name,
                    "name": str(row.get("segment", "")).strip() or "-",
                    "delta": delta,
                    "share_of_total": (delta / denom) * 100.0,
                }
            )

    rows.sort(key=lambda item: abs(float(item.get("delta", 0.0))), reverse=True)
    return rows[: max(1, limit)]


def _query_anomaly_rows(
    query_scope: AnalysisResult | None,
    limit: int = 6,
) -> list[dict[str, object]]:
    if query_scope is None:
        return []
    rows = query_scope.top_winners + query_scope.top_losers
    deltas = [float(row.click_delta_vs_previous) for row in rows if abs(float(row.click_delta_vs_previous)) >= 200]
    if len(deltas) < 5:
        return []
    sorted_abs = sorted(abs(value) for value in deltas)
    mid = len(sorted_abs) // 2
    median_abs = sorted_abs[mid]
    mad = sorted(abs(abs(value) - median_abs) for value in deltas)[mid]
    threshold = median_abs + 2.5 * max(1.0, mad)

    out: list[dict[str, object]] = []
    for row in rows:
        delta = float(row.click_delta_vs_previous)
        if abs(delta) < max(2000.0, threshold):
            continue
        out.append(
            {
                "query": str(row.key).strip(),
                "delta": delta,
                "delta_pct": float(row.click_delta_pct_vs_previous) * 100.0,
                "current_clicks": float(row.current_clicks),
            }
        )
    out.sort(key=lambda item: abs(float(item.get("delta", 0.0))), reverse=True)
    return out[: max(1, limit)]


def _data_quality_score(
    *,
    additional_context: dict[str, object] | None,
    external_signals: list[ExternalSignal],
    senuto_error: str | None,
) -> tuple[int, list[str]]:
    score = 100
    notes: list[str] = []
    ctx = additional_context or {}
    errors = ctx.get("errors", [])
    if isinstance(errors, list):
        penalty = min(25, 4 * len([item for item in errors if str(item).strip()]))
        score -= penalty
        if penalty:
            notes.append(f"{len(errors)} context-source errors detected.")

    freshness_rows = _source_freshness_rows(ctx)
    stale_count = sum(1 for row in freshness_rows if str(row.get("status", "")).lower() == "stale")
    degraded_count = sum(1 for row in freshness_rows if str(row.get("status", "")).lower() == "degraded")
    score -= min(20, stale_count * 5 + degraded_count * 8)
    if stale_count:
        notes.append(f"{stale_count} source(s) served from stale cache.")
    if degraded_count:
        notes.append(f"{degraded_count} source(s) degraded/no live data.")

    if senuto_error:
        score -= 10
        notes.append("Senuto unavailable in this run.")

    if not external_signals:
        score -= 8
        notes.append("No external signals available.")

    score = max(0, min(100, score))
    if not notes:
        notes.append("No material data-quality risks detected.")
    return score, notes[:4]


def _decision_one_pager_rows(hypotheses: list[dict[str, object]], limit: int = 6) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for row in hypotheses[: max(1, limit)]:
        thesis = str(row.get("thesis", "")).strip()
        category = str(row.get("category", "")).strip() or "Unknown"
        confidence = int(row.get("confidence", 0) or 0)
        evidence_rows = row.get("evidence", [])
        evidence = ""
        if isinstance(evidence_rows, list) and evidence_rows:
            evidence = str(evidence_rows[0]).strip()
        if not evidence:
            evidence = "Evidence available in supporting appendix."
        rows.append(
            {
                "what_changed": category,
                "why_it_matters": f"{thesis} (confidence {confidence}/100). Evidence: {evidence}",
            }
        )
    return rows


def _brand_proxy_from_gsc(
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
) -> dict[str, float] | None:
    rows = (segment_diagnostics or {}).get("brand_non_brand") or []
    if not isinstance(rows, list):
        return None
    brand_row = next(
        (
            row
            for row in rows
            if isinstance(row, dict)
            and str(row.get("segment", "")).strip().lower() == "brand"
        ),
        None,
    )
    if not isinstance(brand_row, dict):
        return None
    return {
        "current_clicks": float(brand_row.get("current_clicks", 0.0)),
        "delta_vs_previous": float(brand_row.get("delta_vs_previous", 0.0)),
        "delta_pct_vs_previous": float(brand_row.get("delta_pct_vs_previous", 0.0)) * 100.0,
        "delta_vs_yoy": float(brand_row.get("delta_vs_yoy", 0.0)),
        "delta_pct_vs_yoy": float(brand_row.get("delta_pct_vs_yoy", 0.0)) * 100.0,
        "current_impressions": float(brand_row.get("current_impressions", 0.0)),
        "impressions_delta_vs_previous": float(brand_row.get("impressions_delta_vs_previous", 0.0)),
        "impressions_delta_pct_vs_previous": float(
            brand_row.get("impressions_delta_pct_vs_previous", 0.0)
        )
        * 100.0,
        "impressions_delta_vs_yoy": float(brand_row.get("impressions_delta_vs_yoy", 0.0)),
        "impressions_delta_pct_vs_yoy": float(
            brand_row.get("impressions_delta_pct_vs_yoy", 0.0)
        )
        * 100.0,
    }


def _brand_ctr_proxy_from_gsc(
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
) -> dict[str, float] | None:
    rows = (segment_diagnostics or {}).get("brand_non_brand") or []
    if not isinstance(rows, list):
        return None
    brand_row = next(
        (
            row
            for row in rows
            if isinstance(row, dict)
            and str(row.get("segment", "")).strip().lower() == "brand"
        ),
        None,
    )
    if not isinstance(brand_row, dict):
        return None

    current_clicks = float(brand_row.get("current_clicks", 0.0))
    previous_clicks = float(brand_row.get("previous_clicks", 0.0))
    yoy_clicks = float(brand_row.get("yoy_clicks", 0.0))
    current_impr = float(brand_row.get("current_impressions", 0.0))
    previous_impr = float(brand_row.get("previous_impressions", 0.0))
    yoy_impr = float(brand_row.get("yoy_impressions", 0.0))

    current_ctr = (current_clicks / current_impr) if current_impr else 0.0
    previous_ctr = (previous_clicks / previous_impr) if previous_impr else 0.0
    yoy_ctr = (yoy_clicks / yoy_impr) if yoy_impr else 0.0
    return {
        "current_ctr": current_ctr,
        "previous_ctr": previous_ctr,
        "yoy_ctr": yoy_ctr,
        "delta_pp_vs_previous": (current_ctr - previous_ctr) * 100.0,
        "delta_pp_vs_yoy": (current_ctr - yoy_ctr) * 100.0,
    }


def _trade_plan_paid_pressure(
    additional_context: dict[str, object] | None,
) -> tuple[str, float]:
    trade_plan = (additional_context or {}).get("trade_plan", {})
    if not isinstance(trade_plan, dict) or not trade_plan.get("enabled"):
        return ("unknown", 0.0)
    rows = trade_plan.get("channel_split", [])
    if not isinstance(rows, list):
        return ("unknown", 0.0)
    paid_tokens = ("paid", "google ads", "ads", "cpc", "sem", "search")
    delta_sum = 0.0
    matched = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        channel = str(row.get("channel", "")).strip().lower()
        if not channel:
            continue
        if any(token in channel for token in paid_tokens):
            delta_sum += float(row.get("delta_spend", 0.0) or 0.0)
            matched += 1
    if matched == 0:
        return ("unknown", 0.0)
    if delta_sum > 0:
        return ("up", delta_sum)
    if delta_sum < 0:
        return ("down", delta_sum)
    return ("flat", 0.0)


def _trade_plan_overlap_intensity(
    trade_plan: dict[str, object],
) -> dict[str, object]:
    campaign_rows = trade_plan.get("campaign_rows", [])
    windows = trade_plan.get("windows", {})
    current_window = windows.get("current", {}) if isinstance(windows, dict) else {}
    current_start = _parse_iso_day_safe((current_window or {}).get("start"))
    current_end = _parse_iso_day_safe((current_window or {}).get("end"))
    if not isinstance(campaign_rows, list) or current_start is None or current_end is None:
        return {
            "score": 0,
            "campaigns": 0,
            "overlap_days": 0,
            "recency_days": None,
            "label": "Low",
        }

    overlap_campaigns = 0
    overlap_day_set: set[date] = set()
    latest_overlap_day: date | None = None
    for row in campaign_rows:
        if not isinstance(row, dict):
            continue
        first_day = _parse_iso_day_safe(row.get("first_date"))
        last_day = _parse_iso_day_safe(row.get("last_date"))
        if first_day is None and last_day is None:
            if (
                float(row.get("current_spend", 0.0) or 0.0) > 0.0
                or float(row.get("current_impressions", 0.0) or 0.0) > 0.0
                or float(row.get("current_clicks", 0.0) or 0.0) > 0.0
            ):
                overlap_campaigns += 1
            continue
        if first_day is None:
            first_day = last_day
        if last_day is None:
            last_day = first_day
        if first_day is None or last_day is None:
            continue
        if last_day < current_start or first_day > current_end:
            continue
        overlap_campaigns += 1
        overlap_start = max(first_day, current_start)
        overlap_end = min(last_day, current_end)
        if overlap_start <= overlap_end:
            for offset in range((overlap_end - overlap_start).days + 1):
                overlap_day_set.add(overlap_start + timedelta(days=offset))
            latest_overlap_day = overlap_end if latest_overlap_day is None else max(latest_overlap_day, overlap_end)

    overlap_days = len(overlap_day_set)
    recency_days: int | None = None
    if latest_overlap_day is not None:
        recency_days = max(0, (current_end - latest_overlap_day).days)

    if overlap_campaigns <= 0:
        return {
            "score": 0,
            "campaigns": 0,
            "overlap_days": 0,
            "recency_days": recency_days,
            "label": "Low",
        }

    campaign_component = min(overlap_campaigns, 6) * 9
    days_component = min(overlap_days, 7) * 5
    recency_component = 0
    if recency_days is not None:
        if recency_days <= 1:
            recency_component = 18
        elif recency_days <= 3:
            recency_component = 12
        elif recency_days <= 7:
            recency_component = 7
        else:
            recency_component = 3
    score = max(0, min(100, 25 + campaign_component + days_component + recency_component))
    if score >= 70:
        label = "High"
    elif score >= 45:
        label = "Medium"
    else:
        label = "Low"
    return {
        "score": int(score),
        "campaigns": overlap_campaigns,
        "overlap_days": overlap_days,
        "recency_days": recency_days,
        "label": label,
    }


def _trade_plan_signal(additional_context: dict[str, object] | None) -> dict[str, object] | None:
    trade_plan = (additional_context or {}).get("trade_plan", {})
    if not isinstance(trade_plan, dict) or not trade_plan.get("enabled"):
        return None
    campaign_rows = trade_plan.get("campaign_rows", [])
    channel_rows = trade_plan.get("channel_split", [])
    campaign_names: list[str] = []
    if isinstance(campaign_rows, list):
        for row in campaign_rows:
            if not isinstance(row, dict):
                continue
            raw_name = str(row.get("campaign", "")).strip()
            if not raw_name:
                continue
            # Ignore placeholders/noisy labels like "1", "2", "-".
            if re.fullmatch(r"[-\d\s.]+", raw_name):
                continue
            name = raw_name[:64]
            if name.lower() not in {n.lower() for n in campaign_names}:
                campaign_names.append(name)
            if len(campaign_names) >= 2:
                break

    paid_direction = "unknown"
    if isinstance(channel_rows, list) and channel_rows:
        paid_direction, _ = _trade_plan_paid_pressure(additional_context)
    overlap_intensity = _trade_plan_overlap_intensity(trade_plan)

    confidence = 56
    if campaign_names:
        confidence += 8
    if paid_direction in {"up", "down"}:
        confidence += 10
    if int(overlap_intensity.get("score", 0) or 0) >= 60:
        confidence += 6
    confidence = max(45, min(82, confidence))

    if paid_direction == "up":
        statement = (
            "Trade plan indicates higher paid-search pressure in the active campaign window; "
            "organic demand allocation can shift from SEO to paid placements."
        )
    elif paid_direction == "down":
        statement = (
            "Trade plan indicates lower paid-search pressure in the active campaign window; "
            "organic demand capture can improve for overlapping intents."
        )
    else:
        statement = (
            "Trade plan confirms active campaign windows; campaign timing should be treated as "
            "a demand-allocation context in SEO interpretation."
        )

    summary = "Trade plan context is available for this window."
    if campaign_names:
        summary = (
            "Campaign overlap detected in analyzed week "
            "(" + ", ".join(f"`{name}`" for name in campaign_names) + ")."
        )
    yoy_availability = trade_plan.get("yoy_availability", {})
    yoy_message = str((yoy_availability or {}).get("message", "")).strip() if isinstance(yoy_availability, dict) else ""
    return {
        "summary": summary,
        "statement": statement,
        "confidence": confidence,
        "overlap_intensity_score": int(overlap_intensity.get("score", 0) or 0),
        "overlap_intensity_label": str(overlap_intensity.get("label", "Low")),
        "overlap_campaigns": int(overlap_intensity.get("campaigns", 0) or 0),
        "overlap_days": int(overlap_intensity.get("overlap_days", 0) or 0),
        "overlap_recency_days": overlap_intensity.get("recency_days"),
        "yoy_availability_message": yoy_message,
    }


def _top_platform_pulse_line(additional_context: dict[str, object] | None) -> str:
    pulse = (additional_context or {}).get("platform_regulatory_pulse", {})
    if not isinstance(pulse, dict) or not pulse.get("enabled"):
        return ""
    rows = pulse.get("rows", [])
    if not isinstance(rows, list):
        return ""
    valid_rows = [row for row in rows if isinstance(row, dict)]
    if not valid_rows:
        return ""
    severities = [
        str(row.get("severity", "")).strip().lower()
        for row in valid_rows
        if isinstance(row, dict)
    ]
    if "high" in severities:
        impact = "high"
    elif "medium" in severities:
        impact = "medium"
    else:
        impact = "low"
    examples = []
    for row in valid_rows[:2]:
        title = str(row.get("title", "")).strip()
        if title:
            examples.append(f"`{title}`")
    example_text = "; ".join(examples) if examples else "platform and regulatory signals"
    return (
        "Platform/regulatory context: "
        + example_text
        + ". Likely weekly impact: "
        + impact
        + "; treat this as context, not primary root cause."
    )


def _top_unique_feature_movers(
    rows: list[dict[str, object]],
    *,
    limit: int = 3,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    by_feature: dict[str, dict[str, object]] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        feature_name = str(row.get("feature", "")).strip()
        key = feature_name.lower()
        if not key:
            continue
        prev = by_feature.get(key)
        if prev is None:
            by_feature[key] = row
            continue
        prev_abs = abs(float(prev.get("delta_clicks_vs_previous", 0.0) or 0.0))
        curr_abs = abs(float(row.get("delta_clicks_vs_previous", 0.0) or 0.0))
        if curr_abs > prev_abs:
            by_feature[key] = row

    unique_rows = list(by_feature.values())
    gains = sorted(
        [row for row in unique_rows if float(row.get("delta_clicks_vs_previous", 0.0) or 0.0) > 0.0],
        key=lambda row: float(row.get("delta_clicks_vs_previous", 0.0) or 0.0),
        reverse=True,
    )[: max(1, int(limit))]
    losses = sorted(
        [row for row in unique_rows if float(row.get("delta_clicks_vs_previous", 0.0) or 0.0) < 0.0],
        key=lambda row: float(row.get("delta_clicks_vs_previous", 0.0) or 0.0),
    )[: max(1, int(limit))]
    return gains, losses


def _feature_mover_pairs(
    rows: list[dict[str, object]],
    *,
    delta_key: str,
    limit: int = 1,
) -> tuple[list[str], list[str]]:
    if not rows:
        return [], []
    gains = sorted(
        [
            row for row in rows
            if isinstance(row, dict) and float(row.get(delta_key, 0.0) or 0.0) > 0.0
        ],
        key=lambda row: float(row.get(delta_key, 0.0) or 0.0),
        reverse=True,
    )[: max(1, int(limit))]
    losses = sorted(
        [
            row for row in rows
            if isinstance(row, dict) and float(row.get(delta_key, 0.0) or 0.0) < 0.0
        ],
        key=lambda row: float(row.get(delta_key, 0.0) or 0.0),
    )[: max(1, int(limit))]
    gain_text = [
        f"{str(row.get('feature', '')).strip()} ({_fmt_signed_compact(row.get(delta_key, 0.0))})"
        for row in gains
        if str(row.get("feature", "")).strip()
    ]
    loss_text = [
        f"{str(row.get('feature', '')).strip()} ({_fmt_signed_compact(row.get(delta_key, 0.0))})"
        for row in losses
        if str(row.get("feature", "")).strip()
    ]
    return gain_text, loss_text


def _serp_appearance_summary_text(additional_context: dict[str, object] | None) -> str:
    feature_split = (additional_context or {}).get("gsc_feature_split", {})
    if not isinstance(feature_split, dict) or not feature_split.get("enabled"):
        return ""
    weekly_rows = feature_split.get("rows_weekly", feature_split.get("rows", []))
    monthly_rows = feature_split.get("rows_mom", [])
    overview_rows = feature_split.get("feature_overview", [])
    if not isinstance(weekly_rows, list):
        weekly_rows = []
    if not isinstance(monthly_rows, list):
        monthly_rows = []
    if not isinstance(overview_rows, list):
        overview_rows = []

    wow_gain, wow_loss = _feature_mover_pairs(weekly_rows, delta_key="delta_clicks_vs_previous", limit=1)
    mom_gain, mom_loss = _feature_mover_pairs(monthly_rows, delta_key="delta_clicks_vs_previous", limit=1)

    yoy_reference_rows: list[dict[str, object]] = []
    if overview_rows:
        for row in overview_rows:
            if not isinstance(row, dict):
                continue
            yoy_reference_rows.append(
                {
                    "feature": row.get("feature", ""),
                    "delta_clicks_vs_previous": row.get("yoy_delta_clicks", 0.0),
                }
            )
    else:
        yoy_reference_rows = weekly_rows
    yoy_gain, yoy_loss = _feature_mover_pairs(yoy_reference_rows, delta_key="delta_clicks_vs_previous", limit=1)

    parts: list[str] = []
    if wow_gain or wow_loss:
        parts.append(
            "WoW "
            + (f"up: {wow_gain[0]}; " if wow_gain else "")
            + (f"down: {wow_loss[0]}" if wow_loss else "")
        )
    if mom_gain or mom_loss:
        parts.append(
            "MoM "
            + (f"up: {mom_gain[0]}; " if mom_gain else "")
            + (f"down: {mom_loss[0]}" if mom_loss else "")
        )
    if yoy_gain or yoy_loss:
        parts.append(
            "YoY "
            + (f"up: {yoy_gain[0]}; " if yoy_gain else "")
            + (f"down: {yoy_loss[0]}" if yoy_loss else "")
        )
    if not parts:
        return ""
    return "SERP appearance split (GSC searchAppearance): " + " | ".join(parts) + "."


def _serp_unified_feature_rows(
    additional_context: dict[str, object] | None,
) -> list[dict[str, object]]:
    feature_split = (additional_context or {}).get("gsc_feature_split", {})
    if not isinstance(feature_split, dict) or not feature_split.get("enabled"):
        return []
    rows = feature_split.get("rows_unified", feature_split.get("feature_overview", []))
    if not isinstance(rows, list):
        return []
    out: list[dict[str, object]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        feature = str(row.get("feature", "")).strip()
        if not feature:
            continue
        out.append(row)
    return out


def _serp_unified_compact_table_lines(
    additional_context: dict[str, object] | None,
    *,
    limit: int = 4,
) -> list[str]:
    rows = _serp_unified_feature_rows(additional_context)
    if not rows:
        return []
    ranked = sorted(
        rows,
        key=lambda row: max(
            abs(float(row.get("wow_delta_clicks", 0.0) or 0.0)),
            abs(float(row.get("mom_delta_clicks", 0.0) or 0.0)),
            abs(float(row.get("yoy_delta_clicks", 0.0) or 0.0)),
        ),
        reverse=True,
    )[: max(1, int(limit))]
    lines = [
        "SERP feature split table (aligned windows):",
        "Feature | WoW | MoM | YoY",
    ]
    for row in ranked:
        feature = str(row.get("feature", "")).strip()
        wow = _fmt_signed_compact(row.get("wow_delta_clicks", 0.0))
        mom = _fmt_signed_compact(row.get("mom_delta_clicks", 0.0))
        yoy = _fmt_signed_compact(row.get("yoy_delta_clicks", 0.0))
        lines.append(f"{feature} | {wow} | {mom} | {yoy}")
    return lines


def _daily_serp_feature_shift_line(additional_context: dict[str, object] | None) -> str:
    context = (additional_context or {}).get("daily_serp_feature_shifts", {})
    if not isinstance(context, dict) or not context.get("enabled"):
        return ""
    rows = context.get("anomalies", context.get("rows", []))
    if not isinstance(rows, list) or not rows:
        return ""
    top = rows[:2]
    snippets: list[str] = []
    for row in top:
        if not isinstance(row, dict):
            continue
        feature = str(row.get("feature", "")).strip()
        day = str(row.get("date", "")).strip()
        delta_pp = float(row.get("delta_share_pp_vs_previous_weekday", 0.0) or 0.0)
        if feature and day:
            snippets.append(f"{day}: {feature} ({delta_pp:+.2f} pp vs previous-weekday)")
    if not snippets:
        return ""
    return (
        "Daily SERP feature-share shifts: "
        + "; ".join(snippets)
        + ". This helps explain where clicks moved between result types."
    )


def _daily_anomaly_detector_line(additional_context: dict[str, object] | None) -> str:
    context = (additional_context or {}).get("daily_gsc_anomalies", {})
    if not isinstance(context, dict) or not context.get("enabled"):
        return ""
    rows = context.get("rows", [])
    if not isinstance(rows, list) or not rows:
        return ""
    snippets: list[str] = []
    for row in rows[:2]:
        if not isinstance(row, dict):
            continue
        day = str(row.get("date", "")).strip()
        delta_clicks = _fmt_signed_compact(row.get("delta_clicks_vs_previous_weekday", 0.0))
        delta_ctr = float(row.get("delta_ctr_pp_vs_previous_weekday", 0.0) or 0.0)
        markers = row.get("markers", [])
        marker_text = ""
        if isinstance(markers, list) and markers:
            marker_text = str(markers[0]).strip()
        if day:
            snippets.append(
                f"{day} ({delta_clicks}; CTR {delta_ctr:+.2f} pp)"
                + (f" -> {marker_text}" if marker_text else "")
            )
    if not snippets:
        return ""
    return "Daily KPI anomaly detector: " + "; ".join(snippets) + "."


def _google_updates_timeline_text(additional_context: dict[str, object] | None) -> str:
    updates = (additional_context or {}).get("google_updates_timeline", {})
    if not isinstance(updates, dict) or not updates.get("enabled"):
        return ""
    summary = updates.get("summary", {})
    if not isinstance(summary, dict):
        return ""
    current_30d = int(summary.get("count_current_30d", 0) or 0)
    previous_30d = int(summary.get("count_previous_30d", 0) or 0)
    yoy_30d = int(summary.get("count_yoy_30d", 0) or 0)
    latest_date = str(summary.get("latest_update_date", "")).strip()
    latest_title = str(summary.get("latest_update_title", "")).strip()
    total_count = int(summary.get("total_count_13m", 0) or 0)
    return (
        "Google update timeline (13M): "
        f"{total_count} tracked updates/signals. "
        f"Last 30d vs previous 30d: {current_30d} vs {previous_30d}; "
        f"vs YoY 30d: {current_30d} vs {yoy_30d}. "
        + (
            f"Latest signal: {latest_date} (`{latest_title}`)."
            if latest_date and latest_title
            else ""
        )
    ).strip()


def _serp_case_study_text(additional_context: dict[str, object] | None) -> str:
    studies = (additional_context or {}).get("serp_case_studies", {})
    if not isinstance(studies, dict) or not studies.get("enabled"):
        return ""
    summary = studies.get("summary", {})
    if not isinstance(summary, dict):
        return ""
    topic_counts = summary.get("topic_counts_13m", {})
    top_topics: list[str] = []
    if isinstance(topic_counts, dict):
        ranked = sorted(
            [
                (str(topic).strip(), int(count or 0))
                for topic, count in topic_counts.items()
                if str(topic).strip()
            ],
            key=lambda item: item[1],
            reverse=True,
        )
        top_topics = [f"{topic}" for topic, _ in ranked[:3]]
    latest_date = str(summary.get("latest_case_date", "")).strip()
    latest_title = str(summary.get("latest_case_title", "")).strip()
    return (
        "SERP case-study scanner (13M): recurring external patterns include "
        + (", ".join(top_topics) if top_topics else "CTR and feature-layout shifts")
        + ". "
        + (
            f"Latest relevant publication: {latest_date} (`{latest_title}`)."
            if latest_date and latest_title
            else ""
        )
    ).strip()


def _brand_ads_hypothesis(
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
    external_signals: list[ExternalSignal],
) -> dict[str, object] | None:
    brand_proxy = _brand_proxy_from_gsc(segment_diagnostics)
    brand_ctr = _brand_ctr_proxy_from_gsc(segment_diagnostics)
    if not brand_proxy or not brand_ctr:
        return None

    clicks_wow_pct = float(brand_proxy.get("delta_pct_vs_previous", 0.0))
    ctr_wow_pp = float(brand_ctr.get("delta_pp_vs_previous", 0.0))
    huge_drop = clicks_wow_pct <= -15.0 and ctr_wow_pp <= -0.25
    huge_rise = clicks_wow_pct >= 15.0 and ctr_wow_pp >= 0.25
    if not (huge_drop or huge_rise):
        return None

    paid_direction, paid_delta = _trade_plan_paid_pressure(additional_context)
    campaign_context = _campaign_event_context(
        external_signals=external_signals,
        query_scope=None,
    )
    campaign_count = len(campaign_context.get("allegro_events", [])) if isinstance(campaign_context, dict) else 0

    confidence = 58
    evidence_bits: list[str] = [
        f"brand clicks WoW {clicks_wow_pct:+.2f}%",
        f"brand CTR WoW {ctr_wow_pp:+.2f} pp",
    ]
    if paid_direction == "up":
        confidence += 12
        evidence_bits.append(f"paid/search spend signal up ({_fmt_signed_compact(paid_delta)})")
    elif paid_direction == "down":
        confidence += 10
        evidence_bits.append(f"paid/search spend signal down ({_fmt_signed_compact(paid_delta)})")
    if campaign_count > 0:
        confidence += 6
        evidence_bits.append(f"campaign signals present ({campaign_count})")
    confidence = max(45, min(85, confidence))

    if huge_drop:
        statement = (
            "Large WoW decline in brand organic clicks/CTR is consistent with higher brand paid-ad pressure in SERP "
            "(possible organic cannibalization by Google Ads)."
        )
        if paid_direction == "up":
            statement = (
                "Large WoW decline in brand organic clicks/CTR with rising paid-search pressure indicates likely "
                "Google Ads cannibalization on brand queries."
            )
    else:
        statement = (
            "Large WoW increase in brand organic clicks/CTR is consistent with reduced brand paid-ad pressure "
            "(possible lower Google Ads overlap on brand queries)."
        )
        if paid_direction == "down":
            statement = (
                "Large WoW increase in brand organic clicks/CTR with lower paid-search pressure indicates likely "
                "Google Ads brand coverage reduction or pause."
            )

    return {
        "statement": statement,
        "confidence": confidence,
        "evidence": "; ".join(evidence_bits),
    }


def _has_allegro_days_like_signal(
    external_signals: list[ExternalSignal],
    country_code: str,
) -> bool:
    code = country_code.strip().upper()
    if code not in {"CZ", "SK"}:
        return False
    tokens = ("allegro days", "dni allegro", "smart week", "megaraty")
    for signal in external_signals:
        text = _normalize_text(f"{signal.title} {signal.details} {signal.source}")
        if not any(token in text for token in tokens):
            continue
        # Prefer country-targeted campaign sources if present.
        if f"({code.lower()})" in signal.source.lower() or "campaign tracker" in signal.source.lower():
            return True
        # Fallback: allow generic source mentions as directional evidence.
        return True
    return False


def _format_trend_points(
    rows: list[dict[str, object]],
    value_key: str,
    limit: int = 3,
    with_sign: bool = False,
    compact: bool = False,
) -> str:
    points: list[str] = []
    for row in rows[:limit]:
        trend = str(row.get("trend", "")).strip()
        if not trend:
            continue
        value = _safe_float(row, value_key)
        if with_sign:
            metric = _fmt_signed_compact(value) if compact else _fmt_signed_int(value)
        else:
            metric = _fmt_compact(value) if compact else _fmt_int(value)
        points.append(f"`{trend}` ({metric})")
    return ", ".join(points)


def _product_trend_rows(
    additional_context: dict[str, object] | None,
    scope_results: list[tuple[str, AnalysisResult]],
) -> dict[str, object]:
    product_trends = (additional_context or {}).get("product_trends", {})
    if not isinstance(product_trends, dict):
        product_trends = {}

    top_rows_raw = product_trends.get("top_rows", 12)
    try:
        top_rows = max(1, int(top_rows_raw))
    except (TypeError, ValueError):
        top_rows = 12

    query_scope = _find_scope(scope_results, "query")

    yoy_rows = product_trends.get("top_yoy_non_brand", [])
    if not isinstance(yoy_rows, list):
        yoy_rows = []
    yoy_rows = [
        row
        for row in yoy_rows
        if isinstance(row, dict) and not _is_noise_or_brand_query(str(row.get("trend", "")))
    ]

    yoy_fallback = False
    if not yoy_rows:
        yoy_rows = _fallback_non_brand_yoy_from_gsc(query_scope=query_scope, top_rows=top_rows)
        yoy_fallback = bool(yoy_rows)

    current_rows = product_trends.get("current_non_brand", [])
    if not isinstance(current_rows, list):
        current_rows = []
    current_rows = [
        row
        for row in current_rows
        if isinstance(row, dict) and not _is_noise_or_brand_query(str(row.get("trend", "")))
    ]

    upcoming_rows = product_trends.get("upcoming_31d", [])
    if not isinstance(upcoming_rows, list):
        upcoming_rows = []
    upcoming_rows = [
        row
        for row in upcoming_rows
        if isinstance(row, dict) and not _is_noise_or_brand_query(str(row.get("trend", "")))
    ]

    return {
        "enabled": bool(product_trends.get("enabled", False) or yoy_rows or current_rows or upcoming_rows),
        "top_rows": top_rows,
        "horizon_days": int(product_trends.get("horizon_days", 31) or 31),
        "yoy_rows": yoy_rows,
        "current_rows": current_rows,
        "upcoming_rows": upcoming_rows,
        "yoy_fallback": yoy_fallback,
    }


def _infer_trend_drivers(
    trend_names: list[str],
    external_signals: list[ExternalSignal],
) -> list[dict[str, object]]:
    normalized_names = [_normalize_text(name) for name in trend_names if str(name).strip()]
    signal_blob = " ".join(
        _normalize_text(f"{signal.title} {signal.details} {signal.source}") for signal in external_signals
    )
    drivers: list[dict[str, object]] = []

    for label, tokens in TREND_EVENT_DRIVER_RULES:
        matched_raw: list[str] = []
        for name in trend_names:
            normalized_name = _normalize_text(name)
            if any(token in normalized_name for token in tokens):
                matched_raw.append(name)

        if not matched_raw:
            continue

        unique_matched: list[str] = []
        seen: set[str] = set()
        for name in matched_raw:
            norm = _normalize_text(name)
            if not norm or norm in seen:
                continue
            seen.add(norm)
            unique_matched.append(name)

        if not unique_matched:
            continue

        signal_support = any(token in signal_blob for token in tokens)
        drivers.append(
            {
                "label": label,
                "count": len(unique_matched),
                "examples": unique_matched[:2],
                "signal_support": signal_support,
            }
        )

    drivers.sort(
        key=lambda row: (
            int(row.get("count", 0)),
            1 if bool(row.get("signal_support", False)) else 0,
        ),
        reverse=True,
    )
    return drivers[:3]


def _trend_demand_bucket(trend: str) -> str:
    normalized = _normalize_text(trend)
    if any(token in normalized for token in TREND_EVENT_CAMPAIGN_TOKENS):
        return "Campaign/event-driven"
    if any(token in normalized for token in TREND_SEASONAL_CALENDAR_TOKENS):
        return "Seasonal/calendar-driven"
    return "Evergreen/base-demand"


def _trend_seasonality_decomposition(
    *,
    yoy_rows: list[dict[str, object]],
    upcoming_rows: list[dict[str, object]],
    horizon_days: int,
) -> dict[str, str]:
    if not yoy_rows:
        return {}

    bucket_weights: dict[str, float] = {
        "Seasonal/calendar-driven": 0.0,
        "Campaign/event-driven": 0.0,
        "Evergreen/base-demand": 0.0,
    }
    bucket_examples: dict[str, list[tuple[str, float]]] = {
        "Seasonal/calendar-driven": [],
        "Campaign/event-driven": [],
        "Evergreen/base-demand": [],
    }

    total_abs_delta = 0.0
    net_delta = 0.0
    for row in yoy_rows:
        if not isinstance(row, dict):
            continue
        trend = str(row.get("trend", "")).strip()
        if not trend:
            continue
        delta = float(_safe_float(row, "delta_value"))
        weight = abs(delta)
        if weight <= 0.0:
            continue
        bucket = _trend_demand_bucket(trend)
        bucket_weights[bucket] = bucket_weights.get(bucket, 0.0) + weight
        bucket_examples.setdefault(bucket, []).append((trend, weight))
        total_abs_delta += weight
        net_delta += delta

    if total_abs_delta <= 0.0:
        return {}

    ranking = sorted(bucket_weights.items(), key=lambda item: item[1], reverse=True)
    parts: list[str] = []
    for bucket, weight in ranking:
        if weight <= 0.0:
            continue
        share = (weight / total_abs_delta) * 100.0
        examples_raw = sorted(
            bucket_examples.get(bucket, []),
            key=lambda item: item[1],
            reverse=True,
        )
        examples = ", ".join(
            f"`{name}`" for name, _ in examples_raw[:2] if str(name).strip()
        )
        if examples:
            parts.append(f"{bucket} {share:.0f}% (e.g. {examples})")
        else:
            parts.append(f"{bucket} {share:.0f}%")

    if not parts:
        return {}

    decomposition_line = (
        "Seasonality decomposition (plain language): "
        + "; ".join(parts)
        + "."
    )

    top_bucket = ranking[0][0] if ranking else ""
    top_share = (ranking[0][1] / total_abs_delta * 100.0) if ranking else 0.0
    if top_bucket == "Seasonal/calendar-driven" and top_share >= 50.0:
        plain_line = (
            "Plain-language trend read: most YoY trend movement looks like calendar seasonality, "
            "so this is primarily a demand-timing story rather than a technical SEO deterioration."
        )
    elif top_bucket == "Campaign/event-driven" and top_share >= 45.0:
        plain_line = (
            "Plain-language trend read: YoY trend movement is mainly campaign/event-driven, "
            "so validate campaign overlap and paid pressure before technical SEO escalation."
        )
    else:
        plain_line = (
            "Plain-language trend read: YoY trend movement is mixed, with a meaningful evergreen component; "
            "check assortment/content coverage together with timing effects."
        )

    delta_direction = "up" if net_delta >= 0 else "down"
    plain_line += f" Net non-brand YoY trend direction is {delta_direction} ({_fmt_signed_compact(net_delta)})."

    watchlist_line = ""
    upcoming_candidates = [row for row in upcoming_rows if isinstance(row, dict)]
    if upcoming_candidates:
        seasonal_count = 0
        event_count = 0
        for row in upcoming_candidates[:8]:
            trend_name = str(row.get("trend", "")).strip()
            bucket = _trend_demand_bucket(trend_name)
            if bucket == "Seasonal/calendar-driven":
                seasonal_count += 1
            elif bucket == "Campaign/event-driven":
                event_count += 1
        if seasonal_count > 0 or event_count > 0:
            watchlist_line = (
                f"Seasonality watchlist (next {horizon_days} days): "
                f"{seasonal_count} seasonal/calendar and {event_count} campaign/event trend(s) in top upcoming signals, "
                "so demand rotation can continue even if SEO efficiency stays stable."
            )

    return {
        "seasonality_decomposition_line": decomposition_line,
        "plain_language_line": plain_line,
        "seasonality_watchlist_line": watchlist_line,
    }


def _build_product_trend_summary(
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    additional_context: dict[str, object] | None,
) -> dict[str, str]:
    trend_ctx = _product_trend_rows(
        additional_context=additional_context,
        scope_results=scope_results,
    )
    if not trend_ctx.get("enabled"):
        return {}

    yoy_rows = list(trend_ctx.get("yoy_rows", []))
    current_rows = list(trend_ctx.get("current_rows", []))
    upcoming_rows = list(trend_ctx.get("upcoming_rows", []))
    top_rows = int(trend_ctx.get("top_rows", 12))
    horizon_days = int(trend_ctx.get("horizon_days", 31))

    yoy_line = ""
    if yoy_rows:
        current_sum = sum(_safe_float(row, "current_value") for row in yoy_rows)
        previous_sum = sum(_safe_float(row, "previous_value") for row in yoy_rows)
        delta = current_sum - previous_sum
        delta_pct = (delta / previous_sum * 100.0) if previous_sum else 0.0

        gainers = sorted(
            [row for row in yoy_rows if _safe_float(row, "delta_value") > 0],
            key=lambda row: _safe_float(row, "delta_value"),
            reverse=True,
        )
        decliners = sorted(
            [row for row in yoy_rows if _safe_float(row, "delta_value") < 0],
            key=lambda row: _safe_float(row, "delta_value"),
        )

        decliners_text = _format_trend_points(decliners, "delta_value", limit=2, with_sign=True, compact=True)
        yoy_line = (
            "Non-brand product YoY (top rows): "
            f"net {_fmt_signed_compact(delta)} ({delta_pct:+.2f}%)."
        )
        gainers_text = _format_trend_points(gainers, "delta_value", limit=3, with_sign=True, compact=True)
        if gainers_text:
            yoy_line += f" Top gainers: {gainers_text}."
        if decliners_text:
            yoy_line += f" Top decliners: {decliners_text}."
        if bool(trend_ctx.get("yoy_fallback")):
            yoy_line += " Source fallback: GSC query trends."

    current_line = ""
    if current_rows:
        ranked_current = sorted(
            current_rows,
            key=lambda row: _safe_float(row, "value"),
            reverse=True,
        )
        current_line = (
            "Current non-brand trend leaders: "
            f"{_format_trend_points(ranked_current, 'value', limit=5, compact=True)}."
        )

    upcoming_line = ""
    if upcoming_rows:
        ranked_upcoming = sorted(
            upcoming_rows,
            key=lambda row: _safe_float(row, "value"),
            reverse=True,
        )
        upcoming_line = (
            f"Next {horizon_days} days watchlist: "
            f"{_format_trend_points(ranked_upcoming, 'value', limit=5, compact=True)}."
        )

    trend_names: list[str] = []
    for row in (yoy_rows[:top_rows] + current_rows[:top_rows] + upcoming_rows[:top_rows]):
        if not isinstance(row, dict):
            continue
        trend = str(row.get("trend", "")).strip()
        if trend:
            trend_names.append(trend)
    drivers = _infer_trend_drivers(trend_names=trend_names, external_signals=external_signals)

    drivers_line = ""
    if drivers:
        parts = []
        for row in drivers:
            label = str(row.get("label", "")).strip()
            count = int(row.get("count", 0))
            examples = row.get("examples", [])
            if isinstance(examples, list):
                example_text = ", ".join(f"`{str(item).strip()}`" for item in examples[:2] if str(item).strip())
            else:
                example_text = ""
            support = "supported by external signals" if bool(row.get("signal_support")) else "weak external support"
            if example_text:
                parts.append(f"{label} ({count} trends; {support}; e.g. {example_text})")
            else:
                parts.append(f"{label} ({count} trends; {support})")
        drivers_line = "Likely event/campaign drivers: " + "; ".join(parts) + "."

    executive_line = ""
    if yoy_rows or current_rows or upcoming_rows:
        snippet_parts: list[str] = []
        if yoy_rows:
            current_sum = sum(_safe_float(row, "current_value") for row in yoy_rows)
            previous_sum = sum(_safe_float(row, "previous_value") for row in yoy_rows)
            delta = current_sum - previous_sum
            snippet_parts.append(f"YoY net {_fmt_signed_int(delta)} across top non-brand rows")
        if current_rows:
            snippet_parts.append(
                f"current leaders: {_format_trend_points(sorted(current_rows, key=lambda row: _safe_float(row, 'value'), reverse=True), 'value', limit=2)}"
            )
        if drivers:
            driver_names = ", ".join(str(row.get("label", "")) for row in drivers[:2] if str(row.get("label", "")).strip())
            if driver_names:
                snippet_parts.append(f"likely drivers: {driver_names}")
        executive_line = "Product-trend summary: " + "; ".join(snippet_parts) + "."

    decomposition = _trend_seasonality_decomposition(
        yoy_rows=yoy_rows,
        upcoming_rows=upcoming_rows,
        horizon_days=horizon_days,
    )

    return {
        "executive_line": executive_line,
        "yoy_line": yoy_line,
        "current_line": current_line,
        "upcoming_line": upcoming_line,
        "drivers_line": drivers_line,
        "seasonality_decomposition_line": str(decomposition.get("seasonality_decomposition_line", "")).strip(),
        "plain_language_line": str(decomposition.get("plain_language_line", "")).strip(),
        "seasonality_watchlist_line": str(decomposition.get("seasonality_watchlist_line", "")).strip(),
    }


def _build_appendix_highlights_lines(
    run_date: date,
    windows: dict[str, DateWindow],
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    additional_context: dict[str, object] | None,
) -> list[str]:
    lines: list[str] = []

    ga4 = (additional_context or {}).get("ga4", {})
    if USE_GA4_IN_REPORT and isinstance(ga4, dict) and ga4.get("enabled"):
        summary = ga4.get("summary", {})
        if isinstance(summary, dict):
            current = summary.get("current", {})
            previous = summary.get("previous", {})
            if isinstance(current, dict) and isinstance(previous, dict):
                sessions_current = _ga4_num(current, "sessions")
                sessions_previous = _ga4_num(previous, "sessions")
                revenue_current = _ga4_num(current, "revenue")
                revenue_previous = _ga4_num(previous, "revenue")
                ga4_parts: list[str] = []
                if sessions_current is not None and sessions_previous is not None:
                    sessions_delta = sessions_current - sessions_previous
                    sessions_delta_pct = (
                        (sessions_delta / sessions_previous * 100.0)
                        if sessions_previous
                        else 0.0
                    )
                    ga4_parts.append(
                        f"sessions {_fmt_int(sessions_current)} ({sessions_delta_pct:+.2f}% WoW)"
                    )
                if revenue_current is not None and revenue_previous is not None:
                    revenue_delta = revenue_current - revenue_previous
                    revenue_delta_pct = (
                        (revenue_delta / revenue_previous * 100.0)
                        if revenue_previous
                        else 0.0
                    )
                    ga4_parts.append(
                        f"purchase revenue {_fmt_int(revenue_current)} ({revenue_delta_pct:+.2f}% WoW)"
                    )
                if ga4_parts:
                    lines.append("- GA4 weekly check: " + ", ".join(ga4_parts) + ".")

                channels = ga4.get("channels", {})
                yoy_deltas = channels.get("yoy_deltas", []) if isinstance(channels, dict) else []
                if isinstance(yoy_deltas, list) and yoy_deltas:
                    top_riser = max(
                        (row for row in yoy_deltas if isinstance(row, dict)),
                        key=lambda row: float(row.get("delta_vs_yoy", 0.0) or 0.0),
                        default=None,
                    )
                    top_faller = min(
                        (row for row in yoy_deltas if isinstance(row, dict)),
                        key=lambda row: float(row.get("delta_vs_yoy", 0.0) or 0.0),
                        default=None,
                    )
                    if isinstance(top_riser, dict) and isinstance(top_faller, dict):
                        lines.append(
                            "- GA4 YoY channels: "
                            f"top riser `{top_riser.get('channel', '')}` ({_fmt_signed_int(top_riser.get('delta_vs_yoy', 0.0))}), "
                            f"top decline `{top_faller.get('channel', '')}` ({_fmt_signed_int(top_faller.get('delta_vs_yoy', 0.0))})."
                        )
    elif USE_GA4_IN_REPORT and isinstance(ga4, dict):
        errors = ga4.get("errors", [])
        if isinstance(errors, list) and errors:
            lines.append(f"- GA4 warning: {str(errors[0]).strip()}")

    senuto_intelligence = (additional_context or {}).get("senuto_intelligence", {})
    if isinstance(senuto_intelligence, dict) and senuto_intelligence.get("enabled"):
        competitors = senuto_intelligence.get("competitors_overview", [])
        if isinstance(competitors, list) and competitors:
            top_comp = next(
                (row for row in competitors if isinstance(row, dict)),
                None,
            )
            if isinstance(top_comp, dict):
                lines.append(
                    "- Senuto competitor radar: top overlapping competitor "
                    f"`{top_comp.get('domain', '')}` with "
                    f"{_fmt_int(top_comp.get('common_keywords', 0.0))} common keywords."
                )

    allegro_trends = (additional_context or {}).get("allegro_trends", {})
    if isinstance(allegro_trends, dict) and allegro_trends.get("enabled"):
        rows = allegro_trends.get("rows", [])
        if isinstance(rows, list) and rows:
            top_row = next((row for row in rows if isinstance(row, dict)), None)
            if isinstance(top_row, dict):
                lines.append(
                    "- Allegro Trends API: top tracked query "
                    f"`{str(top_row.get('query', '')).strip()}` "
                    f"with GMV {_fmt_int(top_row.get('gmv', 0.0))} and visits {_fmt_int(top_row.get('visit', 0.0))} "
                    f"in {str(allegro_trends.get('from', '')).strip()} to {str(allegro_trends.get('till', '')).strip()}."
                )
    elif isinstance(allegro_trends, dict):
        errors = allegro_trends.get("errors", [])
        if isinstance(errors, list) and errors:
            lines.append(f"- Allegro Trends warning: {str(errors[0]).strip()}")

    trend_summary = _build_product_trend_summary(
        scope_results=scope_results,
        external_signals=external_signals,
        additional_context=additional_context,
    )
    for key in (
        "yoy_line",
        "seasonality_decomposition_line",
        "plain_language_line",
        "seasonality_watchlist_line",
        "current_line",
        "upcoming_line",
        "drivers_line",
    ):
        text = trend_summary.get(key, "").strip()
        if text:
            lines.append(f"- {text}")

    cluster_line = _query_cluster_summary_line(_find_scope(scope_results, "query"))
    if cluster_line:
        lines.append(f"- {cluster_line}")

    campaign_context = _campaign_event_context(
        external_signals=external_signals,
        query_scope=_find_scope(scope_results, "query"),
    )
    allegro_events = campaign_context.get("allegro_events", [])
    competitor_events = campaign_context.get("competitor_events", [])
    query_events = campaign_context.get("query_events", [])
    allegro_count = len(allegro_events) if isinstance(allegro_events, list) else 0
    competitor_count = len(competitor_events) if isinstance(competitor_events, list) else 0
    query_count = len(query_events) if isinstance(query_events, list) else 0
    if allegro_count or competitor_count or query_count:
        campaign_days: list[date] = []
        if isinstance(allegro_events, list):
            campaign_days.extend(
                row.day for row in allegro_events if isinstance(row, ExternalSignal)
            )
        if isinstance(competitor_events, list):
            campaign_days.extend(
                row[0].day
                for row in competitor_events
                if isinstance(row, tuple) and len(row) == 2 and isinstance(row[0], ExternalSignal)
            )
        period_label = ""
        if campaign_days:
            period_label = (
                f" Window: {min(campaign_days).isoformat()} to {max(campaign_days).isoformat()}."
            )
        line = (
            "- Campaign-event context from appendix: "
            f"Allegro mentions={allegro_count}, competitor mentions={competitor_count}, "
            f"campaign query movers={query_count}.{period_label}"
        )
        if isinstance(allegro_events, list) and allegro_events:
            latest = allegro_events[0]
            line += f" Latest Allegro signal: {latest.day.isoformat()} (`{latest.title}`)."
        elif isinstance(competitor_events, list) and competitor_events:
            latest_signal, competitor = competitor_events[0]
            line += (
                f" Latest competitor signal: {latest_signal.day.isoformat()} "
                f"({competitor}: `{latest_signal.title}`)."
            )
        lines.append(line)

    market_events = (additional_context or {}).get("market_event_calendar", {})
    if isinstance(market_events, dict) and market_events.get("enabled"):
        rows = market_events.get("events", [])
        if isinstance(rows, list) and rows:
            high_count = sum(
                1
                for row in rows
                if isinstance(row, dict)
                and str(row.get("impact_level", "")).strip().upper() == "HIGH"
            )
            event_days: list[date] = []
            for row in rows:
                if not isinstance(row, dict):
                    continue
                day_text = str(row.get("date", "")).strip()
                if not day_text:
                    continue
                try:
                    day = date.fromisoformat(day_text[:10])
                except ValueError:
                    continue
                event_days.append(day)
            period = ""
            if event_days:
                period = f" Window: {min(event_days).isoformat()} to {max(event_days).isoformat()}."
            lines.append(
                "- Market-event API context: "
                f"{_fmt_int(len(rows))} events for {str(market_events.get('country_code', '')).strip() or 'market'}, "
                f"high-impact candidates={_fmt_int(high_count)}.{period}"
            )

    trends = (additional_context or {}).get("google_trends", [])
    if isinstance(trends, list) and trends:
        top_topics: list[str] = []
        for row in trends[:3]:
            if not isinstance(row, dict):
                continue
            topic = str(row.get("topic", "")).strip()
            if not topic:
                continue
            traffic = int(row.get("approx_traffic", 0))
            top_topics.append(f"`{topic}` ({_fmt_int(traffic)})")
        if top_topics:
            country_label = str((additional_context or {}).get("country_code", "")).strip() or "market"
            lines.append(f"- Google Trends ({country_label}) top topics: " + ", ".join(top_topics) + ".")

    pagespeed = (additional_context or {}).get("pagespeed", {})
    if isinstance(pagespeed, dict) and pagespeed:
        mobile = pagespeed.get("mobile", {})
        if isinstance(mobile, dict):
            lcp = float(mobile.get("lcp_ms", 0.0))
            inp = float(mobile.get("inp_ms", 0.0))
            cls = float(mobile.get("cls", 0.0))
            if lcp or inp or cls:
                lines.append(
                    "- CrUX mobile snapshot: "
                    f"LCP={lcp:.0f}ms, INP={inp:.0f}ms, CLS={cls:.2f} "
                    "(use this when traffic shifts are suspiciously page/device-specific)."
                )

    macro = (additional_context or {}).get("macro", {})
    if isinstance(macro, dict):
        nbp = macro.get("nbp_fx", {})
        if isinstance(nbp, dict):
            eur = nbp.get("eur_pln", {})
            usd = nbp.get("usd_pln", {})
            if isinstance(eur, dict) and isinstance(usd, dict):
                eur_delta = float(eur.get("delta_pct_vs_previous", 0.0))
                usd_delta = float(usd.get("delta_pct_vs_previous", 0.0))
                if eur_delta or usd_delta:
                    lines.append(
                        "- FX context vs previous week: "
                        f"EUR/PLN {eur_delta:+.2f}%, USD/PLN {usd_delta:+.2f}%."
                    )

    current_window = windows.get("current_28d")
    previous_window = windows.get("previous_28d")
    if isinstance(current_window, DateWindow) and isinstance(previous_window, DateWindow):
        lines.append(
            "- Date windows used in this run: "
            f"current={current_window.start.isoformat()} to {current_window.end.isoformat()}, "
            f"previous={previous_window.start.isoformat()} to {previous_window.end.isoformat()}."
        )

    return lines[:12]


def _summarize_query_delta_rows(rows: list[KeyDelta], limit: int = 5) -> str:
    points: list[str] = []
    for row in rows[:limit]:
        key = str(row.key).strip()
        if not key:
            continue
        points.append(
            f"`{key}` ({_fmt_signed_int(row.click_delta_vs_previous)} vs prev; prev={_fmt_int(row.previous_clicks)})"
        )
    return "; ".join(points)


def _top_non_brand_trends_line(
    scope_results: list[tuple[str, AnalysisResult]],
    additional_context: dict[str, object] | None,
    limit: int = 5,
) -> str:
    trend_ctx = _product_trend_rows(
        additional_context=additional_context,
        scope_results=scope_results,
    )
    current_rows = trend_ctx.get("current_rows", [])
    if isinstance(current_rows, list) and current_rows:
        ranked = sorted(
            [row for row in current_rows if isinstance(row, dict)],
            key=lambda row: _safe_float(row, "value"),
            reverse=True,
        )
        points = _format_trend_points(ranked, "value", limit=limit)
        if points:
            return f"Top {limit} current non-brand trends: {points}."

    yoy_rows = trend_ctx.get("yoy_rows", [])
    if isinstance(yoy_rows, list) and yoy_rows:
        ranked = sorted(
            [row for row in yoy_rows if isinstance(row, dict)],
            key=lambda row: abs(_safe_float(row, "delta_value")),
            reverse=True,
        )
        points = _format_trend_points(ranked, "delta_value", limit=limit, with_sign=True)
        if points:
            return f"Top {limit} YoY non-brand trend movers: {points}."
    return ""


def _build_executive_summary_lines(
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
    hypotheses: list[dict[str, object]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
    senuto_summary: dict[str, float] | None,
    senuto_error: str | None,
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    yoy = totals["yoy_52w"]

    click_wow_pct = _signed_pct(_ratio_delta(current.clicks, previous.clicks))
    click_yoy_pct = _signed_pct(_ratio_delta(current.clicks, yoy.clicks))
    impression_wow_pct = _signed_pct(_ratio_delta(current.impressions, previous.impressions))
    impression_yoy_pct = _signed_pct(_ratio_delta(current.impressions, yoy.impressions))
    ctr_wow_pp = (current.ctr - previous.ctr) * 100.0
    ctr_yoy_pp = (current.ctr - yoy.ctr) * 100.0
    query_scope = _find_scope(scope_results, "query")
    growth_themes = ""
    decline_themes = ""
    if query_scope:
        growth_themes, decline_themes = _executive_yoy_theme_summary(query_scope, limit=3)
    pos_wow_delta = current.position - previous.position
    pos_yoy_delta = current.position - yoy.position

    top_driver_rows = _top_hypotheses_for_actions(hypotheses, limit=2)
    top_driver = top_driver_rows[0] if top_driver_rows else {}
    top_driver_name = str(top_driver.get("category", "Demand timing")).strip() or "Demand timing"
    top_driver_priority = int(top_driver.get("driver_priority_score", 0) or 0)

    why_parts: list[str] = []
    if growth_themes:
        why_parts.append(f"gainers: {growth_themes}")
    if decline_themes:
        why_parts.append(f"decliners: {decline_themes}")
    why_text = "; ".join(why_parts) if why_parts else "traffic changed mostly because demand rotated between themes and pages"

    risk_bits: list[str] = []
    if _ratio_delta(current.clicks, yoy.clicks) <= -0.08:
        risk_bits.append("YoY traffic gap is still material")
    if ctr_yoy_pp <= -0.20:
        risk_bits.append("YoY click efficiency is weaker")
    if pos_yoy_delta >= 0.20:
        risk_bits.append("average ranking quality is below last year")
    if not risk_bits:
        risk_bits.append("main risk is traffic allocation between page types, not a broad technical outage")

    source_quality = _source_quality_summary(additional_context)
    comparability = _comparability_summary(
        additional_context,
        brand_trends_available=bool(((additional_context or {}).get("google_trends_brand", {}) or {}).get("enabled")),
        brand_proxy_available=isinstance(_brand_proxy_from_gsc(segment_diagnostics), dict),
    )
    reliability_line = " Data reliability: baseline acceptable for decision use."
    if source_quality and comparability:
        reliability_line = f" Data reliability: {source_quality} {comparability}"
    elif source_quality:
        reliability_line = f" Data reliability: {source_quality}"
    elif comparability:
        reliability_line = f" Data reliability: {comparability}"

    opportunity = (
        f"`{top_driver_name}` is the biggest lever this week"
        + (f" ({top_driver_priority}/100 priority)" if top_driver_priority else "")
        + "; quick execution can recover traffic share in the next run."
    )
    decision_text = (
        "Decision this week: keep demand/timing and routing actions first, and escalate technical SEO only if "
        "CTR/position deteriorate in the next validation window."
    )
    actions = _top_priority_actions(hypotheses, limit=2)
    next_action_text = "; ".join(actions) if actions else "SEO Team: validate top hypothesis on refreshed weekly data (SEO Team | next run)"

    return [
        "- **What changed**: clicks "
        f"{_fmt_compact(current.clicks)} ({click_wow_pct} WoW; {click_yoy_pct} YoY), impressions "
        f"{_fmt_compact(current.impressions)} ({impression_wow_pct} WoW; {impression_yoy_pct} YoY), "
        f"CTR {_pct(current.ctr)} ({ctr_wow_pp:+.2f} pp WoW; {ctr_yoy_pp:+.2f} pp YoY), avg position "
        f"{current.position:.2f} ({_position_delta_label(pos_wow_delta)} WoW; {_position_delta_label(pos_yoy_delta)} YoY). [E1]",
        "- **Why**: In plain language, this week looks like demand and page-mix rotation, not one broad technical SEO failure mode; "
        + why_text
        + ".",
        "- **Risk**: Business implication: "
        + "; ".join(risk_bits[:2])
        + ".",
        "- **Opportunity**: "
        + opportunity
        + "."
        + reliability_line,
        "- **Decision**: " + decision_text,
        "- **Next action**: Priority actions (owner | ETA): " + next_action_text + ".",
    ]


def _build_leadership_snapshot_lines(
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    yoy = totals["yoy_52w"]
    wow_pct = _signed_pct(_ratio_delta(current.clicks, previous.clicks))
    yoy_pct = _signed_pct(_ratio_delta(current.clicks, yoy.clicks))

    query_scope = _find_scope(scope_results, "query")
    cluster_rows = _query_cluster_rows(query_scope)
    gains = sorted(
        [row for row in cluster_rows if float(row.get("delta_vs_previous", 0.0)) > 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
        reverse=True,
    )[:2]
    losses = sorted(
        [row for row in cluster_rows if float(row.get("delta_vs_previous", 0.0)) < 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
    )[:2]

    gains_text = "; ".join(
        f"{row.get('cluster', '')} ({_fmt_signed_compact(row.get('delta_vs_previous', 0.0))})"
        for row in gains
        if str(row.get("cluster", "")).strip()
    )
    losses_text = "; ".join(
        f"{row.get('cluster', '')} ({_fmt_signed_compact(row.get('delta_vs_previous', 0.0))})"
        for row in losses
        if str(row.get("cluster", "")).strip()
    )

    return [
        (
            "- **What changed**: **GSC organic clicks** moved "
            f"**{_fmt_signed_compact(current.clicks - previous.clicks)} WoW ({wow_pct})** and "
            f"**{_fmt_signed_compact(current.clicks - yoy.clicks)} YoY ({yoy_pct})**. "
            + (f"Main gains: {gains_text}. " if gains_text else "")
            + (f"Main declines: {losses_text}." if losses_text else "")
        ).strip(),
        (
            "- **Why it matters**: the current pattern is mainly **category-level demand rotation** "
            "(seasonality/event timing) rather than broad SEO efficiency deterioration, "
            "so decisions should prioritize demand timing and paid/campaign impact quantification before heavy technical escalation."
        ),
    ]


def _confidence_for_category(
    hypotheses: list[dict[str, object]],
    category_tokens: tuple[str, ...],
) -> int | None:
    for row in hypotheses:
        category = _normalize_text(str(row.get("category", "")))
        if any(token in category for token in category_tokens):
            try:
                return int(row.get("confidence", 0))
            except (TypeError, ValueError):
                return None
    return None


def _latest_google_update_signal(
    external_signals: list[ExternalSignal],
) -> ExternalSignal | None:
    updates = [
        signal
        for signal in external_signals
        if signal.source in {"Google Search Status", "Google Search Central Blog"}
        and "update" in _normalize_text(f"{signal.title} {signal.details}")
    ]
    if not updates:
        return None
    return max(updates, key=lambda signal: signal.day)


def _build_what_is_happening_lines(
    totals: dict[str, MetricSummary],
    windows: dict[str, DateWindow],
    scope_results: list[tuple[str, AnalysisResult]],
    hypotheses: list[dict[str, object]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
) -> list[str]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    yoy = totals["yoy_52w"]
    current_window = windows.get("current_28d")
    previous_window = windows.get("previous_28d")
    query_scope = _find_scope(scope_results, "query")
    cluster_rows = _query_cluster_rows(query_scope)

    seasonality_conf = _confidence_for_category(hypotheses, ("seasonality",))
    events_conf = _confidence_for_category(hypotheses, ("events",))
    template_conf = _confidence_for_category(hypotheses, ("template",))
    algorithm_conf = _confidence_for_category(hypotheses, ("algorithm",))
    weather_conf = _confidence_for_category(hypotheses, ("macro", "weather"))

    period_label = ""
    if isinstance(current_window, DateWindow) and isinstance(previous_window, DateWindow):
        period_label = (
            f"{current_window.start.isoformat()} to {current_window.end.isoformat()} "
            f"vs {previous_window.start.isoformat()} to {previous_window.end.isoformat()}"
        )

    def _cluster_change_text(row: dict[str, object], field: str = "delta_vs_previous") -> str:
        delta = float(row.get(field, 0.0))
        current_clicks = float(row.get("current_clicks", 0.0))
        baseline = current_clicks - delta if field == "delta_vs_previous" else current_clicks - delta
        pct = _ratio_delta(current_clicks, baseline) if baseline else 0.0
        return f"{str(row.get('cluster', '')).strip()} ({_fmt_signed_compact(delta)}; {_signed_pct(pct)})"

    gains = sorted(
        [row for row in cluster_rows if float(row.get("delta_vs_previous", 0.0)) > 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
        reverse=True,
    )[:3]
    losses = sorted(
        [row for row in cluster_rows if float(row.get("delta_vs_previous", 0.0)) < 0],
        key=lambda row: float(row.get("delta_vs_previous", 0.0)),
    )[:3]
    gain_text = "; ".join(_cluster_change_text(row) for row in gains if str(row.get("cluster", "")).strip())
    loss_text = "; ".join(_cluster_change_text(row) for row in losses if str(row.get("cluster", "")).strip())

    winter_cluster = next(
        (row for row in cluster_rows if str(row.get("cluster", "")).strip() == "Winter & weather demand"),
        None,
    )
    event_cluster = next(
        (row for row in cluster_rows if str(row.get("cluster", "")).strip() == "WOŚP & charity events"),
        None,
    )
    non_brand_cluster = next(
        (row for row in cluster_rows if str(row.get("cluster", "")).strip() == "Other non-brand demand"),
        None,
    )
    winter_yoy_delta = float(winter_cluster.get("delta_vs_yoy", 0.0)) if isinstance(winter_cluster, dict) else 0.0
    event_yoy_delta = float(event_cluster.get("delta_vs_yoy", 0.0)) if isinstance(event_cluster, dict) else 0.0
    non_brand_yoy_delta = float(non_brand_cluster.get("delta_vs_yoy", 0.0)) if isinstance(non_brand_cluster, dict) else 0.0

    total_yoy_delta = float(current.clicks - yoy.clicks)
    non_brand_share = 0.0
    if total_yoy_delta != 0:
        non_brand_share = (non_brand_yoy_delta / total_yoy_delta) * 100.0

    home_wow = 0.0
    home_yoy = 0.0
    home_wow_pct = ""
    home_yoy_pct = ""
    page_name_line = ""
    template_rows = (segment_diagnostics or {}).get("page_template")
    if template_rows:
        home_row = next(
            (
                row
                for row in template_rows
                if str(row.get("segment", "")).strip().lower() == "home"
            ),
            None,
        )
        if home_row and float(home_row.get("delta_vs_previous", 0.0)) < 0:
            home_wow = float(home_row.get("delta_vs_previous", 0.0))
            home_yoy = float(home_row.get("delta_vs_yoy", 0.0))
            home_prev = float(home_row.get("previous_clicks", 0.0))
            home_yoy_clicks = float(home_row.get("yoy_clicks", 0.0))
            home_wow_pct = _signed_pct(_ratio_delta(home_prev + home_wow, home_prev))
            home_yoy_pct = _signed_pct(_ratio_delta(home_yoy_clicks + home_yoy, home_yoy_clicks))
            page_name_line = (
                f"home is down {_fmt_signed_compact(home_wow)} WoW ({home_wow_pct}) "
                f"and {_fmt_signed_compact(home_yoy)} YoY ({home_yoy_pct})."
            )

    brand_note = ""
    brand_wow = 0.0
    brand_yoy = 0.0
    brand_enabled = False
    brand_context = (additional_context or {}).get("google_trends_brand", {})
    brand_summary = brand_context.get("summary", {}) if isinstance(brand_context, dict) else {}
    if isinstance(brand_summary, dict):
        brand_wow = float(brand_summary.get("delta_pct_vs_previous", 0.0))
        brand_yoy = float(brand_summary.get("delta_pct_vs_yoy", 0.0))
        brand_enabled = bool(brand_context.get("enabled")) if isinstance(brand_context, dict) else False
        if brand_wow > 0:
            brand_note = (
                f"Google Trends brand demand is up WoW ({brand_wow:+.2f}%), "
                "which supports a click-distribution effect in search results rather than pure demand collapse."
            )
        elif brand_wow < 0:
            brand_note = (
                f"Google Trends brand demand is down WoW ({brand_wow:+.2f}%), "
                "so demand-side softness is part of the homepage decline."
            )
        if home_yoy < 0 and brand_yoy >= 0:
            brand_note += " YoY brand interest is stable/up, so URL allocation remains a high-priority check."

    wow_pct = _signed_pct(_ratio_delta(current.clicks, previous.clicks))
    yoy_pct = _signed_pct(_ratio_delta(current.clicks, yoy.clicks))
    ctr_wow_pp = (current.ctr - previous.ctr) * 100.0
    pos_wow_delta = current.position - previous.position

    lines: list[str] = []
    lines.append(
        "In summary, this is a category-level demand rotation week: seasonal normalization and event-driven spikes explain most movement, "
        "while core SEO efficiency remains stable. Current data points to a demand/exposure effect first, so technical escalation should be secondary unless efficiency weakens."
    )

    lines.append(
        (
            "**WoW diagnosis**: **clicks changed by "
            f"{_fmt_signed_compact(current.clicks - previous.clicks)} ({wow_pct})**, with CTR {ctr_wow_pp:+.2f} pp and avg position {_position_delta_label(pos_wow_delta)}. [E1] "
            "This pattern indicates category-level demand rotation, not broad quality collapse. "
            + (f"Growth clusters: {gain_text}. " if gain_text else "")
            + (f"Decline clusters: {loss_text}. " if loss_text else "")
            + (f"Window: {period_label}." if period_label else "")
        ).strip()
    )
    gsc_coverage = (additional_context or {}).get("gsc_data_coverage", {})
    if isinstance(gsc_coverage, dict) and gsc_coverage.get("p52w_mode") == "masked_to_days_with_data":
        days_with_data = int(gsc_coverage.get("days_with_data", 0) or 0)
        days_total = int(gsc_coverage.get("days_total", 0) or 0)
        lines.append(
            "P52W comparison was computed only for days with available GSC data in the analyzed week: "
            f"{days_with_data}/{days_total} days."
        )
    daily_story = _build_daily_gsc_storyline(
        additional_context=additional_context,
        external_signals=external_signals,
        weather_summary=weather_summary,
        top_n=3,
    )
    daily_narrative_lines = daily_story.get("narrative_lines", [])
    if isinstance(daily_narrative_lines, list):
        for row in daily_narrative_lines[:5]:
            text = str(row).strip()
            if text:
                lines.append(text)
    serp_daily_line = _daily_serp_feature_shift_line(additional_context)
    if serp_daily_line:
        lines.append(serp_daily_line)
    daily_anomaly_line = _daily_anomaly_detector_line(additional_context)
    if daily_anomaly_line:
        lines.append(daily_anomaly_line)

    yoy_line = (
        f"**YoY diagnosis**: **clicks are {_fmt_signed_compact(current.clicks - yoy.clicks)} ({yoy_pct})**. "
        f"Winter cluster contribution: {_fmt_signed_compact(winter_yoy_delta)}; event cluster contribution: {_fmt_signed_compact(event_yoy_delta)}. "
    )
    if total_yoy_delta < 0 and non_brand_yoy_delta > 0:
        yoy_line += (
            f"Other non-brand demand adds {_fmt_signed_compact(non_brand_yoy_delta)} and **offsets {abs(non_brand_share):.1f}% of total YoY decline**."
        )
    else:
        yoy_line += (
            f"Other non-brand demand contributes {_fmt_signed_compact(non_brand_yoy_delta)} "
            f"({non_brand_share:+.1f}% of total YoY click delta)."
        )
    lines.append(yoy_line.strip())

    trend_summary = _build_product_trend_summary(
        scope_results=scope_results,
        external_signals=external_signals,
        additional_context=additional_context,
    )
    trend_decomposition_line = str(trend_summary.get("seasonality_decomposition_line", "")).strip()
    if trend_decomposition_line:
        lines.append(trend_decomposition_line)
    trend_plain_line = str(trend_summary.get("plain_language_line", "")).strip()
    if trend_plain_line:
        lines.append(trend_plain_line)

    brand_proxy = _brand_proxy_from_gsc(segment_diagnostics)
    if brand_enabled and brand_proxy:
        lines.append(
            "**Brand demand context (Google Trends + GSC)**: Google Trends branded-search interest is "
            f"{brand_wow:+.2f}% WoW and {brand_yoy:+.2f}% YoY; "
            f"GSC brand clicks are {brand_proxy.get('delta_pct_vs_previous', 0.0):+.2f}% WoW "
            f"({_fmt_signed_compact(brand_proxy.get('delta_vs_previous', 0.0))}) and "
            f"{brand_proxy.get('delta_pct_vs_yoy', 0.0):+.2f}% YoY "
            f"({_fmt_signed_compact(brand_proxy.get('delta_vs_yoy', 0.0))}); "
            f"GSC brand impressions are {brand_proxy.get('impressions_delta_pct_vs_previous', 0.0):+.2f}% WoW "
            f"({_fmt_signed_compact(brand_proxy.get('impressions_delta_vs_previous', 0.0))}) and "
            f"{brand_proxy.get('impressions_delta_pct_vs_yoy', 0.0):+.2f}% YoY "
            f"({_fmt_signed_compact(brand_proxy.get('impressions_delta_vs_yoy', 0.0))})."
        )
    elif brand_enabled:
        lines.append(
            "**Brand demand context**: Google Trends branded-search interest is "
            f"{brand_wow:+.2f}% WoW and {brand_yoy:+.2f}% YoY, "
            "and GSC brand-proxy segment is unavailable in this run."
        )
    elif brand_proxy:
        lines.append(
            "**Brand demand context (GSC proxy)**: Google Trends brand series is unavailable in this run, "
            f"so we use GSC brand-query proxy: WoW {brand_proxy.get('delta_pct_vs_previous', 0.0):+.2f}% "
            f"({_fmt_signed_compact(brand_proxy.get('delta_vs_previous', 0.0))}), "
            f"impressions WoW {brand_proxy.get('impressions_delta_pct_vs_previous', 0.0):+.2f}% "
            f"({_fmt_signed_compact(brand_proxy.get('impressions_delta_vs_previous', 0.0))}), "
            f"YoY {brand_proxy.get('delta_pct_vs_yoy', 0.0):+.2f}% "
            f"({_fmt_signed_compact(brand_proxy.get('delta_vs_yoy', 0.0))}), "
            f"impressions YoY {brand_proxy.get('impressions_delta_pct_vs_yoy', 0.0):+.2f}% "
            f"({_fmt_signed_compact(brand_proxy.get('impressions_delta_vs_yoy', 0.0))})."
        )
    else:
        lines.append(
            "**Brand demand context**: unavailable in this run (Google Trends brand time-series missing), "
            "and no GSC brand-proxy segment was available."
        )

    country_code = str((additional_context or {}).get("country_code", "")).strip().upper()
    brand_ctr_proxy = _brand_ctr_proxy_from_gsc(segment_diagnostics)
    if (
        country_code in {"CZ", "SK"}
        and _has_allegro_days_like_signal(external_signals, country_code=country_code)
        and isinstance(brand_ctr_proxy, dict)
        and float(brand_ctr_proxy.get("delta_pp_vs_previous", 0.0)) <= -0.10
    ):
        lines.append(
            "**Brand CTR during Allegro Days (CZ/SK)**: brand CTR fell "
            f"{float(brand_ctr_proxy.get('delta_pp_vs_previous', 0.0)):+.2f} pp WoW "
            f"(YoY {float(brand_ctr_proxy.get('delta_pp_vs_yoy', 0.0)):+.2f} pp) while campaign signals were active. "
            "Interpretation: part of brand demand may be reallocated to paid brand ads in SERP, so organic brand CTR softening is not automatically a technical SEO issue."
        )
    brand_ads_hyp = _brand_ads_hypothesis(
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
        external_signals=external_signals,
    )
    if isinstance(brand_ads_hyp, dict):
        lines.append(
            "**Brand Ads overlap hypothesis (WoW)**: "
            f"{str(brand_ads_hyp.get('statement', '')).strip()} "
            f"Evidence: {str(brand_ads_hyp.get('evidence', '')).strip()}. "
            f"Confidence: {_confidence_bucket(brand_ads_hyp.get('confidence', 0))}."
        )
    contradiction_rows = _contradiction_reconciliation_lines(
        totals=totals,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
    )
    if contradiction_rows:
        lines.append(
            "Contradiction check: mixed directional signals were detected; each contradiction requires explicit reconciliation before escalation."
        )
        for row in contradiction_rows[:2]:
            lines.append("Reconciliation: " + row)

    movement_direction = "up" if (current.clicks - previous.clicks) >= 0 else "down"
    causal_chain = (
        f"**Causal chain**: Observation -> traffic is {movement_direction} while efficiency signals are mixed-but-stable WoW. "
        "Evidence -> cluster deltas and CTR/position behavior indicate a theme-level demand rotation. "
        "Hypothesis -> seasonality/event timing plus paid/campaign timing explain most movement. "
        "Decision -> treat this as demand/exposure-led first and quantify paid/campaign incremental effect before escalating technical SEO."
    )
    if page_name_line:
        causal_chain += f" {page_name_line}"
    if brand_note:
        causal_chain += f" {brand_note}"
    if template_conf is not None:
        causal_chain += f" Confidence (template/routing): {_confidence_bucket(template_conf)}."
    lines.append(causal_chain)
    causality_guardrail = _causality_guardrail_summary(hypotheses)
    if causality_guardrail:
        lines.append(causality_guardrail)

    causal_confidence = int(hypotheses[0].get("confidence", 65)) if hypotheses else 65
    conflict_flags: list[str] = []
    if (current.clicks - previous.clicks) > 0 and home_wow < 0:
        conflict_flags.append("overall clicks up while home template is down")
    if brand_enabled and brand_wow < 0 and (current.clicks - previous.clicks) > 0:
        conflict_flags.append("brand demand down while total clicks are up")
    freshness_rows = _source_freshness_rows(additional_context)
    stale_or_degraded = sum(
        1
        for row in freshness_rows
        if str(row.get("status", "")).strip().lower() in {"stale", "degraded"}
    )
    causal_confidence -= 6 * len(conflict_flags)
    if stale_or_degraded > 0:
        causal_confidence -= 4
    causal_confidence = max(45, min(95, causal_confidence))
    if conflict_flags:
        lines.append(
            "Causality confidence: "
            f"{_confidence_bucket(causal_confidence)} ({causal_confidence}/100). Mixed signals detected ({'; '.join(conflict_flags[:2])}); "
            "keep this as a working explanation until next-run validation."
        )
    else:
        lines.append(
            f"Causality confidence: {_confidence_bucket(causal_confidence)} ({causal_confidence}/100). Main signals are directionally consistent in this run."
        )

    update_signal = _latest_google_update_signal(external_signals)
    if update_signal:
        update_line = (
            f"Algorithm context: Google update published on {update_signal.day.isoformat()} (`{update_signal.title}`). "
            "Treat as a contributing context signal and validate only where cluster/page behavior changed after the update date."
        )
        if algorithm_conf is not None:
            update_line += f" Confidence: {_confidence_bucket(algorithm_conf)}."
        lines.append(update_line)

    campaign_context = _campaign_event_context(external_signals=external_signals, query_scope=query_scope)
    allegro_campaign_events = campaign_context.get("allegro_events", [])
    competitor_campaign_events = campaign_context.get("competitor_events", [])
    campaign_query_events = campaign_context.get("query_events", [])
    allegro_count = len(allegro_campaign_events) if isinstance(allegro_campaign_events, list) else 0
    competitor_count = len(competitor_campaign_events) if isinstance(competitor_campaign_events, list) else 0
    query_count = len(campaign_query_events) if isinstance(campaign_query_events, list) else 0
    if allegro_count >= 3 or competitor_count > 0 or query_count > 0:
        campaign_examples: list[str] = []
        if isinstance(allegro_campaign_events, list):
            for row in allegro_campaign_events[:2]:
                if isinstance(row, ExternalSignal) and str(row.title).strip():
                    campaign_examples.append(f"`{row.title}`")
        if isinstance(competitor_campaign_events, list):
            for row in competitor_campaign_events[:1]:
                if isinstance(row, tuple) and len(row) == 2 and isinstance(row[0], ExternalSignal):
                    campaign_examples.append(f"`{row[0].title}`")
        if isinstance(campaign_query_events, list):
            for row in campaign_query_events[:1]:
                query_name = str(getattr(row, "key", "")).strip()
                if query_name:
                    campaign_examples.append(f"query mover `{query_name}`")
        examples_text = "; ".join(campaign_examples[:3]) if campaign_examples else "active sales campaign signals"
        lines.append(
            "Campaign context: this week included signals such as "
            + examples_text
            + ". Likely impact on demand allocation: medium."
        )
    unified_timeline = _marketplace_timeline_rows(
        additional_context=additional_context,
        campaign_context=campaign_context,
        report_country_code=country_code or "PL",
        max_rows=24,
    )
    if unified_timeline:
        preview = "; ".join(
            f"{str(row.get('day', '')).strip()} [{row.get('country', '')}] {row.get('track', '')}: {row.get('event', '')}"
            for row in unified_timeline[:4]
        )
        if preview:
            lines.append(
                "Weekly market storyline on one timeline: "
                + preview
                + ". Use this timeline as context for demand shifts, not as a standalone root cause."
            )
    market_events_ctx = (additional_context or {}).get("market_event_calendar", {})
    if isinstance(market_events_ctx, dict):
        counts = market_events_ctx.get("counts", {})
        if isinstance(counts, dict):
            current_events = int(counts.get("current", 0) or 0)
            yoy_events = int(counts.get("yoy", 0) or 0)
            delta_events = int(counts.get("delta_vs_yoy", 0) or 0)
            delta_events_pct = float(counts.get("delta_pct_vs_yoy", 0.0) or 0.0)
            if current_events > 0 or yoy_events > 0:
                lines.append(
                    "Marketplace events YoY context: "
                    f"{current_events} events now vs {yoy_events} YoY ({delta_events:+d}; {delta_events_pct:+.1f}%)."
                )

    pulse_line = _top_platform_pulse_line(additional_context)
    if pulse_line:
        lines.append(pulse_line)

    trade_plan_signal = _trade_plan_signal(additional_context)
    trade_plan = (additional_context or {}).get("trade_plan", {})
    if isinstance(trade_plan, dict) and trade_plan.get("enabled"):
        if isinstance(trade_plan_signal, dict):
            summary = str(trade_plan_signal.get("summary", "")).strip()
            statement = str(trade_plan_signal.get("statement", "")).strip()
            confidence = int(trade_plan_signal.get("confidence", 0) or 0)
            if summary:
                lines.append("Trade-plan signal: " + summary)
            if statement:
                lines.append(
                    "Trade-plan interpretation: "
                    + statement
                    + f" Likelihood: {_confidence_bucket(confidence)}."
                )
            overlap_score = int(trade_plan_signal.get("overlap_intensity_score", 0) or 0)
            overlap_campaigns = int(trade_plan_signal.get("overlap_campaigns", 0) or 0)
            overlap_days = int(trade_plan_signal.get("overlap_days", 0) or 0)
            overlap_recency = trade_plan_signal.get("overlap_recency_days")
            if overlap_score > 0:
                lines.append(
                    "Trade-plan overlap intensity: "
                    f"{overlap_score}/100 (campaigns={overlap_campaigns}, overlap days={overlap_days}"
                    + (
                        f", recency={int(overlap_recency)}d)."
                        if isinstance(overlap_recency, (int, float))
                        else ")."
                    )
                )
            yoy_availability_message = str(trade_plan_signal.get("yoy_availability_message", "")).strip()
            if yoy_availability_message:
                lines.append("Trade-plan YoY availability: " + yoy_availability_message)
        channel_split = trade_plan.get("channel_split", [])
        if isinstance(channel_split, list) and channel_split:
            yoy_channel_hypotheses: list[str] = []
            for row in channel_split[:3]:
                if not isinstance(row, dict):
                    continue
                name = str(row.get("channel", "")).strip()
                if not name:
                    continue
                impact = str(row.get("yoy_hypothesis_impact", "")).strip().lower()
                reason = str(row.get("yoy_hypothesis_reason", "")).strip()
                confidence = int(row.get("yoy_hypothesis_confidence", 0) or 0)
                delta_pct_yoy = row.get("delta_spend_pct_vs_yoy")
                if not impact or not reason:
                    continue
                yoy_channel_hypotheses.append(
                    f"`{name}`: {impact} impact"
                    + (
                        f" ({float(delta_pct_yoy):+.2f}% spend vs YoY)"
                        if isinstance(delta_pct_yoy, (int, float))
                        else ""
                    )
                    + f" -> {reason} (confidence {confidence}/100)"
                )
            if yoy_channel_hypotheses:
                lines.append("Trade-plan YoY hypotheses (channels): " + "; ".join(yoy_channel_hypotheses) + ".")
        campaign_rows = trade_plan.get("campaign_rows", [])
        if isinstance(campaign_rows, list) and campaign_rows:
            yoy_campaign_hypotheses_current: list[str] = []
            yoy_campaign_hypotheses_upcoming: list[str] = []
            for row in campaign_rows[:6]:
                if not isinstance(row, dict):
                    continue
                campaign_name = str(row.get("campaign", "")).strip()
                if not campaign_name:
                    continue
                impact = str(row.get("yoy_hypothesis_impact", "")).strip().lower()
                reason = str(row.get("yoy_hypothesis_reason", "")).strip()
                confidence = int(row.get("yoy_hypothesis_confidence", 0) or 0)
                delta_pct_yoy = row.get("delta_spend_pct_vs_yoy")
                if not impact or not reason:
                    continue
                hypothesis_line = (
                    f"`{campaign_name}`: {impact} impact"
                    + (
                        f" ({float(delta_pct_yoy):+.2f}% spend vs YoY)"
                        if isinstance(delta_pct_yoy, (int, float))
                        else ""
                    )
                    + f" -> {reason} (confidence {confidence}/100)"
                )
                in_current_window = (
                    float(row.get("current_spend", 0.0) or 0.0) > 0.0
                    or float(row.get("current_impressions", 0.0) or 0.0) > 0.0
                    or float(row.get("current_clicks", 0.0) or 0.0) > 0.0
                )
                if in_current_window:
                    yoy_campaign_hypotheses_current.append(hypothesis_line)
                else:
                    yoy_campaign_hypotheses_upcoming.append(hypothesis_line)
            if yoy_campaign_hypotheses_current:
                lines.append("Trade-plan YoY hypotheses (campaigns active this week): " + "; ".join(yoy_campaign_hypotheses_current[:3]) + ".")
            if yoy_campaign_hypotheses_upcoming:
                lines.append("Trade-plan YoY hypotheses (upcoming campaigns, not active this week): " + "; ".join(yoy_campaign_hypotheses_upcoming[:2]) + ".")
    elif isinstance(trade_plan, dict):
        errors = trade_plan.get("errors", [])
        if isinstance(errors, list) and errors:
            lines.append("Trade plan context unavailable in this run: " + str(errors[0]).strip())
        elif trade_plan:
            lines.append(
                "Trade plan context unavailable in this run: no rows matched current/previous weekly windows."
            )

    feature_split = (additional_context or {}).get("gsc_feature_split", {})
    if isinstance(feature_split, dict) and feature_split.get("enabled"):
        rows_weekly = feature_split.get("rows_weekly", feature_split.get("rows", []))
        rows_mom = feature_split.get("rows_mom", [])
        feature_overview = feature_split.get("feature_overview", [])
        if not isinstance(rows_weekly, list):
            rows_weekly = []
        if not isinstance(rows_mom, list):
            rows_mom = []
        if not isinstance(feature_overview, list):
            feature_overview = []

        wow_gains, wow_losses = _feature_mover_pairs(
            [row for row in rows_weekly if isinstance(row, dict)],
            delta_key="delta_clicks_vs_previous",
            limit=1,
        )
        mom_gains, mom_losses = _feature_mover_pairs(
            [row for row in rows_mom if isinstance(row, dict)],
            delta_key="delta_clicks_vs_previous",
            limit=1,
        )
        yoy_rows = []
        for row in feature_overview:
            if not isinstance(row, dict):
                continue
            yoy_rows.append(
                {
                    "feature": row.get("feature", ""),
                    "delta_clicks_vs_previous": row.get("yoy_delta_clicks", 0.0),
                }
            )
        yoy_gains, yoy_losses = _feature_mover_pairs(yoy_rows, delta_key="delta_clicks_vs_previous", limit=1)

        serp_parts: list[str] = []
        if wow_gains or wow_losses:
            serp_parts.append(
                "WoW "
                + (f"up: {wow_gains[0]}; " if wow_gains else "")
                + (f"down: {wow_losses[0]}" if wow_losses else "")
            )
        if mom_gains or mom_losses:
            serp_parts.append(
                "MoM "
                + (f"up: {mom_gains[0]}; " if mom_gains else "")
                + (f"down: {mom_losses[0]}" if mom_losses else "")
            )
        if yoy_gains or yoy_losses:
            serp_parts.append(
                "YoY "
                + (f"up: {yoy_gains[0]}; " if yoy_gains else "")
                + (f"down: {yoy_losses[0]}" if yoy_losses else "")
            )
        if serp_parts:
            lines.append(
                "SERP appearance deltas (GSC searchAppearance) indicate traffic moved between result types, "
                "not one uniform decline. "
                + " | ".join(serp_parts)
                + "."
            )
        for row in _serp_unified_compact_table_lines(additional_context, limit=4):
            lines.append(row)

    updates_timeline = _google_updates_timeline_text(additional_context)
    if updates_timeline:
        lines.append(updates_timeline)

    case_study_context = _serp_case_study_text(additional_context)
    if case_study_context:
        lines.append(
            case_study_context
            + " Use this as hypothesis support, not as standalone root cause."
        )

    weekly_news = (additional_context or {}).get("weekly_news_digest", {})
    if isinstance(weekly_news, dict) and weekly_news.get("enabled"):
        rows = weekly_news.get("rows", [])
        headlines: list[str] = []
        if isinstance(rows, list):
            for row in rows[:3]:
                if not isinstance(row, dict):
                    continue
                title = str(row.get("title", "")).strip()
                if title:
                    headlines.append(f"`{title}`")
        if headlines:
            news_yoy_text = ""
            current_count = weekly_news.get("total_count")
            yoy_count = weekly_news.get("total_count_yoy")
            if isinstance(current_count, (int, float)) and isinstance(yoy_count, (int, float)):
                news_yoy_text = " YoY volume context: " + _format_count_yoy_change(float(current_count), float(yoy_count))
            lines.append(
                "SEO/GEO publication context: "
                + "; ".join(headlines)
                + "."
                + news_yoy_text
                + " Direct impact this week is likely low-to-medium; use as supporting context only."
            )

    temp_diff = float(weather_summary.get("avg_temp_diff_c", 0.0))
    precip_change = float(weather_summary.get("precip_change_pct", 0.0))
    temp_diff_yoy = float(weather_summary.get("avg_temp_diff_yoy_c", 0.0))
    precip_change_yoy = float(weather_summary.get("precip_change_pct_yoy", 0.0))
    lines.append(
        "Weather context: week-over-week conditions changed "
        f"({temp_diff:+.1f}C; precipitation {precip_change:+.1f}%). "
        f"YoY weather delta: {temp_diff_yoy:+.1f}C; precipitation {precip_change_yoy:+.1f}%. "
        "This supports demand-timing interpretation "
        "rather than a direct technical SEO issue. "
        + (
            f"Likely impact on demand timing: {_confidence_bucket(weather_conf)}."
            if weather_conf is not None
            else "Likely impact on demand timing: Medium."
        )
    )

    macro = (additional_context or {}).get("macro", {})
    if isinstance(macro, dict):
        nbp = macro.get("nbp_fx", {})
        if isinstance(nbp, dict):
            eur = nbp.get("eur_pln", {})
            usd = nbp.get("usd_pln", {})
            if isinstance(eur, dict) and isinstance(usd, dict):
                eur_delta = float(eur.get("delta_pct_vs_previous", 0.0) or 0.0)
                usd_delta = float(usd.get("delta_pct_vs_previous", 0.0) or 0.0)
                if abs(eur_delta) >= 1.0 or abs(usd_delta) >= 1.0:
                    lines.append(
                        "Macro context: FX moved materially vs previous week "
                        f"(EUR/PLN {eur_delta:+.2f}%, USD/PLN {usd_delta:+.2f}%). "
                        "This can affect price competitiveness and conversion in import-heavy categories."
                    )
    macro_backdrop = (additional_context or {}).get("macro_backdrop", {})
    if isinstance(macro_backdrop, dict):
        macro_rows = macro_backdrop.get("rows", {})
        if isinstance(macro_rows, dict):
            inflation = macro_rows.get("inflation_cpi_pct", {})
            unemployment = macro_rows.get("unemployment_pct", {})
            if isinstance(inflation, dict) and inflation.get("latest_value") is not None:
                inf_latest = float(inflation.get("latest_value", 0.0) or 0.0)
                inf_prev = float(inflation.get("previous_value", 0.0) or 0.0)
                lines.append(
                    "Macro backdrop (annual): "
                    f"inflation is {inf_latest:.2f}% vs {inf_prev:.2f}% in prior available year. "
                    "Higher inflation usually means more price-sensitive shopping behavior."
                )
            if isinstance(unemployment, dict) and unemployment.get("latest_value") is not None:
                un_latest = float(unemployment.get("latest_value", 0.0) or 0.0)
                lines.append(
                    "Labor-market backdrop: "
                    f"unemployment is {un_latest:.2f}% in latest available year; "
                    "treat as supporting context for medium-term demand."
                )

    promo_radar = (additional_context or {}).get("competitor_promo_radar", {})
    if isinstance(promo_radar, dict) and promo_radar.get("enabled"):
        promo_rows = promo_radar.get("rows", [])
        if isinstance(promo_rows, list) and promo_rows:
            top_titles: list[str] = []
            for row in promo_rows[:2]:
                if not isinstance(row, dict):
                    continue
                title = str(row.get("title", "")).strip()
                if title:
                    top_titles.append(f"`{title}`")
            lines.append(
                "Competitor promo radar: "
                f"{len(promo_rows)} relevant promo mentions in the analyzed period"
                + (f" (examples: {', '.join(top_titles)})." if top_titles else ".")
                + " This can redirect clicks between marketplaces."
            )

    long_window = (additional_context or {}).get("long_window_context", {})
    if isinstance(long_window, dict) and long_window.get("enabled"):
        kpi = long_window.get("kpi", {})
        if isinstance(kpi, dict):
            lw_clicks_pct = kpi.get("clicks_delta_pct_vs_previous")
            lw_clicks_yoy_pct = kpi.get("clicks_delta_pct_vs_yoy")
            if isinstance(lw_clicks_pct, (int, float)) and isinstance(lw_clicks_yoy_pct, (int, float)):
                lines.append(
                    "Broader 28d context: "
                    f"clicks are {float(lw_clicks_pct):+.2f}% vs previous 28d and {float(lw_clicks_yoy_pct):+.2f}% YoY. "
                    "This helps separate short weekly noise from longer trend direction."
                )

    tracker = (additional_context or {}).get("hypothesis_tracker", {})
    if isinstance(tracker, dict):
        counts = tracker.get("counts", {})
        if isinstance(counts, dict):
            new_n = int(counts.get("new", 0) or 0)
            confirmed_n = int(counts.get("confirmed", 0) or 0)
            rejected_n = int(counts.get("rejected", 0) or 0)
            lines.append(
                "Hypothesis continuity: "
                f"{confirmed_n} recurring signals, {new_n} new signals, {rejected_n} rejected since last run. "
                "Recurring signals are usually more reliable than one-off observations."
            )

    forecast_start = str(weather_summary.get("forecast_start", "")).strip()
    forecast_end = str(weather_summary.get("forecast_end", "")).strip()
    if forecast_start and forecast_end:
        lines.append(
            "**Forward 7d**: weather forecast "
            f"{forecast_start} to {forecast_end} (avg {float(weather_summary.get('forecast_avg_temp_c', 0.0)):+.1f}C, "
            f"precip {float(weather_summary.get('forecast_precip_mm', 0.0)):.1f}mm). "
            "Use this as timing context for near-term category demand."
        )

    confirmed_points: list[str] = []
    hypothesis_points: list[str] = []
    confirmed_points.append(
        f"GSC clicks {_fmt_signed_compact(current.clicks - previous.clicks)} WoW ({wow_pct}) with near-stable CTR ({ctr_wow_pp:+.2f} pp)."
    )
    if page_name_line:
        confirmed_points.append(page_name_line)
    if isinstance(brand_proxy, dict):
        confirmed_points.append(
            "Brand in GSC: clicks "
            f"{float(brand_proxy.get('delta_pct_vs_previous', 0.0)):+.2f}% WoW, "
            f"{float(brand_proxy.get('delta_pct_vs_yoy', 0.0)):+.2f}% YoY."
        )
    if brand_note:
        hypothesis_points.append(brand_note)
    if isinstance(brand_ads_hyp, dict):
        hypothesis_points.append(
            str(brand_ads_hyp.get("statement", "")).strip()
            + f" (confidence: {_confidence_bucket(brand_ads_hyp.get('confidence', 0))})."
        )
    if update_signal:
        hypothesis_points.append(
            f"Post-update SERP behavior check after {update_signal.day.isoformat()} is required before root-cause lock."
        )
    if isinstance(trade_plan_signal, dict):
        hypothesis_points.append(
            "Paid-channel timing may reallocate brand/category clicks between organic and ads."
        )
    open_questions: list[str] = []
    if not brand_enabled and not isinstance(brand_proxy, dict):
        open_questions.append(
            "Do we have a complete brand-demand baseline (Google Trends brand series or GSC brand proxy) for this market?"
        )
    if isinstance(trade_plan_signal, dict):
        trade_conf = int(trade_plan_signal.get("confidence", 0) or 0)
        if trade_conf < 70:
            open_questions.append(
                "How much of weekly movement is explained by campaign/paid overlap versus underlying organic demand rotation?"
            )
    else:
        open_questions.append(
            "Do we need additional campaign or paid-channel inputs to explain this week's movement with higher confidence?"
        )
    source_errors = (additional_context or {}).get("errors", [])
    if isinstance(source_errors, list) and source_errors:
        open_questions.append(
            "Which missing/degraded external source had the biggest impact on interpretation confidence this week?"
        )
    if not open_questions:
        open_questions.append(
            "Which single hypothesis should be falsified first in the next run before escalation?"
        )

    if confirmed_points or hypothesis_points:
        lines.append("")
        lines.append("**Confirmed facts vs plausible drivers vs open questions**")
        if confirmed_points:
            lines.append("Confirmed facts from data:")
            for item in confirmed_points[:2]:
                lines.append(f"- {item}")
        if hypothesis_points:
            lines.append("Plausible drivers (need validation):")
            for item in hypothesis_points[:2]:
                lines.append(f"- {item}")
        lines.append("Open questions for next run:")
        for item in open_questions[:2]:
            lines.append(f"- {item}")

    # Dynamic driver table (only active/contextual drivers for this run).
    driver_rows: list[tuple[str, str, str, str, str]] = []
    demand_direction = "Down" if (current.clicks - previous.clicks) < 0 else "Up"
    demand_conf = seasonality_conf if seasonality_conf is not None else 70
    demand_impact = "High" if abs(_ratio_delta(current.clicks, previous.clicks)) >= 0.03 else "Medium"
    driver_rows.append(
        (
            "Demand mix",
            demand_direction,
            f"Organic clicks {wow_pct} WoW; efficiency change limited (CTR {ctr_wow_pp:+.2f} pp).",
            _confidence_bucket(demand_conf),
            demand_impact,
        )
    )
    if page_name_line:
        driver_rows.append(
            (
                "Page mix",
                "Down" if home_wow < 0 else "Up",
                page_name_line,
                _confidence_bucket(template_conf if template_conf is not None else 72),
                "High",
            )
        )
    if isinstance(brand_ads_hyp, dict):
        driver_rows.append(
            (
                "Brand paid overlap",
                "Mixed",
                str(brand_ads_hyp.get("statement", "")).strip(),
                _confidence_bucket(brand_ads_hyp.get("confidence", 0)),
                "Medium",
            )
        )
    if isinstance(trade_plan_signal, dict):
        driver_rows.append(
            (
                "Trade plan timing",
                "Mixed",
                str(trade_plan_signal.get("statement", "")).strip(),
                _confidence_bucket(trade_plan_signal.get("confidence", 0)),
                "Medium",
            )
        )
    if update_signal:
        driver_rows.append(
            (
                "Algorithm context",
                "Mixed",
                f"Post-update monitoring needed after {update_signal.day.isoformat()}.",
                _confidence_bucket(algorithm_conf if algorithm_conf is not None else 68),
                "Medium",
            )
        )
    if abs(temp_diff) >= 2.0 or abs(precip_change) >= 20.0:
        driver_rows.append(
            (
                "Weather timing",
                "Mixed",
                f"Weather shift vs previous week: temp {temp_diff:+.1f}C, precip {precip_change:+.1f}%.",
                _confidence_bucket(weather_conf if weather_conf is not None else 64),
                "Low-Med",
            )
        )

    if driver_rows:
        top_driver_bits = [
            f"{row[0]} ({row[1].lower()}, confidence {row[3]})"
            for row in driver_rows[:3]
        ]
        lines.append("")
        lines.append("Top drivers this week: " + "; ".join(top_driver_bits) + ".")
    top_action_hypotheses = _top_hypotheses_for_actions(hypotheses, limit=3)
    if top_action_hypotheses:
        priority_bits = []
        for row in top_action_hypotheses:
            category = str(row.get("category", "Unknown")).strip() or "Unknown"
            priority = int(row.get("driver_priority_score", 0) or 0)
            impact = int(row.get("impact_score", 0) or 0)
            controllability = int(row.get("controllability_score", 0) or 0)
            confidence = int(row.get("confidence", 0) or 0)
            priority_bits.append(
                f"`{category}`={priority}/100 (impact {impact}/5 x confidence {confidence}/100 x controllability {controllability}/5)"
            )
        if priority_bits:
            lines.append(
                "Driver priority model (impact x confidence x controllability): "
                + "; ".join(priority_bits[:3])
                + "."
            )
        lines.append("Uncertainty framing (actionable):")
        for row in top_action_hypotheses[:3]:
            category = str(row.get("category", "Unknown")).strip() or "Unknown"
            confidence = int(row.get("confidence", 0) or 0)
            lines.append(
                f"- `{category}`: {_uncertainty_action_template(confidence)}"
            )
        validation_plan = _next_week_validation_plan_lines(hypotheses, limit=3)
        if validation_plan:
            lines.append("Validation plan (next week, top hypotheses):")
            lines.extend(validation_plan[:3])
        counterfactual_lines = _counterfactual_check_lines(hypotheses, limit=3)
        if counterfactual_lines:
            lines.append("Counterfactual checks (what would disprove top hypotheses?):")
            lines.extend(counterfactual_lines[:3])
    escalation_gate = _technical_seo_escalation_gate(
        totals=totals,
        hypotheses=hypotheses,
        contradiction_count=len(contradiction_rows),
    )
    lines.append(
        "Technical SEO escalation gate: "
        f"{str(escalation_gate.get('status', '')).strip()} "
        f"-> {str(escalation_gate.get('reason', '')).strip()} "
        f"Next action: {str(escalation_gate.get('next_action', '')).strip()}"
    )

    lines.append("")
    lines.append("**Reasoning ledger (facts -> hypotheses -> validation)**")
    lines.append(
        "- Facts observed: "
        f"clicks {_fmt_signed_compact(current.clicks - previous.clicks)} WoW ({wow_pct}), "
        f"CTR {ctr_wow_pp:+.2f} pp WoW, avg position {_position_delta_label(pos_wow_delta)} WoW. [E1]"
    )
    if hypotheses:
        ranked_hypotheses = "; ".join(
            f"{str(row.get('category', 'Unknown')).strip()} ({_confidence_bucket(row.get('confidence', 0))})"
            for row in hypotheses[:3]
            if isinstance(row, dict)
        )
        if ranked_hypotheses:
            lines.append("- Primary hypotheses: " + ranked_hypotheses + ".")
    alt_explanations: list[str] = []
    if query_count == 0 and allegro_count == 0 and competitor_count == 0:
        alt_explanations.append("campaign impact may be underestimated because no explicit campaign mentions were captured")
    if brand_enabled is False and not isinstance(brand_proxy, dict):
        alt_explanations.append("brand-demand baseline is weak (Google Trends + GSC brand proxy unavailable)")
    if alt_explanations:
        lines.append("- Alternative explanations to falsify: " + "; ".join(alt_explanations[:2]) + ".")
    source_errors = (additional_context or {}).get("errors", [])
    if isinstance(source_errors, list) and source_errors:
        sanitized = [_sanitize_data_gap_message(row) for row in source_errors[:3]]
        sanitized = [row for row in sanitized if row]
        if sanitized:
            lines.append("- Data gaps affecting certainty: " + "; ".join(sanitized[:2]) + ".")
    lines.append("- Escalation logic is applied only after validation checks and contradiction reconciliation.")

    return _compact_manager_section(lines, max_lines=40)


def _focus_terms_from_query_scope(query_scope: AnalysisResult | None) -> list[str]:
    if query_scope is None:
        return []

    stop_words = {
        "allegro",
        "allegropl",
        "listing",
        "kategoria",
        "oferta",
        "www",
        "https",
        "query",
        "smart",
        "strona",
        "glowna",
    }

    counts: dict[str, int] = {}
    for row in query_scope.top_losers + query_scope.top_winners:
        text = _normalize_text(row.key)
        for token in re.findall(r"[a-z0-9]{4,}", text):
            if token in stop_words:
                continue
            if token.isdigit():
                continue
            counts[token] = counts.get(token, 0) + 1

    ranked = sorted(counts.items(), key=lambda item: item[1], reverse=True)
    return [token for token, _ in ranked[:20]]


def _collect_context_notes(rows: list[dict[str, object]], fallback_key: str = "excerpt") -> list[str]:
    notes: list[str] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        highlights = row.get("highlights", [])
        if isinstance(highlights, list):
            for note in highlights:
                text = str(note).strip()
                if text:
                    notes.append(text)
        if not highlights:
            fallback = str(row.get(fallback_key, "")).strip()
            if fallback:
                notes.append(fallback)
    return notes


def _find_overlap_terms(texts: list[str], terms: list[str]) -> list[str]:
    if not texts or not terms:
        return []
    blob = " ".join(_normalize_text(text) for text in texts if text)
    overlaps = [term for term in terms if term in blob]
    return overlaps[:6]


def _build_reasoning_hypotheses(
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    ferie_context: dict[str, object],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
    senuto_summary: dict[str, float] | None,
    senuto_error: str | None,
) -> list[dict[str, object]]:
    current = totals["current_28d"]
    previous = totals["previous_28d"]
    hypotheses: list[dict[str, object]] = []
    query_scope = _find_scope(scope_results, "query")

    if query_scope:
        winter_terms = (
            "fajerwerk",
            "petard",
            "sanki",
            "odsniez",
            "snieg",
            "kalendarz 2026",
            "dzien babci",
            "dziadka",
        )
        event_terms = (
            "wosp",
            "wieniawa",
            "nocowanka",
            "licytac",
            "walentyn",
        )
        winter_drop = _sum_losses(query_scope.top_losers, winter_terms)
        event_gain = _sum_gains(query_scope.top_winners, event_terms)
        ctr_up = current.ctr >= previous.ctr
        position_better = current.position <= previous.position

        if winter_drop >= 5000:
            score = 86 if (ctr_up and position_better) else 73
            evidence = [
                f"Drop in winter-intent query clicks vs previous week: ~{_fmt_int(winter_drop)}.",
                f"CTR current vs prev: {_pct(current.ctr)} vs {_pct(previous.ctr)}.",
                f"Avg position current vs prev: {current.position:.2f} vs {previous.position:.2f}.",
            ]
            hypotheses.append(
                {
                    "category": "Seasonality",
                    "confidence": score,
                    "thesis": (
                        "Declines are concentrated in winter-intent demand clusters (for example sled/firework-type intents), "
                        "while CTR and average position improved. This pattern is consistent with post-peak seasonality "
                        "and does not indicate broad ranking/indexing degradation."
                    ),
                    "evidence": evidence,
                    "owner": "SEO + Merchandising",
                }
            )

        if event_gain >= 3000:
            hypotheses.append(
                {
                    "category": "Events",
                    "confidence": 82,
                    "thesis": "Event-driven demand topics are offsetting part of the natural seasonal declines.",
                    "evidence": [
                        f"Increase in event-query clicks vs previous week: ~{_fmt_int(event_gain)}.",
                        "Top winners include queries tied to current events.",
                    ],
                    "owner": "SEO + Content",
                }
            )

    yoy_ferie = ferie_context.get("yoy_comparison", {})
    ferie_delta_pp = float(yoy_ferie.get("avg_daily_delta_pp", 0.0))
    if abs(ferie_delta_pp) >= 0.3:
        top_region = ""
        yoy_rows = yoy_ferie.get("rows")
        if isinstance(yoy_rows, list) and yoy_rows:
            row = yoy_rows[0]
            if isinstance(row, dict):
                top_region = (
                    f"{row.get('name', 'Region')}: {int(row.get('current_days', 0))} days vs "
                    f"{int(row.get('yoy_days', 0))} days ({float(row.get('contribution_pp', 0.0)):+.2f} pp)."
                )
        hypotheses.append(
            {
                "category": "Seasonality",
                "confidence": 74,
                "thesis": "Regional winter-break distribution changes YoY comparability for winter demand.",
                "evidence": [
                    f"YoY delta in winter-break exposure (avg daily GMV proxy): {ferie_delta_pp:+.2f} pp.",
                    top_region or "Material regional winter-break difference in the YoY window.",
                ],
                "owner": "SEO + BI",
            }
        )

    temp_diff = float(weather_summary.get("avg_temp_diff_c", 0.0))
    if abs(temp_diff) >= 2.0:
        hypotheses.append(
            {
                "category": "Macro/Weather",
                "confidence": 68,
                "thesis": "Weather anomaly may have shifted demand between product categories.",
                "evidence": [
                    f"Avg temperature diff current vs previous: {temp_diff:+.1f}C.",
                    f"Precipitation change current vs previous: {float(weather_summary.get('precip_change_pct', 0.0)):+.1f}%.",
                ],
                "owner": "SEO + Commercial",
            }
        )

    google_updates = [
        signal
        for signal in external_signals
        if signal.source in {"Google Search Status", "Google Search Central Blog"}
        and "update" in _normalize_text(f"{signal.title} {signal.details}")
    ]
    if google_updates:
        latest = max(google_updates, key=lambda row: row.day)
        hypotheses.append(
            {
                "category": "Algorithm",
                "confidence": 71,
                "thesis": "Google algorithm changes may have affected selected query/page clusters.",
                "evidence": [
                    f"Latest update: {latest.day.isoformat()} ({latest.title}).",
                    "Query/page segment monitoring is required after the update date.",
                ],
                "owner": "SEO",
            }
        )

    if senuto_summary:
        avg_delta = float(senuto_summary.get("avg_delta_pct", 0.0))
        if avg_delta < -5:
            hypotheses.append(
                {
                    "category": "SEO Visibility",
                    "confidence": 79,
                    "thesis": "Senuto visibility decline suggests a ranking component, not only seasonality.",
                    "evidence": [
                        f"Senuto avg delta vs previous week: {avg_delta:+.2f}%.",
                        f"Senuto latest delta vs previous week: {float(senuto_summary.get('latest_delta_pct', 0.0)):+.2f}%.",
                    ],
                    "owner": "SEO",
                }
            )
    elif senuto_error:
        hypotheses.append(
            {
                "category": "Data quality",
                "confidence": 92,
                "thesis": "Missing Senuto data lowers confidence in ranking diagnosis.",
                "evidence": [f"Senuto error: {senuto_error}"],
                "owner": "SEO Ops",
            }
        )

    if segment_diagnostics:
        non_brand = _top_negative_segment(segment_diagnostics.get("brand_non_brand"), min_clicks=3000.0)
        if non_brand and str(non_brand.get("segment")) == "non_brand":
            hypotheses.append(
                {
                    "category": "Demand mix",
                    "confidence": 76,
                    "thesis": "Non-brand decline is larger than brand decline, which may indicate competitive pressure beyond branded traffic.",
                    "evidence": [
                        f"Non-brand delta vs prev: {_fmt_signed_int(non_brand.get('delta_vs_previous', 0.0))} ({_signed_pct(float(non_brand.get('delta_pct_vs_previous', 0.0)))}).",
                    ],
                    "owner": "SEO + Acquisition",
                }
            )

        device_loser = _top_negative_segment(segment_diagnostics.get("device"), min_clicks=3000.0)
        if device_loser:
            hypotheses.append(
                {
                    "category": "Technical/UX",
                    "confidence": 67,
                    "thesis": "One device segment drives the decline more than others and needs targeted UX/SEO audit.",
                    "evidence": [
                        f"Device {device_loser.get('segment')}: delta vs prev {_fmt_signed_int(device_loser.get('delta_vs_previous', 0.0))} ({_signed_pct(float(device_loser.get('delta_pct_vs_previous', 0.0)))}).",
                    ],
                    "owner": "SEO + Web Performance",
                }
            )

        template_loser = _top_negative_segment(segment_diagnostics.get("page_template"), min_clicks=3000.0)
        if template_loser:
            hypotheses.append(
                {
                    "category": "Page Name performance",
                    "confidence": 72,
                    "thesis": "Decline concentrated in one page-name segment suggests page-level issue or section-specific demand drop.",
                    "evidence": [
                        f"Page Name segment {template_loser.get('segment')}: delta vs prev {_fmt_signed_int(template_loser.get('delta_vs_previous', 0.0))} ({_signed_pct(float(template_loser.get('delta_pct_vs_previous', 0.0)))}).",
                    ],
                    "owner": "SEO + Product",
                }
            )

    if additional_context:
        focus_terms = _focus_terms_from_query_scope(query_scope)
        pagespeed = additional_context.get("pagespeed", {})
        if isinstance(pagespeed, dict):
            mobile = pagespeed.get("mobile", {})
            if isinstance(mobile, dict):
                lcp = float(mobile.get("lcp_ms", 0.0))
                inp = float(mobile.get("inp_ms", 0.0))
                cls = float(mobile.get("cls", 0.0))
                if lcp >= 2500 or inp >= 200 or cls >= 0.1:
                    hypotheses.append(
                        {
                            "category": "Technical/UX",
                            "confidence": 66,
                            "thesis": "Mobile CWV are near/below \"good\" thresholds and can depress organic performance.",
                            "evidence": [
                                f"PageSpeed mobile LCP={lcp:.0f}ms, INP={inp:.0f}ms, CLS={cls:.2f}.",
                            ],
                            "owner": "Web Performance",
                        }
                    )

        macro = additional_context.get("macro", {})
        if isinstance(macro, dict):
            nbp = macro.get("nbp_fx", {})
            if isinstance(nbp, dict):
                for code_key, label in (("eur_pln", "EUR/PLN"), ("usd_pln", "USD/PLN")):
                    row = nbp.get(code_key)
                    if isinstance(row, dict):
                        delta = float(row.get("delta_pct_vs_previous", 0.0))
                        if abs(delta) >= 2.0:
                            hypotheses.append(
                                {
                                    "category": "Macro",
                                    "confidence": 63,
                                    "thesis": "FX changes can affect pricing/offers and demand in selected categories.",
                                    "evidence": [
                                        f"{label} avg delta vs previous week: {delta:+.2f}%.",
                                    ],
                                    "owner": "Commercial + Pricing",
                                }
                            )
                            break

        product_trends = additional_context.get("product_trends", {})
        if isinstance(product_trends, dict):
            trend_rows = product_trends.get("top_yoy_non_brand", [])
            fallback_source = False
            if not (isinstance(trend_rows, list) and trend_rows):
                trend_rows = _fallback_non_brand_yoy_from_gsc(
                    query_scope=query_scope,
                    top_rows=10,
                )
                fallback_source = bool(trend_rows)
            if isinstance(trend_rows, list) and trend_rows:
                current_sum = 0.0
                previous_sum = 0.0
                for row in trend_rows:
                    if not isinstance(row, dict):
                        continue
                    current_sum += float(row.get("current_value", 0.0))
                    previous_sum += float(row.get("previous_value", 0.0))
                delta = current_sum - previous_sum
                delta_pct = (delta / previous_sum * 100.0) if previous_sum else 0.0

                lead = next((row for row in trend_rows if isinstance(row, dict)), None)
                lead_text = ""
                if isinstance(lead, dict):
                    lead_text = (
                        f"Top trend: {lead.get('trend', '')} "
                        f"({_fmt_int(lead.get('current_value', 0.0))} vs {_fmt_int(lead.get('previous_value', 0.0))})."
                    )
                confidence = 74 if abs(delta_pct) >= 10 else 66
                direction = "upward" if delta >= 0 else "downward"
                hypotheses.append(
                    {
                        "category": "Non-brand product trends",
                        "confidence": confidence,
                        "thesis": (
                            "Non-brand product trend sheets indicate "
                            f"{direction} demand pressure vs previous year, which should be considered in SEO diagnosis."
                        ),
                        "evidence": [
                            f"Top non-brand trends sum: current={_fmt_int(current_sum)}, previous={_fmt_int(previous_sum)}, delta={_fmt_signed_int(delta)} ({delta_pct:+.2f}%).",
                            "Source: GSC query fallback (sheet YoY unavailable)." if fallback_source else "Source: product trend sheets.",
                            lead_text or "Top non-brand trend rows were detected in the configured sheets.",
                        ],
                        "owner": "SEO + Category/Merchandising",
                    }
                )

        senuto_intelligence = additional_context.get("senuto_intelligence", {})
        if isinstance(senuto_intelligence, dict) and senuto_intelligence.get("enabled"):
            competitors = senuto_intelligence.get("competitors_overview", [])
            if isinstance(competitors, list) and competitors:
                top_comp = next((row for row in competitors if isinstance(row, dict)), None)
                if isinstance(top_comp, dict):
                    hypotheses.append(
                        {
                            "category": "Competitive pressure",
                            "confidence": 67,
                            "thesis": (
                                "Competitor overlap indicates elevated competitive pressure in shared non-brand SERP space."
                            ),
                            "evidence": [
                                f"Top competitor overlap domain: {top_comp.get('domain', '')}.",
                                f"Common keywords: {_fmt_int(top_comp.get('common_keywords', 0.0))}.",
                                f"Top10 keywords (competitor): {_fmt_int(top_comp.get('top10_current', 0.0))}.",
                            ],
                            "owner": "SEO + Category Owners",
                        }
                    )

            direct_answers = senuto_intelligence.get("direct_answers", [])
            if isinstance(direct_answers, list) and direct_answers:
                hypotheses.append(
                    {
                        "category": "SERP features",
                        "confidence": 64,
                        "thesis": (
                            "Direct-answer SERP features are present for tracked queries and may reduce classic blue-link CTR."
                        ),
                        "evidence": [
                            f"Direct-answer rows in Senuto snapshot: {_fmt_int(len(direct_answers))}.",
                        ],
                        "owner": "SEO",
                    }
                )

            seasonality = senuto_intelligence.get("seasonality", {})
            if isinstance(seasonality, dict):
                trend_values = seasonality.get("trend_values", [])
                if isinstance(trend_values, list) and trend_values:
                    month_idx = date.today().month - 1
                    current_idx = (
                        float(trend_values[month_idx])
                        if 0 <= month_idx < len(trend_values)
                        else 0.0
                    )
                    hypotheses.append(
                        {
                            "category": "Seasonality",
                            "confidence": 73,
                            "thesis": (
                                "Senuto monthly seasonality supports demand-timing interpretation and should be read together with weekly weather shifts."
                            ),
                            "evidence": [
                                f"Senuto current-month seasonality index: {current_idx:.2f}.",
                                f"Senuto peak month/value: M{int(seasonality.get('peak_month', 0) or 0)} / {float(seasonality.get('peak_value', 0.0) or 0.0):.2f}.",
                                f"Senuto low month/value: M{int(seasonality.get('low_month', 0) or 0)} / {float(seasonality.get('low_value', 0.0) or 0.0):.2f}.",
                            ],
                            "owner": "SEO + BI",
                        }
                    )

        updates_timeline = additional_context.get("google_updates_timeline", {})
        if isinstance(updates_timeline, dict) and updates_timeline.get("enabled"):
            summary = updates_timeline.get("summary", {})
            if isinstance(summary, dict):
                current_30d = int(summary.get("count_current_30d", 0) or 0)
                previous_30d = int(summary.get("count_previous_30d", 0) or 0)
                yoy_30d = int(summary.get("count_yoy_30d", 0) or 0)
                latest_date = str(summary.get("latest_update_date", "")).strip()
                latest_title = str(summary.get("latest_update_title", "")).strip()
                if current_30d > 0:
                    delta_vs_previous = current_30d - previous_30d
                    confidence = 62
                    if abs(delta_vs_previous) >= 3:
                        confidence += 8
                    if yoy_30d > 0 and abs(current_30d - yoy_30d) >= 3:
                        confidence += 4
                    hypotheses.append(
                        {
                            "category": "Algorithm/SERP context",
                            "confidence": min(82, confidence),
                            "thesis": (
                                "Google update cadence in recent weeks can influence SERP behavior and click allocation; "
                                "treat update timing as a plausible context layer for weekly movement."
                            ),
                            "evidence": [
                                f"Update timeline 30d: {current_30d} vs previous 30d {previous_30d}, vs YoY 30d {yoy_30d}.",
                                (
                                    f"Latest tracked update signal: {latest_date} ({latest_title})."
                                    if latest_date and latest_title
                                    else "Latest update signal date was not available in this run."
                                ),
                            ],
                            "owner": "SEO",
                        }
                    )

        case_studies = additional_context.get("serp_case_studies", {})
        if isinstance(case_studies, dict) and case_studies.get("enabled"):
            summary = case_studies.get("summary", {})
            if isinstance(summary, dict):
                topic_counts = summary.get("topic_counts_13m", {})
                top_topics: list[str] = []
                if isinstance(topic_counts, dict):
                    ranked_topics = sorted(
                        [
                            (str(topic).strip(), int(count or 0))
                            for topic, count in topic_counts.items()
                            if str(topic).strip()
                        ],
                        key=lambda item: item[1],
                        reverse=True,
                    )
                    top_topics = [topic for topic, _ in ranked_topics[:3]]
                current_30d = int(summary.get("count_current_30d", 0) or 0)
                previous_30d = int(summary.get("count_previous_30d", 0) or 0)
                total_13m = int(summary.get("total_count_13m", 0) or 0)
                latest_case_date = str(summary.get("latest_case_date", "")).strip()
                latest_case_title = str(summary.get("latest_case_title", "")).strip()
                if total_13m > 0:
                    confidence = 60
                    if current_30d >= previous_30d:
                        confidence += 6
                    if any("ctr" in _normalize_text(topic) or "serp" in _normalize_text(topic) for topic in top_topics):
                        confidence += 8
                    hypotheses.append(
                        {
                            "category": "SERP behavior context",
                            "confidence": min(82, confidence),
                            "thesis": (
                                "External case studies repeatedly describe CTR and placement-mix shifts across SERP features, "
                                "which supports a traffic-allocation interpretation before technical root-cause escalation."
                            ),
                            "evidence": [
                                "Top recurring topics (13M): "
                                + (", ".join(top_topics) if top_topics else "CTR and SERP feature changes")
                                + ".",
                                f"Case-study volume 30d: {current_30d} vs previous 30d {previous_30d}; total in 13M: {total_13m}.",
                                (
                                    f"Latest relevant case signal: {latest_case_date} ({latest_case_title})."
                                    if latest_case_date and latest_case_title
                                    else "Latest case-study title was not available in this run."
                                ),
                            ],
                            "owner": "SEO + BI",
                        }
                    )

        seo_presentations = additional_context.get("seo_presentations", {})
        if isinstance(seo_presentations, dict) and seo_presentations.get("enabled"):
            highlights = seo_presentations.get("highlights", [])
            if isinstance(highlights, list) and highlights:
                top_notes = [
                    str(row.get("note", "")).strip()
                    for row in highlights[:2]
                    if isinstance(row, dict) and str(row.get("note", "")).strip()
                ]
                years_covered = [
                    str(row.get("year", "")).strip()
                    for row in seo_presentations.get("years", [])
                    if isinstance(row, dict) and str(row.get("year", "")).strip()
                ]
                years_label = ", ".join(years_covered[:2]) if years_covered else "current and previous year"
                hypotheses.append(
                    {
                        "category": "Internal initiatives",
                        "confidence": 61,
                        "thesis": "Recent SEO team presentation topics may explain part of the observed movement in monitored segments.",
                        "evidence": [
                            f"Presentation archive analyzed for years: {years_label}.",
                            top_notes[0] if top_notes else "Highlights extracted from team presentations.",
                        ],
                        "owner": "SEO Team",
                    }
                )

        historical_reports = additional_context.get("historical_reports", {})
        if isinstance(historical_reports, dict) and historical_reports.get("enabled"):
            recent_reports = historical_reports.get("recent_reports", [])
            recent_notes = []
            if isinstance(recent_reports, list):
                recent_notes = _collect_context_notes(recent_reports)
            yoy_report = historical_reports.get("yoy_report", {})
            if isinstance(yoy_report, dict):
                recent_notes.extend(
                    _collect_context_notes([yoy_report], fallback_key="excerpt")
                )

            overlaps = _find_overlap_terms(recent_notes, focus_terms)
            if recent_notes:
                thesis = (
                    "Current query/page shifts are consistent with themes highlighted in prior weekly reports."
                    if overlaps
                    else "Historical reports add context, but overlap with current movers is limited."
                )
                confidence = 72 if overlaps else 58
                evidence = [
                    f"Recent reports analyzed: {len(recent_reports)}.",
                ]
                if isinstance(yoy_report, dict) and yoy_report.get("id"):
                    evidence.append(
                        f"YoY report loaded: {str(yoy_report.get('date', '')).strip() or 'available'}."
                    )
                if overlaps:
                    evidence.append(
                        f"Shared terms between history and current movers: {', '.join(overlaps[:5])}."
                    )
                hypotheses.append(
                    {
                        "category": "Continuity",
                        "confidence": confidence,
                        "thesis": thesis,
                        "evidence": evidence,
                        "owner": "SEO Team",
                    }
                )

        status_log = additional_context.get("status_log", {})
        if isinstance(status_log, dict) and status_log.get("enabled"):
            status_entries = status_log.get("entries", [])
            notes: list[str] = []
            latest_date = ""
            if isinstance(status_entries, list):
                for row in status_entries[:8]:
                    if not isinstance(row, dict):
                        continue
                    topic = str(row.get("topic", "")).strip()
                    summary = str(row.get("summary", "")).strip()
                    if topic:
                        notes.append(topic)
                    if summary:
                        notes.append(summary)
                    if not latest_date:
                        latest_date = str(row.get("date", "")).strip()
            overlaps = _find_overlap_terms(notes, focus_terms)
            if notes:
                confidence = 70 if overlaps else 55
                thesis = (
                    "Recently discussed status topics align with this week's winners/losers."
                    if overlaps
                    else "Status updates are available, but direct linkage to current movers is weak."
                )
                evidence = [
                    f"Status entries analyzed: {len(status_entries) if isinstance(status_entries, list) else 0}.",
                ]
                if latest_date:
                    evidence.append(f"Latest status date in source: {latest_date}.")
                if overlaps:
                    evidence.append(f"Overlap terms: {', '.join(overlaps[:5])}.")
                hypotheses.append(
                    {
                        "category": "Execution continuity",
                        "confidence": confidence,
                        "thesis": thesis,
                        "evidence": evidence,
                        "owner": "SEO Team + PM",
                    }
                )

        campaign_context = _campaign_event_context(
            external_signals=external_signals,
            query_scope=query_scope,
        )
        allegro_events = campaign_context.get("allegro_events", [])
        competitor_events = campaign_context.get("competitor_events", [])
        query_events = campaign_context.get("query_events", [])
        allegro_count = len(allegro_events) if isinstance(allegro_events, list) else 0
        competitor_count = len(competitor_events) if isinstance(competitor_events, list) else 0
        query_count = len(query_events) if isinstance(query_events, list) else 0
        if allegro_count or competitor_count or query_count:
            evidence: list[str] = [
                f"Campaign signals: allegro={allegro_count}, competitors={competitor_count}, query_movers={query_count}.",
            ]
            if isinstance(allegro_events, list) and allegro_events:
                first = allegro_events[0]
                evidence.append(
                    f"Latest Allegro campaign mention: {first.day.isoformat()} | {first.title}."
                )
            if isinstance(competitor_events, list) and competitor_events:
                first_signal, first_competitor = competitor_events[0]
                evidence.append(
                    f"Latest competitor campaign mention: {first_signal.day.isoformat()} | {first_competitor} | {first_signal.title}."
                )
            if isinstance(query_events, list) and query_events:
                top = query_events[0]
                evidence.append(
                    f"Top campaign query mover: {top.key} ({_fmt_signed_int(top.click_delta_vs_previous)} vs prev)."
                )
            confidence = 76 if (allegro_count and competitor_count) else 69
            hypotheses.append(
                {
                    "category": "Campaign events",
                    "confidence": confidence,
                    "thesis": (
                        "Observed changes may be influenced by marketplace campaign periods on Allegro and/or competitors, "
                        "not only by technical SEO factors."
                    ),
                    "evidence": evidence,
                    "owner": "SEO + Commercial",
                }
            )

    if not hypotheses:
        hypotheses.append(
            {
                "category": "Unknown",
                "confidence": 40,
                "thesis": "No strong signals for a single dominant thesis; deeper segmentation and monitoring required.",
                "evidence": ["No dominant signals above heuristic thresholds."],
                "owner": "SEO",
            }
        )

    default_protocol = {
        "falsifier": "If next run shows opposite direction and no supporting source signals, reject this hypothesis.",
        "validation_metric": "GSC clicks/CTR/position by affected segment",
        "validation_date": (date.today() + timedelta(days=7)).isoformat(),
    }
    for row in hypotheses:
        if not isinstance(row, dict):
            continue
        category = str(row.get("category", "")).strip().lower()
        if "seasonality" in category:
            row.setdefault("falsifier", "If seasonal-cluster deltas explain <30% of total click delta for 2 runs, seasonality is not primary.")
            row.setdefault("validation_metric", "Share of total click delta explained by seasonal clusters + weather shift")
        elif "campaign" in category:
            row.setdefault("falsifier", "If campaign windows overlap <2 days with affected intents or deltas move opposite, campaign effect is unlikely.")
            row.setdefault("validation_metric", "Window-overlap days + aligned direction in trade-plan and query clusters")
        elif "algorithm" in category:
            row.setdefault("falsifier", "If impacted segments are not at least 10% worse than control segments post-update, algorithm effect is unlikely.")
            row.setdefault("validation_metric", "Post-update delta on impacted segments vs control segments")
        elif "technical" in category:
            row.setdefault("falsifier", "If technical metrics improve by >=10% with no traffic recovery, technical root-cause is weak.")
            row.setdefault("validation_metric", "CWV/device/template change vs segment traffic response")
        elif "continuity" in category:
            row.setdefault("falsifier", "If repeated themes are absent for 2 consecutive runs, continuity signal is not explanatory.")
            row.setdefault("validation_metric", "Share of repeated winner/loser terms vs prior reports")
        elif "non-brand" in category or "demand mix" in category:
            row.setdefault("falsifier", "If non-brand recovers >=5% while external demand proxies remain weak, this hypothesis is unlikely.")
            row.setdefault("validation_metric", "Non-brand clicks/CTR/position vs external demand proxies")
        else:
            row.setdefault("falsifier", default_protocol["falsifier"])
            row.setdefault("validation_metric", default_protocol["validation_metric"])
        row.setdefault("validation_date", default_protocol["validation_date"])
        evidence_count = len(row.get("evidence", [])) if isinstance(row.get("evidence"), list) else 0
        if evidence_count >= 2 and not any(token in category for token in SUPPORTING_CONTEXT_CATEGORY_TOKENS):
            row["causality_level"] = "triangulated"
        else:
            row["causality_level"] = "correlation-only"
            row["causality_guardrail_note"] = (
                "Use as contextual signal only; do not treat as standalone root cause."
            )
        if any(token in category for token in SUPPORTING_CONTEXT_CATEGORY_TOKENS):
            row["supporting_context_only"] = True
            row["confidence"] = min(63, int(row.get("confidence", 0) or 0))
            row["thesis"] = (
                str(row.get("thesis", "")).strip()
                + " (supporting context only; not a primary fact)."
            ).strip()

    return _apply_driver_priority_model(hypotheses)


def _build_integrated_reasoning(
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    ferie_context: dict[str, object],
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None,
    additional_context: dict[str, object] | None,
    senuto_summary: dict[str, float] | None,
    senuto_error: str | None,
) -> list[str]:
    precomputed = (additional_context or {}).get("precomputed_hypotheses", [])
    if isinstance(precomputed, list) and precomputed:
        hypotheses = [row for row in precomputed if isinstance(row, dict)]
    else:
        hypotheses = _build_reasoning_hypotheses(
            totals=totals,
            scope_results=scope_results,
            external_signals=external_signals,
            weather_summary=weather_summary,
            ferie_context=ferie_context,
            segment_diagnostics=segment_diagnostics,
            additional_context=additional_context,
            senuto_summary=senuto_summary,
            senuto_error=senuto_error,
        )

    lines: list[str] = []
    for row in hypotheses:
        score = int(row.get("confidence", 0))
        category = str(row.get("category", "Unknown"))
        thesis = str(row.get("thesis", ""))
        evidence = row.get("evidence")
        evidence_text = ""
        if isinstance(evidence, list):
            evidence_text = " | ".join(str(item) for item in evidence[:3])
        lines.append(f"- [{score}/100] {category}: {thesis} Evidence: {evidence_text}")
    return lines


def _build_upcoming_trends(
    run_date: date,
    external_signals: list[ExternalSignal],
    ferie_trends: list[tuple[date, str, str, str, str]] | None = None,
    additional_context: dict[str, object] | None = None,
) -> list[tuple[str, str, str, str]]:
    horizon = run_date + timedelta(days=60)
    rows: list[tuple[date, str, str, str, str]] = []

    if ferie_trends:
        rows.extend(ferie_trends)

    for signal in external_signals:
        if signal.source == "Public Holidays" and run_date < signal.day <= horizon:
            rows.append(
                (
                    signal.day,
                    signal.day.isoformat(),
                    signal.title,
                    "Holiday timing can shift purchase intent and weekly traffic distribution.",
                    "Validate campaign calendar and traffic expectations for days before/after the holiday.",
                )
            )

    product_trends = (additional_context or {}).get("product_trends", {})
    if isinstance(product_trends, dict) and product_trends.get("enabled"):
        upcoming_rows = product_trends.get("upcoming_31d", [])
        if isinstance(upcoming_rows, list):
            for row in upcoming_rows[:10]:
                if not isinstance(row, dict):
                    continue
                day_label = str(row.get("date", "")).strip()
                trend_day: date
                if day_label:
                    try:
                        trend_day = date.fromisoformat(day_label[:10])
                    except ValueError:
                        continue
                    if not (run_date < trend_day <= horizon):
                        continue
                else:
                    # next31D sheet can have no explicit date; treat as horizon-level signal
                    trend_day = run_date + timedelta(days=1)
                trend_name = str(row.get("trend", "")).strip()
                if not trend_name:
                    continue
                value = float(row.get("value", 0.0))
                rows.append(
                    (
                        trend_day,
                        trend_day.isoformat(),
                        f"Product trend window: {trend_name}",
                        (
                            f"Upcoming non-brand product trend in next {int(product_trends.get('horizon_days', 31))} days "
                            f"(score/value: {value:.2f})."
                        ),
                        "Prepare category pages/content and campaign timing before the trend window starts.",
                    )
                )

    deduped: list[tuple[date, str, str, str, str]] = []
    seen: set[tuple[str, str]] = set()
    for row in sorted(rows, key=lambda item: item[0]):
        dedupe_key = (row[1], row[2])
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        deduped.append(row)

    return [(label, title, impact, action) for _, label, title, impact, action in deduped[:10]]


def build_markdown_report(
    run_date: date,
    report_country_code: str,
    windows: dict[str, DateWindow],
    totals: dict[str, MetricSummary],
    scope_results: list[tuple[str, AnalysisResult]],
    external_signals: list[ExternalSignal],
    weather_summary: dict[str, float],
    ferie_context: dict[str, object] | None = None,
    upcoming_ferie_trends: list[tuple[date, str, str, str, str]] | None = None,
    segment_diagnostics: dict[str, list[dict[str, float | str]]] | None = None,
    additional_context: dict[str, object] | None = None,
    senuto_summary: dict[str, float] | None = None,
    senuto_error: str | None = None,
    query_filter_stats: dict[str, dict[str, int]] | None = None,
) -> str:
    if ferie_context is None:
        ferie_context = {
            "source": "OpenHolidays API (SchoolHolidays + Subdivisions)",
            "source_url": "",
            "profiles_ranked": [],
            "window_stats": {},
            "missing_years": [],
            "yoy_comparison": {"avg_daily_delta_pp": 0.0, "rows": []},
            "errors": ["Ferie context was not provided."],
        }

    current = totals["current_28d"]
    previous = totals["previous_28d"]
    yoy = totals["yoy_52w"]

    click_delta_prev = current.clicks - previous.clicks
    click_delta_prev_pct = _ratio_delta(current.clicks, previous.clicks)
    click_delta_yoy = current.clicks - yoy.clicks
    click_delta_yoy_pct = _ratio_delta(current.clicks, yoy.clicks)

    impression_delta_prev = current.impressions - previous.impressions
    impression_delta_prev_pct = _ratio_delta(current.impressions, previous.impressions)

    hypotheses = _build_reasoning_hypotheses(
        totals=totals,
        scope_results=scope_results,
        external_signals=external_signals,
        weather_summary=weather_summary,
        ferie_context=ferie_context,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
        senuto_summary=senuto_summary,
        senuto_error=senuto_error,
    )

    lines: list[str] = []
    country_label = report_country_code.strip().upper() or "PL"
    lines.append(f"# Weekly SEO Intelligence Report ({run_date.isoformat()} | {country_label})")
    lines.append("")
    lines.append("## Executive summary")
    executive_lines = _build_executive_summary_lines(
        totals=totals,
        scope_results=scope_results,
        hypotheses=hypotheses,
        external_signals=external_signals,
        weather_summary=weather_summary,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
        senuto_summary=senuto_summary,
        senuto_error=senuto_error,
    )
    lines.extend(executive_lines)

    lines.append("")
    lines.extend(
        _kpi_snapshot_table_lines(
            totals=totals,
            additional_context=additional_context,
        )
    )
    lines.append("")
    lines.extend(_driver_scoreboard_lines(hypotheses=hypotheses, limit=4))
    lines.append("")
    lines.extend(
        _meeting_ready_talking_points(
            totals=totals,
            hypotheses=hypotheses,
        )
    )
    lines.append("")
    lines.extend(
        _context_snapshot_lines(
            scope_results=scope_results,
            external_signals=external_signals,
            weather_summary=weather_summary,
            additional_context=additional_context,
        )
    )
    lines.append("")
    lines.append("## What is happening and why")
    narrative_lines = _build_what_is_happening_lines(
        totals=totals,
        windows=windows,
        scope_results=scope_results,
        hypotheses=hypotheses,
        external_signals=external_signals,
        weather_summary=weather_summary,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
    )
    lines.extend(narrative_lines)
    evidence_ledger = _build_evidence_ledger(
        windows=windows,
        external_signals=external_signals,
        additional_context=additional_context,
        weather_summary=weather_summary,
        limit=12,
    )
    if evidence_ledger:
        anchor_ids = ", ".join(f"[{row.get('id', '')}]" for row in evidence_ledger if row.get("id"))
        if anchor_ids:
            lines.append("")
            lines.append(f"- Evidence anchors used in this report: {anchor_ids}.")
            lines.append("- Key claim mapping: KPI movement [E1]; external context and timeline claims use subsequent evidence IDs from the ledger.")
    claim_evidence_rows = _claim_evidence_completeness_rows(
        executive_lines=executive_lines,
        narrative_lines=narrative_lines,
        evidence_ledger=evidence_ledger,
    )
    if claim_evidence_rows:
        mapped = sum(1 for row in claim_evidence_rows if str(row.get("status", "")).strip() == "mapped")
        total_claims = len(claim_evidence_rows)
        coverage_pct = (mapped / total_claims * 100.0) if total_claims else 0.0
        lines.append("")
        lines.append("## Evidence coverage check")
        lines.append(
            f"- Major-claim evidence coverage: {mapped}/{total_claims} ({coverage_pct:.0f}%) claims are mapped to evidence anchors."
        )
        if mapped < total_claims:
            lines.append("- Unmapped claims are treated as hypotheses (not confirmed facts) until evidence mapping is available.")
            missing = [row for row in claim_evidence_rows if str(row.get("status", "")).strip() != "mapped"]
            for row in missing[:3]:
                claim_text = _shorten(str(row.get("claim", "")).strip(), 140)
                lines.append(f"- Evidence gap: `{claim_text}`")
        lines.append("| Claim snippet | Evidence anchor | Status |")
        lines.append("|---|---|---|")
        for row in claim_evidence_rows[:6]:
            claim_text = _shorten(str(row.get("claim", "")).strip(), 110)
            anchor = str(row.get("anchor_id", "")).strip() or "-"
            status = str(row.get("status", "")).strip() or "missing"
            lines.append(f"| {claim_text} | {anchor} | {status} |")

    lines.append("")
    lines.append("## Hypothesis protocol")
    lines.append("- Protocol markers: falsifier | validation metric | validation date.")
    lines.append("| Hypothesis | Priority | Confidence | Falsifier | Validation metric | Validation date |")
    lines.append("|---|---|---|---|---|---|")
    emitted = 0
    for row in hypotheses:
        if not isinstance(row, dict):
            continue
        category = str(row.get("category", "")).strip() or "Unknown"
        thesis = str(row.get("thesis", "")).strip()
        confidence = int(row.get("confidence", 0) or 0)
        priority_score = int(row.get("driver_priority_score", 0) or 0)
        falsifier = str(row.get("falsifier", "")).strip() or "-"
        validation_metric = str(row.get("validation_metric", "")).strip() or "-"
        validation_date = str(row.get("validation_date", "")).strip() or "-"
        if not thesis:
            continue
        lines.append(
            f"| {category}: {thesis} | {priority_score}/100 | {_confidence_bucket(confidence)} ({confidence}/100) | "
            f"{falsifier} | {validation_metric} | {validation_date} |"
        )
        emitted += 1
        if emitted >= 3:
            break
    if emitted == 0:
        lines.append("| (no tracked hypotheses) | - | - | - | - | - |")

    lines.append("")
    lines.append("## Validation plan (next week)")
    validation_lines = _next_week_validation_plan_lines(hypotheses, limit=3)
    if validation_lines:
        lines.extend(validation_lines)
    else:
        lines.append("- No prioritized hypotheses available for next-week validation plan.")

    lines.append("")
    lines.append("## Counterfactual checks")
    counterfactual_lines = _counterfactual_check_lines(hypotheses, limit=3)
    if counterfactual_lines:
        lines.extend(counterfactual_lines)
    else:
        lines.append("- Counterfactual checks unavailable (no active hypotheses).")

    lines.append("")
    lines.append("## Causality guardrail")
    guardrail_summary = _causality_guardrail_summary(hypotheses)
    lines.append("- " + (guardrail_summary or "Causality guardrail status unavailable."))

    lines.append("")
    lines.append("## Escalation rule")
    contradiction_rows = _contradiction_reconciliation_lines(
        totals=totals,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
    )
    escalation_gate = _technical_seo_escalation_gate(
        totals=totals,
        hypotheses=hypotheses,
        contradiction_count=len(contradiction_rows),
    )
    lines.append(
        "- Technical SEO escalation gate: "
        f"{str(escalation_gate.get('status', '')).strip()} | "
        f"{str(escalation_gate.get('reason', '')).strip()} | "
        f"Next action: {str(escalation_gate.get('next_action', '')).strip()}"
    )

    lines.append("")
    lines.extend(
        _governance_lines(
            run_date=run_date,
            report_country_code=country_label,
            windows=windows,
            additional_context=additional_context,
        )
    )
    if evidence_ledger:
        lines.append("")
        lines.append("## Evidence ledger")
        lines.append("| ID | Source | Date | Note |")
        lines.append("|---|---|---|---|")
        for row in evidence_ledger:
            lines.append(
                f"| {row.get('id','')} | {row.get('source','')} | {row.get('date','')} | {row.get('note','')} |"
            )

    # Decision-layer report only (Executive summary + narrative).
    lines = _enforce_section_line_limits(_dedupe_report_lines(lines))
    if not INCLUDE_APPENDIX_IN_REPORT:
        return "\n".join(lines).strip() + "\n"

    lines.append("")
    lines.append("## Appendix (optional technical details)")
    lines.append("Technical diagnostics below are optional and separated from the manager decision brief.")
    lines.append("")
    lines.append("## Date windows")
    lines.append("| Window | Start | End | Days |")
    lines.append("|---|---:|---:|---:|")
    for key in [
        "current_28d",
        "previous_28d",
        "yoy_52w",
        "current_28d_context",
        "previous_28d_context",
        "yoy_28d_context_52w",
    ]:
        window = windows.get(key)
        if not isinstance(window, DateWindow):
            continue
        lines.append(
            f"| {window.name} | {window.start.isoformat()} | {window.end.isoformat()} | {window.days} |"
        )

    freshness_rows = _source_freshness_rows(additional_context=additional_context)
    if freshness_rows:
        lines.append("")
        lines.append("## Source freshness and fallback status")
        lines.append("| Source | Status | Latest data day | TTL (hours) | Cache mode | Notes |")
        lines.append("|---|---|---|---:|---|---|")
        for row in freshness_rows:
            lines.append(
                f"| {row.get('source', '')} | {str(row.get('status', '')).upper()} | "
                f"{row.get('last_day', '-') or '-'} | {float(row.get('ttl_hours', 0.0)):.1f} | "
                f"{row.get('cache_mode', '-')} | {row.get('note', '-')} |"
            )

    lines.append("")
    lines.append("## KPI summary")
    lines.append("| Metric | Current week | Previous week | YoY (52w) |")
    lines.append("|---|---:|---:|---:|")
    lines.append(f"| Clicks | {_fmt_int(current.clicks)} | {_fmt_int(previous.clicks)} | {_fmt_int(yoy.clicks)} |")
    lines.append(
        f"| Impressions | {_fmt_int(current.impressions)} | {_fmt_int(previous.impressions)} | {_fmt_int(yoy.impressions)} |"
    )
    lines.append(f"| CTR | {_pct(current.ctr)} | {_pct(previous.ctr)} | {_pct(yoy.ctr)} |")
    lines.append(
        f"| Avg position | {current.position:.2f} | {previous.position:.2f} | {yoy.position:.2f} |"
    )

    lines.append("")
    lines.append("## Main deltas")
    lines.append("| Metric | WoW | YoY (52w) |")
    lines.append("|---|---:|---:|")
    lines.append(
        f"| Clicks | {_fmt_signed_int(click_delta_prev)} ({_signed_pct(click_delta_prev_pct)}) | {_fmt_signed_int(click_delta_yoy)} ({_signed_pct(click_delta_yoy_pct)}) |"
    )
    lines.append(
        f"| Impressions | {_fmt_signed_int(impression_delta_prev)} ({_signed_pct(impression_delta_prev_pct)}) | {_fmt_signed_int(current.impressions - yoy.impressions)} ({_signed_pct(_ratio_delta(current.impressions, yoy.impressions))}) |"
    )

    lines.append("")
    lines.append("## Decision one-pager")
    lines.append("| What changed | Why it matters |")
    lines.append("|---|---|")
    for row in _decision_one_pager_rows(hypotheses, limit=6):
        lines.append(f"| {row.get('what_changed', '')} | {row.get('why_it_matters', '')} |")

    impact_rows = _impact_attribution_rows(
        totals=totals,
        scope_results=scope_results,
        segment_diagnostics=segment_diagnostics,
        limit=10,
    )
    lines.append("")
    lines.append("## Impact attribution (share of WoW click movement)")
    lines.append("| Dimension | Segment | Delta vs WoW | Share of total WoW delta |")
    lines.append("|---|---|---:|---:|")
    if impact_rows:
        for row in impact_rows:
            lines.append(
                f"| {row.get('dimension', '')} | {row.get('name', '')} | "
                f"{_fmt_signed_int(row.get('delta', 0.0))} | {float(row.get('share_of_total', 0.0)):+.1f}% |"
            )
    else:
        lines.append("| - | - | 0 | 0.0% |")

    lines.append("")
    lines.append("## Query anomaly detection (WoW)")
    lines.append("| Query | Delta vs WoW | Delta % vs WoW | Current clicks |")
    lines.append("|---|---:|---:|---:|")
    anomaly_rows = _query_anomaly_rows(_find_scope(scope_results, "query"), limit=6)
    if anomaly_rows:
        for row in anomaly_rows:
            lines.append(
                f"| `{row.get('query', '')}` | {_fmt_signed_int(row.get('delta', 0.0))} | "
                f"{float(row.get('delta_pct', 0.0)):+.2f}% | {_fmt_int(row.get('current_clicks', 0.0))} |"
            )
    else:
        lines.append("| (no outlier anomalies) | 0 | 0.00% | 0 |")

    data_quality_score, data_quality_notes = _data_quality_score(
        additional_context=additional_context,
        external_signals=external_signals,
        senuto_error=senuto_error,
    )
    lines.append("")
    lines.append("## Data quality score")
    lines.append(f"- Run data-quality score: **{data_quality_score}/100**.")
    for note in data_quality_notes:
        lines.append(f"- {note}")

    long_window = (additional_context or {}).get("long_window_context", {})
    if isinstance(long_window, dict) and long_window.get("enabled"):
        kpi = long_window.get("kpi", {})
        windows_map = long_window.get("windows", {})
        lines.append("")
        lines.append("## Long-window context (last 28 days overlay)")
        if isinstance(windows_map, dict):
            current_ctx = windows_map.get("current_28d_context", {})
            previous_ctx = windows_map.get("previous_28d_context", {})
            if isinstance(current_ctx, dict) and isinstance(previous_ctx, dict):
                lines.append(
                    "- Window overlay: "
                    f"current {current_ctx.get('start', '')} to {current_ctx.get('end', '')} "
                    f"vs previous {previous_ctx.get('start', '')} to {previous_ctx.get('end', '')}."
                )
        if isinstance(kpi, dict):
            lines.append(
                f"- Clicks 28d: {_fmt_signed_int(kpi.get('clicks_delta_vs_previous', 0.0))} "
                f"({_signed_pct(float(kpi.get('clicks_delta_pct_vs_previous', 0.0) or 0.0) / 100.0)}) vs previous 28d; "
                f"{_fmt_signed_int(kpi.get('clicks_delta_vs_yoy', 0.0))} "
                f"({_signed_pct(float(kpi.get('clicks_delta_pct_vs_yoy', 0.0) or 0.0) / 100.0)}) vs 52W."
            )
            lines.append(
                f"- Impressions 28d: {_fmt_signed_int(kpi.get('impressions_delta_vs_previous', 0.0))} "
                f"({_signed_pct(float(kpi.get('impressions_delta_pct_vs_previous', 0.0) or 0.0) / 100.0)}) vs previous 28d; "
                f"{_fmt_signed_int(kpi.get('impressions_delta_vs_yoy', 0.0))} "
                f"({_signed_pct(float(kpi.get('impressions_delta_pct_vs_yoy', 0.0) or 0.0) / 100.0)}) vs 52W."
            )
        movers = long_window.get("query_movers", {})
        if isinstance(movers, dict):
            winners = movers.get("winners", [])
            losers = movers.get("losers", [])
            if isinstance(winners, list) and winners:
                top = winners[:5]
                lines.append(
                    "- Top 28d query winners: "
                    + "; ".join(
                        f"`{row.get('key', '')}` ({_fmt_signed_int(row.get('delta_vs_previous', 0.0))} vs prev28d)"
                        for row in top
                        if isinstance(row, dict)
                    )
                    + "."
                )
            if isinstance(losers, list) and losers:
                top = losers[:5]
                lines.append(
                    "- Top 28d query losers: "
                    + "; ".join(
                        f"`{row.get('key', '')}` ({_fmt_signed_int(row.get('delta_vs_previous', 0.0))} vs prev28d)"
                        for row in top
                        if isinstance(row, dict)
                    )
                    + "."
                )

    lines.append("")
    lines.append("## YoY click contribution by query cluster")
    contribution_rows = _query_cluster_contribution_rows(
        query_scope=_find_scope(scope_results, "query"),
        total_click_delta_yoy=click_delta_yoy,
        limit=8,
    )
    lines.append(
        "| Cluster | Current clicks | Delta vs YoY | Contribution vs total YoY click delta | Movers | Example queries |"
    )
    lines.append("|---|---:|---:|---:|---:|---|")
    if contribution_rows:
        for row in contribution_rows:
            samples = row.get("samples", [])
            sample_text = (
                ", ".join(f"`{str(item).strip()}`" for item in samples[:2] if str(item).strip())
                if isinstance(samples, list)
                else ""
            )
            lines.append(
                f"| {row.get('cluster', '')} | {_fmt_int(row.get('current_clicks', 0.0))} | {_fmt_signed_int(row.get('delta_vs_yoy', 0.0))} | {float(row.get('contribution_pct', 0.0)):+.2f}% | {int(row.get('rows', 0))} | {sample_text or '-'} |"
            )
    else:
        lines.append("| (no query clusters) | 0 | +0 | +0.00% | 0 | - |")

    lines.append("")
    lines.append("## Additional source snapshots")
    pagespeed = (additional_context or {}).get("pagespeed", {})
    lines.append("### CrUX/PageSpeed (origin field data)")
    if isinstance(pagespeed, dict) and pagespeed:
        lines.append("| Strategy | LCP (ms) | INP (ms) | CLS | Category |")
        lines.append("|---|---:|---:|---:|---|")
        for strategy in ("mobile", "desktop"):
            row = pagespeed.get(strategy, {})
            if not isinstance(row, dict):
                continue
            lines.append(
                f"| {strategy} | {float(row.get('lcp_ms', 0.0)):.0f} | {float(row.get('inp_ms', 0.0)):.0f} | {float(row.get('cls', 0.0)):.2f} | {row.get('overall_category', '')} |"
            )
    else:
        lines.append("- No PageSpeed/CrUX data in this run.")

    ga4 = (additional_context or {}).get("ga4", {})
    lines.append("")
    lines.append("### Google Analytics 4 (country slice)")
    if isinstance(ga4, dict) and ga4.get("enabled"):
        lines.append(
            f"- Property: {str(ga4.get('property_id', '')).strip() or '(configured)'} | Country: {str(ga4.get('country_code', '')).strip() or '-'}"
        )
        summary = ga4.get("summary", {})
        if isinstance(summary, dict):
            current_ga4 = summary.get("current", {})
            previous_ga4 = summary.get("previous", {})
            yoy_ga4 = summary.get("yoy", {})
            if (
                isinstance(current_ga4, dict)
                and isinstance(previous_ga4, dict)
                and isinstance(yoy_ga4, dict)
            ):
                metric_rows = [
                    ("sessions", "Sessions"),
                    ("users", "Users"),
                    ("engaged_sessions", "Engaged sessions"),
                    ("transactions", "Transactions"),
                    ("revenue", "Purchase revenue"),
                ]
                lines.append("| Metric | Current week | Previous week | YoY (52w) |")
                lines.append("|---|---:|---:|---:|")
                added_rows = 0
                for metric, label in metric_rows:
                    if not _ga4_metric_available(ga4, metric):
                        continue
                    current_value = _ga4_num(current_ga4, metric)
                    previous_value = _ga4_num(previous_ga4, metric)
                    yoy_value = _ga4_num(yoy_ga4, metric)
                    if current_value is None or previous_value is None or yoy_value is None:
                        continue
                    lines.append(
                        f"| {label} | {_fmt_int(current_value)} | {_fmt_int(previous_value)} | {_fmt_int(yoy_value)} |"
                    )
                    added_rows += 1
                if added_rows == 0:
                    lines.append("| (no complete GA4 metrics in this run) | - | - | - |")

        channels = ga4.get("channels", {})
        yoy_deltas = channels.get("yoy_deltas", []) if isinstance(channels, dict) else []
        lines.append("")
        lines.append("| Channel | Sessions current | Sessions YoY | Delta vs YoY | Delta % vs YoY |")
        lines.append("|---|---:|---:|---:|---:|")
        if isinstance(yoy_deltas, list) and yoy_deltas:
            for row in yoy_deltas[: int(ga4.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                delta_pct = row.get("delta_vs_yoy_pct")
                delta_pct_text = (
                    f"{float(delta_pct):+.2f}%"
                    if isinstance(delta_pct, (float, int))
                    else "-"
                )
                lines.append(
                    f"| {row.get('channel', '')} | {_fmt_int(row.get('sessions_current', 0.0))} | {_fmt_int(row.get('sessions_yoy', 0.0))} | {_fmt_signed_int(row.get('delta_vs_yoy', 0.0))} | {delta_pct_text} |"
                )
        else:
            lines.append("| (no channel rows) | 0 | 0 | +0 | - |")

        if isinstance(yoy_deltas, list) and yoy_deltas:
            rising = sorted(
                [row for row in yoy_deltas if isinstance(row, dict)],
                key=lambda row: float(row.get("delta_vs_yoy", 0.0) or 0.0),
                reverse=True,
            )[:5]
            falling = sorted(
                [row for row in yoy_deltas if isinstance(row, dict)],
                key=lambda row: float(row.get("delta_vs_yoy", 0.0) or 0.0),
            )[:5]
            if rising:
                lines.append(
                    "- Top channel growth YoY: "
                    + "; ".join(
                        f"{row.get('channel', '')} ({_fmt_signed_int(row.get('delta_vs_yoy', 0.0))})"
                        for row in rising
                    )
                    + "."
                )
            if falling:
                lines.append(
                    "- Top channel decline YoY: "
                    + "; ".join(
                        f"{row.get('channel', '')} ({_fmt_signed_int(row.get('delta_vs_yoy', 0.0))})"
                        for row in falling
                    )
                    + "."
                )

        cannibalization = ga4.get("cannibalization", {})
        if isinstance(cannibalization, dict):
            lines.append(
                "- Potential channel cannibalization (GA4): "
                f"{str(cannibalization.get('note', '')).strip()}"
            )

        top_pages = ga4.get("top_landing_pages", [])
        lines.append("")
        show_transactions = _ga4_metric_available(ga4, "transactions")
        show_revenue = _ga4_metric_available(ga4, "revenue")
        page_header = "| Landing page | Sessions |"
        page_divider = "|---|---:|"
        if show_transactions:
            page_header += " Transactions |"
            page_divider += "---:|"
        if show_revenue:
            page_header += " Revenue |"
            page_divider += "---:|"
        lines.append(page_header)
        lines.append(page_divider)
        if isinstance(top_pages, list) and top_pages:
            for row in top_pages[: int(ga4.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                line = (
                    f"| {row.get('landing_page', '')} | {_fmt_int(row.get('sessions', 0.0))} |"
                )
                if show_transactions:
                    line += f" {_fmt_int(row.get('transactions', 0.0))} |"
                if show_revenue:
                    line += f" {_fmt_int(row.get('revenue', 0.0))} |"
                lines.append(line)
        else:
            empty = "| (no rows) | 0 |"
            if show_transactions:
                empty += " 0 |"
            if show_revenue:
                empty += " 0 |"
            lines.append(empty)
    else:
        lines.append("- GA4 not configured or unavailable in this run.")
        if isinstance(ga4, dict):
            errors = ga4.get("errors", [])
            if isinstance(errors, list):
                for err in errors[:3]:
                    lines.append(f"- GA4 warning: {err}")

    allegro_trends = (additional_context or {}).get("allegro_trends", {})
    lines.append("")
    lines.append("### Allegro Trends API (marketplace demand)")
    if isinstance(allegro_trends, dict) and allegro_trends.get("enabled"):
        source_label = str(allegro_trends.get("source", "Allegro Trends API")).strip()
        window_start = str(allegro_trends.get("from", "")).strip()
        window_end = str(allegro_trends.get("till", "")).strip()
        country_trends = str(allegro_trends.get("country_code", country_label)).strip() or country_label
        interval_label = str(allegro_trends.get("interval", "day")).strip() or "day"
        lines.append(
            f"- Source: {source_label} | Country: {country_trends} | Window: {window_start} to {window_end} | Interval: {interval_label}"
        )
        rows = allegro_trends.get("rows", [])
        top_rows = int(allegro_trends.get("top_rows", 10) or 10)
        lines.append("| Query | Visits | PV | Offers | Deals | GMV | Data points | HTTP |")
        lines.append("|---|---:|---:|---:|---:|---:|---:|---:|")
        if isinstance(rows, list) and rows:
            total_visits = 0.0
            total_pv = 0.0
            total_gmv = 0.0
            for row in rows[:max(1, top_rows)]:
                if not isinstance(row, dict):
                    continue
                visits = float(row.get("visit", 0.0) or 0.0)
                pv = float(row.get("pv", 0.0) or 0.0)
                offers = float(row.get("offers", 0.0) or 0.0)
                deals = float(row.get("deals", 0.0) or 0.0)
                gmv = float(row.get("gmv", 0.0) or 0.0)
                points = int(float(row.get("points", 0.0) or 0.0))
                http_code = int(float(row.get("http_code", 0.0) or 0.0))
                total_visits += visits
                total_pv += pv
                total_gmv += gmv
                lines.append(
                    f"| {row.get('query', '')} | {_fmt_int(visits)} | {_fmt_int(pv)} | {_fmt_int(offers)} | {_fmt_int(deals)} | {_fmt_int(gmv)} | {_fmt_int(points)} | {_fmt_int(http_code)} |"
                )
            lines.append(
                f"- Aggregated tracked-query totals: visits {_fmt_int(total_visits)}, PV {_fmt_int(total_pv)}, GMV {_fmt_int(total_gmv)}."
            )
        else:
            lines.append("| (no rows) | 0 | 0 | 0 | 0 | 0 | 0 | 0 |")
        errors = allegro_trends.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:3]:
                lines.append(f"- Allegro Trends warning: {err}")
    else:
        lines.append("- Allegro Trends API not configured or unavailable in this run.")
        if isinstance(allegro_trends, dict):
            errors = allegro_trends.get("errors", [])
            if isinstance(errors, list):
                for err in errors[:3]:
                    lines.append(f"- Allegro Trends warning: {err}")

    trends = (additional_context or {}).get("google_trends", [])
    lines.append("")
    lines.append(f"### Google Trends ({country_label})")
    if isinstance(trends, list) and trends:
        lines.append("| Date | Topic | Approx traffic |")
        lines.append("|---|---|---:|")
        for row in trends[:12]:
            if not isinstance(row, dict):
                continue
            day = row.get("day")
            day_label = day.isoformat() if isinstance(day, date) else str(day or "")
            lines.append(
                f"| {day_label} | {row.get('topic', '')} | {int(row.get('approx_traffic', 0))} |"
            )
    else:
        lines.append("- No Google Trends items in this run.")
    brand_trends = (additional_context or {}).get("google_trends_brand", {})
    if isinstance(brand_trends, dict) and brand_trends.get("enabled"):
        rows = brand_trends.get("rows", [])
        summary = brand_trends.get("summary", {})
        lines.append("")
        lines.append("#### Brand Search Interest (Google Trends API)")
        lines.append("| Keyword | Avg interest (current) | Avg interest (previous) | Avg interest (YoY) | WoW delta % | YoY delta % |")
        lines.append("|---|---:|---:|---:|---:|---:|")
        if isinstance(rows, list) and rows:
            for row in rows[:6]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('keyword', '')} | {float(row.get('current_avg', 0.0)):.1f} | {float(row.get('previous_avg', 0.0)):.1f} | {float(row.get('yoy_avg', 0.0)):.1f} | {float(row.get('delta_pct_vs_previous', 0.0)):+.2f}% | {float(row.get('delta_pct_vs_yoy', 0.0)):+.2f}% |"
                )
        else:
            lines.append("| (no rows) | 0.0 | 0.0 | 0.0 | +0.00% | +0.00% |")
        if isinstance(summary, dict):
            lines.append(
                "- Blended brand-interest trend: "
                f"WoW {float(summary.get('delta_pct_vs_previous', 0.0)):+.2f}%, "
                f"YoY {float(summary.get('delta_pct_vs_yoy', 0.0)):+.2f}%."
            )
        errors = brand_trends.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:3]:
                lines.append(f"- Google Trends brand warning: {err}")

    lines.append("")
    lines.append("### DuckDuckGo context scan")
    ddg_context = (additional_context or {}).get("duckduckgo_context", {})
    if isinstance(ddg_context, dict) and ddg_context.get("enabled"):
        ddg_rows = ddg_context.get("rows", [])
        lines.append("| Query | Heading | Abstract/Hint |")
        lines.append("|---|---|---|")
        if isinstance(ddg_rows, list) and ddg_rows:
            for row in ddg_rows[:8]:
                if not isinstance(row, dict):
                    continue
                query = str(row.get("query", "")).strip()
                heading = str(row.get("heading", "")).strip()
                abstract = str(row.get("abstract", "")).strip()
                if len(abstract) > 220:
                    abstract = abstract[:217] + "..."
                related = row.get("related_topics", [])
                if not abstract and isinstance(related, list) and related:
                    abstract = "; ".join(str(item) for item in related[:2])
                lines.append(f"| {query} | {heading or '-'} | {abstract or '-'} |")
        else:
            lines.append("| (no rows) | - | - |")
        errors = ddg_context.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:3]:
                lines.append(f"- DuckDuckGo warning: {err}")
    else:
        lines.append("- DuckDuckGo context not available in this run.")

    lines.append("")
    lines.append("### Market event calendar (API)")
    market_events = (additional_context or {}).get("market_event_calendar", {})
    if isinstance(market_events, dict) and market_events.get("enabled"):
        lines.append(
            f"- Source: {market_events.get('source', 'Market Events API')} | Country: {market_events.get('country_code', country_label)}"
        )
        rows = market_events.get("events", [])
        lines.append("| Date | Event type | GMV impact | Source | Event | Why it may affect GMV |")
        lines.append("|---|---|---|---|---|---|")
        if isinstance(rows, list) and rows:
            max_rows = int(market_events.get("top_rows", 12) or 12)
            for row in rows[:max_rows]:
                if not isinstance(row, dict):
                    continue
                impact = (
                    f"{row.get('impact_level', 'LOW')} / {row.get('impact_direction', 'Mixed')} "
                    f"(conf {int(row.get('confidence', 0) or 0)}/100)"
                )
                lines.append(
                    f"| {row.get('date', '')} | {row.get('event_type', '')} | {impact} | {row.get('source', '')} | {row.get('title', '')} | {row.get('gmv_reason', '')} |"
                )
        else:
            lines.append("| (no API events) | - | - | - | - | - |")
        errors = market_events.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:3]:
                lines.append(f"- Market event calendar warning: {err}")
    else:
        lines.append("- Market event calendar not configured or unavailable in this run.")

    free_public_hub = (additional_context or {}).get("free_public_source_hub", {})
    lines.append("")
    lines.append("### Free public source hub (20-source inventory)")
    if isinstance(free_public_hub, dict) and free_public_hub.get("enabled"):
        hub_rows = free_public_hub.get("rows", [])
        lines.append("| Source | Type | Status | Details | URL |")
        lines.append("|---|---|---|---|---|")
        if isinstance(hub_rows, list) and hub_rows:
            for row in hub_rows[:30]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('source', '')} | {row.get('type', '')} | {row.get('status', '')} | "
                    f"{row.get('details', '')} | {row.get('url', '')} |"
                )
            if len(hub_rows) > 30:
                lines.append(f"- +{_fmt_int(len(hub_rows) - 30)} additional source-status rows.")
        else:
            lines.append("| (no source rows) | - | - | - | - |")
        errors = free_public_hub.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:5]:
                lines.append(f"- Free-source hub warning: {err}")
    else:
        lines.append("- Free public source hub disabled or unavailable in this run.")

    lines.append("")
    lines.append("### SEO update early analyses (external SEO media)")
    seo_update_rows = [
        signal
        for signal in external_signals
        if "seo update analysis" in _normalize_text(signal.source)
    ]
    if seo_update_rows:
        lines.append("| Date | Source | Update article | Early takeaways |")
        lines.append("|---|---|---|---|")
        for signal in seo_update_rows[:8]:
            details = str(signal.details or "").strip()
            if len(details) > 220:
                details = details[:217] + "..."
            lines.append(
                f"| {signal.day.isoformat()} | {signal.source} | {signal.title} | {details} |"
            )
    else:
        lines.append("- No SEO-media early analyses detected in this run window.")

    product_trends = (additional_context or {}).get("product_trends", {})
    lines.append("")
    lines.append("### Product trends (non-brand, Sheets)")
    if isinstance(product_trends, dict) and product_trends.get("enabled"):
        lines.append(f"- Source: {product_trends.get('source', 'Google Sheets product trend trackers')}")
        lines.append(f"- Top rows used: {int(product_trends.get('top_rows', 0))}")
        lines.append(f"- Upcoming horizon: {int(product_trends.get('horizon_days', 0))} days")

        yoy_rows = product_trends.get("top_yoy_non_brand", [])
        if not (isinstance(yoy_rows, list) and yoy_rows):
            yoy_rows = _fallback_non_brand_yoy_from_gsc(
                query_scope=_find_scope(scope_results, "query"),
                top_rows=int(product_trends.get("top_rows", 12)),
            )
            if yoy_rows:
                lines.append("- YoY trend source fallback: GSC non-brand queries (sheet YoY snapshot unavailable).")
        lines.append("")
        if isinstance(yoy_rows, list) and yoy_rows:
            first_yoy_row = next((row for row in yoy_rows if isinstance(row, dict)), None)
            if isinstance(first_yoy_row, dict):
                current_snapshot = str(first_yoy_row.get("current_snapshot_date", "")).strip()
                previous_snapshot = str(first_yoy_row.get("previous_snapshot_date", "")).strip()
                if current_snapshot or previous_snapshot:
                    lines.append(
                        f"- YoY snapshot dates: current={current_snapshot or '-'} | previous={previous_snapshot or '-'}"
                    )
            lines.append("| Trend | Current year | Previous year | Delta | Delta % | Sheet |")
            lines.append("|---|---:|---:|---:|---:|---|")
            for row in yoy_rows[: int(product_trends.get("top_rows", 12))]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('trend', '')} | {_fmt_int(row.get('current_value', 0.0))} | {_fmt_int(row.get('previous_value', 0.0))} | {_fmt_signed_int(row.get('delta_value', 0.0))} | {float(row.get('delta_pct', 0.0)):+.2f}% | {row.get('sheet', '')} |"
                )
        else:
            lines.append("| Trend | Current year | Previous year | Delta | Delta % | Sheet |")
            lines.append("|---|---:|---:|---:|---:|---|")
            lines.append("| (no data) | 0 | 0 | +0 | +0.00% | - |")

        upcoming_rows = product_trends.get("upcoming_31d", [])
        lines.append("")
        lines.append("| Upcoming date | Trend | Value | Sheet |")
        lines.append("|---|---|---:|---|")
        if isinstance(upcoming_rows, list) and upcoming_rows:
            for row in upcoming_rows[: int(product_trends.get("top_rows", 12))]:
                if not isinstance(row, dict):
                    continue
                day_label = str(row.get("date", "")).strip() or "(next31d)"
                lines.append(
                    f"| {day_label} | {row.get('trend', '')} | {float(row.get('value', 0.0)):.2f} | {row.get('sheet', '')} |"
                )
        else:
            lines.append("| (no data) | - | 0.00 | - |")

        current_rows = product_trends.get("current_non_brand", [])
        lines.append("")
        lines.append("| Trend date | Current trend | Value | Sheet |")
        lines.append("|---|---|---:|---|")
        if isinstance(current_rows, list) and current_rows:
            for row in current_rows[: int(product_trends.get("top_rows", 12))]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('date', '')} | {row.get('trend', '')} | {float(row.get('value', 0.0)):.2f} | {row.get('sheet', '')} |"
                )
        else:
            lines.append("| (no data) | - | 0.00 | - |")

        trend_errors = product_trends.get("errors", [])
        if isinstance(trend_errors, list):
            for err in trend_errors[:5]:
                lines.append(f"- Product trend warning: {err}")
    else:
        lines.append("- No product trend sheet data in this run.")

    senuto_intelligence = (additional_context or {}).get("senuto_intelligence", {})
    lines.append("")
    lines.append("### Senuto competitor and SERP intelligence")
    if isinstance(senuto_intelligence, dict) and senuto_intelligence.get("enabled"):
        lines.append(
            f"- Country id: {int(senuto_intelligence.get('country_id', 0) or 0)} | Top rows: {int(senuto_intelligence.get('top_rows', 10) or 10)}"
        )

        competitors_overview = senuto_intelligence.get("competitors_overview", [])
        if isinstance(competitors_overview, list) and competitors_overview:
            lines.append("| Competitor domain | Common keywords | Visibility (current) | Visibility delta % | Top10 keywords | Domain rank |")
            lines.append("|---|---:|---:|---:|---:|---:|")
            for row in competitors_overview[: int(senuto_intelligence.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('domain', '')} | {_fmt_int(row.get('common_keywords', 0.0))} | {_fmt_int(row.get('visibility_current', 0.0))} | {float(row.get('visibility_diff_pct', 0.0)):+.2f}% | {_fmt_int(row.get('top10_current', 0.0))} | {_fmt_int(row.get('domain_rank_current', 0.0))} |"
                )

        wins_losses = senuto_intelligence.get("wins_losses", {})
        wins_rows = wins_losses.get("wins", []) if isinstance(wins_losses, dict) else []
        losses_rows = wins_losses.get("losses", []) if isinstance(wins_losses, dict) else []
        if isinstance(wins_rows, list) and wins_rows:
            lines.append("")
            lines.append("| Keyword movers | Visibility current | Visibility previous | Visibility diff | Visibility diff % |")
            lines.append("|---|---:|---:|---:|---:|")
            for row in wins_rows[:5]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| WIN: {row.get('keyword', '')} | {_fmt_int(row.get('visibility_current', 0.0))} | {_fmt_int(row.get('visibility_previous', 0.0))} | {_fmt_signed_int(row.get('visibility_diff', 0.0))} | {float(row.get('visibility_diff_pct', 0.0)):+.2f}% |"
                )
        if isinstance(losses_rows, list) and losses_rows:
            if not (isinstance(wins_rows, list) and wins_rows):
                lines.append("")
                lines.append("| Keyword movers | Visibility current | Visibility previous | Visibility diff | Visibility diff % |")
                lines.append("|---|---:|---:|---:|---:|")
            for row in losses_rows[:5]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| LOSS: {row.get('keyword', '')} | {_fmt_int(row.get('visibility_current', 0.0))} | {_fmt_int(row.get('visibility_previous', 0.0))} | {_fmt_signed_int(row.get('visibility_diff', 0.0))} | {float(row.get('visibility_diff_pct', 0.0)):+.2f}% |"
                )

        acquired_lost = senuto_intelligence.get("acquired_lost", {})
        kw_acquired = acquired_lost.get("keywords_acquired", []) if isinstance(acquired_lost, dict) else []
        kw_lost = acquired_lost.get("keywords_lost", []) if isinstance(acquired_lost, dict) else []
        url_acquired = acquired_lost.get("urls_acquired", []) if isinstance(acquired_lost, dict) else []
        url_lost = acquired_lost.get("urls_lost", []) if isinstance(acquired_lost, dict) else []
        lines.append("")
        lines.append(
            f"- Acquired/Lost snapshot: keywords acquired={_fmt_int(len(kw_acquired) if isinstance(kw_acquired, list) else 0)}, "
            f"keywords lost={_fmt_int(len(kw_lost) if isinstance(kw_lost, list) else 0)}, "
            f"urls acquired={_fmt_int(len(url_acquired) if isinstance(url_acquired, list) else 0)}, "
            f"urls lost={_fmt_int(len(url_lost) if isinstance(url_lost, list) else 0)}."
        )

        direct_answers = senuto_intelligence.get("direct_answers", [])
        if isinstance(direct_answers, list) and direct_answers:
            lines.append("")
            lines.append("| Direct-answer keyword | Feature | Searches | Position |")
            lines.append("|---|---|---:|---:|")
            for row in direct_answers[: int(senuto_intelligence.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('keyword', '')} | {row.get('feature', '') or '-'} | {_fmt_int(row.get('searches', 0.0))} | {float(row.get('position', 0.0)):.1f} |"
                )

        seasonality = senuto_intelligence.get("seasonality", {})
        if isinstance(seasonality, dict):
            trend_values = seasonality.get("trend_values", [])
            if isinstance(trend_values, list) and trend_values:
                lines.append("")
                lines.append("| Seasonality (month) | Index |")
                lines.append("|---|---:|")
                for idx, value in enumerate(trend_values, start=1):
                    lines.append(f"| M{idx:02d} | {float(value):.2f} |")
                lines.append(
                    f"- Peak month M{int(seasonality.get('peak_month', 0) or 0)} ({float(seasonality.get('peak_value', 0.0) or 0.0):.2f}), "
                    f"low month M{int(seasonality.get('low_month', 0) or 0)} ({float(seasonality.get('low_value', 0.0) or 0.0):.2f})."
                )

        market_ranking = senuto_intelligence.get("market_ranking", [])
        if isinstance(market_ranking, list) and market_ranking:
            lines.append("")
            lines.append("| Market ranking domain | Category | Rank current | Rank previous | Visibility current | Visibility previous | Share |")
            lines.append("|---|---|---:|---:|---:|---:|---:|")
            for row in market_ranking[: int(senuto_intelligence.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('domain', '')} | {row.get('category', '')} | {_fmt_int(row.get('rank_current', 0.0))} | {_fmt_int(row.get('rank_previous', 0.0))} | {_fmt_int(row.get('visibility_current', 0.0))} | {_fmt_int(row.get('visibility_previous', 0.0))} | {float(row.get('share', 0.0)):.2f} |"
                )

        keyword_trending = senuto_intelligence.get("keyword_trending", [])
        if isinstance(keyword_trending, list) and keyword_trending:
            lines.append("")
            lines.append("| Senuto trending keyword | Searches | Growth metric |")
            lines.append("|---|---:|---:|")
            for row in keyword_trending[: int(senuto_intelligence.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('keyword', '')} | {_fmt_int(row.get('searches', 0.0))} | {float(row.get('growth', 0.0)):.2f} |"
                )

        serp_volatility = senuto_intelligence.get("serp_volatility", [])
        if isinstance(serp_volatility, list) and serp_volatility:
            lines.append("")
            lines.append("| SERP volatility keyword | Domain | Volatility | Position points |")
            lines.append("|---|---|---:|---:|")
            for row in serp_volatility[: int(senuto_intelligence.get("top_rows", 10) or 10)]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('keyword', '')} | {row.get('domain', '')} | {float(row.get('volatility', 0.0)):.2f} | {_fmt_int(row.get('positions_points', 0.0))} |"
                )

        errors = senuto_intelligence.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:8]:
                lines.append(f"- Senuto intelligence warning: {err}")
    else:
        lines.append("- Senuto intelligence not configured in this run.")

    lines.append("")
    lines.append("### Trend-to-page coverage check")
    coverage_rows = _trend_page_coverage_rows(
        scope_results=scope_results,
        additional_context=additional_context,
        limit=8,
    )
    lines.append(
        "| Trend | Landing page | Trend score/value | Median position (sheet) | Page delta WoW | Page delta YoY | Status | Gap interpretation |"
    )
    lines.append("|---|---|---:|---:|---:|---:|---|---|")
    if coverage_rows:
        for row in coverage_rows:
            lines.append(
                f"| {row.get('trend', '')} | {row.get('page', '-')} | {float(row.get('sheet_value', 0.0)):.2f} | {float(row.get('median_position', 0.0)):.2f} | {_fmt_signed_int(row.get('page_delta_vs_prev', 0.0))} | {_fmt_signed_int(row.get('page_delta_vs_yoy', 0.0))} | {row.get('status', '')} | {row.get('gap_reason', '')} |"
            )
    else:
        lines.append("| (no trend coverage rows) | - | 0.00 | 0.00 | +0 | +0 | - | No trend rows with landing page metadata in this run. |")

    lines.append("")
    lines.append("### Query clusters (top movers grouped)")
    cluster_rows = _query_cluster_rows(_find_scope(scope_results, "query"))
    lines.append("| Cluster | Movers | Current clicks | Delta vs prev | Delta vs YoY | Example queries |")
    lines.append("|---|---:|---:|---:|---:|---|")
    if cluster_rows:
        for row in cluster_rows[:8]:
            samples = row.get("samples", [])
            sample_text = ""
            if isinstance(samples, list):
                sample_text = ", ".join(
                    f"`{str(item).strip()}`" for item in samples[:3] if str(item).strip()
                )
            lines.append(
                f"| {row.get('cluster', '')} | {int(row.get('rows', 0))} | {_fmt_int(row.get('current_clicks', 0.0))} | {_fmt_signed_int(row.get('delta_vs_previous', 0.0))} | {_fmt_signed_int(row.get('delta_vs_yoy', 0.0))} | {sample_text or '-'} |"
            )
    else:
        lines.append("| (no query clusters) | 0 | 0 | +0 | +0 | - |")

    campaign_context = _campaign_event_context(
        external_signals=external_signals,
        query_scope=_find_scope(scope_results, "query"),
    )
    allegro_campaign_events = campaign_context.get("allegro_events", [])
    competitor_campaign_events = campaign_context.get("competitor_events", [])
    campaign_query_events = campaign_context.get("query_events", [])
    lines.append("")
    lines.append("### Campaign events (Allegro vs competitors)")
    has_rows = False
    if isinstance(allegro_campaign_events, list) and allegro_campaign_events:
        lines.append("| Side | Date | Source | Event |")
        lines.append("|---|---|---|---|")
    if isinstance(allegro_campaign_events, list):
        for row in allegro_campaign_events[:8]:
            if not isinstance(row, ExternalSignal):
                continue
            has_rows = True
            lines.append(
                f"| Allegro | {row.day.isoformat()} | {row.source} | {row.title} |"
            )
    if isinstance(competitor_campaign_events, list):
        for row in competitor_campaign_events[:8]:
            if (
                not isinstance(row, tuple)
                or len(row) != 2
                or not isinstance(row[0], ExternalSignal)
            ):
                continue
            if not has_rows:
                lines.append("| Side | Date | Source | Event |")
                lines.append("|---|---|---|---|")
            signal, competitor = row
            has_rows = True
            lines.append(
                f"| Competitor ({competitor}) | {signal.day.isoformat()} | {signal.source} | {signal.title} |"
            )
    if not has_rows:
        lines.append("- No detected campaign-event mentions in current source set.")

    lines.append("")
    if isinstance(campaign_query_events, list) and campaign_query_events:
        lines.append("| Campaign query | Delta vs prev | Delta vs YoY |")
        lines.append("|---|---:|---:|")
        for row in campaign_query_events[:8]:
            lines.append(
                f"| {row.key} | {_fmt_signed_int(row.click_delta_vs_previous)} ({_signed_pct(row.click_delta_pct_vs_previous)}) | {_fmt_signed_int(row.click_delta_vs_yoy)} ({_signed_pct(row.click_delta_pct_vs_yoy)}) |"
            )
    else:
        lines.append("- No campaign queries among top movers in this run.")

    macro = (additional_context or {}).get("macro", {})
    lines.append("")
    lines.append("### Macro context (NBP / IMGW)")
    if isinstance(macro, dict) and macro:
        note = str(macro.get("note", "")).strip()
        if note:
            lines.append(f"- {note}")
        nbp_fx = macro.get("nbp_fx", {})
        if isinstance(nbp_fx, dict):
            lines.append("| Pair | Avg current | Avg previous | Delta vs previous | Latest |")
            lines.append("|---|---:|---:|---:|---:|")
            for key, label in (("eur_pln", "EUR/PLN"), ("usd_pln", "USD/PLN")):
                row = nbp_fx.get(key, {})
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {label} | {float(row.get('avg_current', 0.0)):.4f} | {float(row.get('avg_previous', 0.0)):.4f} | {float(row.get('delta_pct_vs_previous', 0.0)):+.2f}% | {float(row.get('latest', 0.0)):.4f} |"
                )
        warnings = macro.get("imgw_warnings", [])
        if isinstance(warnings, list):
            total_warnings = int(macro.get("imgw_warnings_total", len(warnings)))
            high_warnings = int(
                macro.get(
                    "imgw_high_severity_count",
                    sum(
                        1
                        for row in warnings
                        if isinstance(row, dict)
                        and str(row.get("severity", "")).strip() in {"2", "3"}
                    ),
                )
            )
            lines.append(
                f"- Active IMGW warnings: {total_warnings} (showing {len(warnings)}), level 2-3: {high_warnings}."
            )
            if warnings:
                lines.append("| Severity | Event | Areas | Valid to |")
                lines.append("|---:|---|---|---|")
                for row in warnings[:8]:
                    if not isinstance(row, dict):
                        continue
                    lines.append(
                        f"| {row.get('severity', '') or '-'} | {row.get('event', '') or '-'} | {row.get('areas', '') or '-'} | {row.get('to', '') or '-'} |"
                    )
    else:
        lines.append("- No macro context data in this run.")

    macro_backdrop = (additional_context or {}).get("macro_backdrop", {})
    lines.append("")
    lines.append("### Macro backdrop (country-level, World Bank)")
    if isinstance(macro_backdrop, dict) and macro_backdrop.get("rows"):
        lines.append(
            f"- Source: {macro_backdrop.get('source', 'World Bank API')} | Country: {macro_backdrop.get('country_code', country_label)}"
        )
        rows = macro_backdrop.get("rows", {})
        if isinstance(rows, dict):
            lines.append("| Indicator | Latest year | Latest value | Previous year | Previous value |")
            lines.append("|---|---:|---:|---:|---:|")
            for key, label in (
                ("inflation_cpi_pct", "Inflation CPI %"),
                ("unemployment_pct", "Unemployment %"),
                ("gdp_growth_pct", "GDP growth %"),
            ):
                row = rows.get(key, {})
                if not isinstance(row, dict):
                    continue
                latest_val = row.get("latest_value")
                prev_val = row.get("previous_value")
                latest_text = f"{float(latest_val):.2f}" if isinstance(latest_val, (int, float)) else "-"
                prev_text = f"{float(prev_val):.2f}" if isinstance(prev_val, (int, float)) else "-"
                lines.append(
                    f"| {label} | {row.get('latest_year', '') or '-'} | {latest_text} | {row.get('previous_year', '') or '-'} | {prev_text} |"
                )
        errors = macro_backdrop.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:3]:
                lines.append(f"- Macro backdrop warning: {err}")
    else:
        lines.append("- No country-level macro backdrop rows in this run.")

    operational = (additional_context or {}).get("operational_risks", {})
    lines.append("")
    lines.append("### Operational risks (logistics / payments)")
    if isinstance(operational, dict) and operational.get("enabled"):
        logistics = operational.get("logistics", [])
        payments = operational.get("payments", [])
        lines.append(
            f"- Signals captured: logistics={_fmt_int(len(logistics) if isinstance(logistics, list) else 0)}, "
            f"payments={_fmt_int(len(payments) if isinstance(payments, list) else 0)}."
        )
        lines.append("| Type | Date | Headline |")
        lines.append("|---|---|---|")
        if isinstance(logistics, list):
            for row in logistics[:5]:
                if not isinstance(row, dict):
                    continue
                lines.append(f"| Logistics | {row.get('date', '')} | {row.get('title', '')} |")
        if isinstance(payments, list):
            for row in payments[:5]:
                if not isinstance(row, dict):
                    continue
                lines.append(f"| Payments | {row.get('date', '')} | {row.get('title', '')} |")
    else:
        lines.append("- No operational-risk signals in this run.")

    promo_radar = (additional_context or {}).get("competitor_promo_radar", {})
    lines.append("")
    lines.append("### Competitor promo radar")
    if isinstance(promo_radar, dict) and promo_radar.get("enabled"):
        rows = promo_radar.get("rows", [])
        lines.append(
            f"- Source: {promo_radar.get('source', 'Google News RSS query')} | Rows: {_fmt_int(len(rows) if isinstance(rows, list) else 0)}"
        )
        lines.append("| Date | Signal |")
        lines.append("|---|---|")
        if isinstance(rows, list):
            for row in rows[:10]:
                if not isinstance(row, dict):
                    continue
                lines.append(f"| {row.get('date', '')} | {row.get('title', '')} |")
    else:
        lines.append("- No competitor promo-radar mentions in this run.")

    unified_timeline_rows = _marketplace_timeline_rows(
        additional_context=additional_context,
        campaign_context=campaign_context,
        report_country_code=country_label,
        max_rows=50,
    )
    lines.append("")
    lines.append("### Unified marketplace timeline (market events + promo calendar)")
    lines.append("| Date | Country | Track | Event | Impact/confidence | Source |")
    lines.append("|---|---|---|---|---|---|")
    if unified_timeline_rows:
        for row in unified_timeline_rows[:30]:
            day = row.get("day")
            day_label = day.isoformat() if isinstance(day, date) else str(day or "")
            lines.append(
                f"| {day_label} | {row.get('country', '')} | {row.get('track', '')} | "
                f"{row.get('event', '')} | {row.get('impact', '')} | {row.get('source', '')} |"
            )
        if len(unified_timeline_rows) > 30:
            lines.append(f"- +{_fmt_int(len(unified_timeline_rows) - 30)} additional timeline events in raw context.")
    else:
        lines.append("| (no timeline rows) | - | - | - | - | - |")

    source_errors = (additional_context or {}).get("errors", [])
    if isinstance(source_errors, list):
        for err in source_errors[:5]:
            lines.append(f"- Source warning: {err}")

    seo_presentations = (additional_context or {}).get("seo_presentations", {})
    lines.append("")
    lines.append("### SEO team presentations (Drive)")
    if isinstance(seo_presentations, dict) and seo_presentations.get("enabled"):
        lines.append(f"- Source: {seo_presentations.get('source', 'Google Drive')}")
        year_rows = seo_presentations.get("years", [])
        if isinstance(year_rows, list) and year_rows:
            lines.append("| Year | Files in folder | Files with insights | Top insight sample |")
            lines.append("|---|---:|---:|---|")
            for year_row in year_rows[:2]:
                if not isinstance(year_row, dict):
                    continue
                year_label = str(year_row.get("year", ""))
                file_count = int(year_row.get("file_count", 0))
                file_rows = year_row.get("files", [])
                files_with_highlights = 0
                top_note = ""
                if isinstance(file_rows, list):
                    for file_row in file_rows:
                        if isinstance(file_row, dict):
                            notes = file_row.get("highlights", []) or []
                            note_count = len(notes)
                            if note_count:
                                files_with_highlights += 1
                                if not top_note:
                                    top_note = str(notes[0]).strip()
                if not top_note:
                    top_note = "No concrete slide highlights extracted."
                lines.append(
                    f"| {year_label} | {file_count} | {files_with_highlights} | {top_note} |"
                )
        else:
            lines.append("- No year folders found for current/previous year in the configured Drive root.")

        highlights = seo_presentations.get("highlights", [])
        if isinstance(highlights, list) and highlights:
            lines.append("")
            lines.append("| Date | File | Highlight |")
            lines.append("|---|---|---|")
            for row in highlights[:10]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('date', '')} | {row.get('file', '')} | {row.get('note', '')} |"
                )
        errors = seo_presentations.get("errors", [])
        if isinstance(errors, list):
            for err in errors[:5]:
                lines.append(f"- Presentation warning: {err}")
    else:
        lines.append("- No SEO presentation data in this run.")

    historical_reports = (additional_context or {}).get("historical_reports", {})
    lines.append("")
    lines.append("### Previous report continuity (Drive)")
    if isinstance(historical_reports, dict) and historical_reports.get("enabled"):
        lines.append(f"- Source: {historical_reports.get('source', 'Google Drive')}")
        lines.append(f"- Available reports in archive: {int(historical_reports.get('available_reports', 0))}")

        recent_reports = historical_reports.get("recent_reports", [])
        if isinstance(recent_reports, list) and recent_reports:
            lines.append("| Report date | File | Highlight sample |")
            lines.append("|---|---|---|")
            for row in recent_reports[:3]:
                if not isinstance(row, dict):
                    continue
                highlights = row.get("highlights", [])
                highlight = ""
                if isinstance(highlights, list) and highlights:
                    highlight = str(highlights[0]).strip()
                if not highlight:
                    highlight = str(row.get("excerpt", "")).strip()[:180]
                lines.append(
                    f"| {row.get('date', '')} | {row.get('name', '')} | {highlight} |"
                )
        else:
            lines.append("- No previous reports found before current run date.")

        yoy_report = historical_reports.get("yoy_report", {})
        if isinstance(yoy_report, dict) and yoy_report.get("id"):
            yoy_highlights = yoy_report.get("highlights", [])
            yoy_note = ""
            if isinstance(yoy_highlights, list) and yoy_highlights:
                yoy_note = str(yoy_highlights[0]).strip()
            if not yoy_note:
                yoy_note = str(yoy_report.get("excerpt", "")).strip()[:180]
            lines.append(
                f"- YoY reference report: {yoy_report.get('date', '')} | {yoy_report.get('name', '')} | {yoy_note}"
            )
        else:
            lines.append("- YoY reference report not found within configured tolerance.")

        history_errors = historical_reports.get("errors", [])
        if isinstance(history_errors, list):
            for err in history_errors[:5]:
                lines.append(f"- History warning: {err}")
    else:
        lines.append("- No historical report context in this run.")

    lines.append("")
    lines.append("### What changed vs prior-report hypotheses")
    continuity_rows = _hypothesis_continuity_rows(
        hypotheses=hypotheses,
        additional_context=additional_context,
        limit=6,
    )
    lines.append("| Current hypothesis category | Continuity status | Evidence |")
    lines.append("|---|---|---|")
    if continuity_rows:
        for row in continuity_rows:
            lines.append(
                f"| {row.get('category', '')} | {row.get('status', '')} | {row.get('evidence', '')} |"
            )
    else:
        lines.append("| (no continuity rows) | No prior-report context | Historical weekly reports were not available in this run. |")

    tracker = (additional_context or {}).get("hypothesis_tracker", {})
    lines.append("")
    lines.append("### Hypothesis tracker")
    if isinstance(tracker, dict):
        counts = tracker.get("counts", {})
        if isinstance(counts, dict):
            lines.append(
                "- Status counts this run: "
                f"new={_fmt_int(counts.get('new', 0))}, "
                f"confirmed={_fmt_int(counts.get('confirmed', 0))}, "
                f"rejected={_fmt_int(counts.get('rejected', 0))}."
            )
        active = tracker.get("active", [])
        rejected = tracker.get("rejected", [])
        lines.append("| Hypothesis ID | Status | Category | Confidence |")
        lines.append("|---|---|---|---:|")
        has_rows = False
        if isinstance(active, list):
            for row in active[:8]:
                if not isinstance(row, dict):
                    continue
                has_rows = True
                lines.append(
                    f"| {row.get('id', '')} | {row.get('status', '')} | {row.get('category', '')} | "
                    f"{int(row.get('confidence', 0) or 0)}/100 |"
                )
        if isinstance(rejected, list):
            for row in rejected[:4]:
                if not isinstance(row, dict):
                    continue
                has_rows = True
                lines.append(
                    f"| {row.get('id', '')} | {row.get('status', '')} | {row.get('category', '')} | "
                    f"{int(row.get('confidence', 0) or 0)}/100 |"
                )
        if not has_rows:
            lines.append("| (no tracked hypotheses) | - | - | 0/100 |")
    else:
        lines.append("- Hypothesis tracker not available in this run.")

    status_log = (additional_context or {}).get("status_log", {})
    lines.append("")
    lines.append("### SEO status log (Sheets)")
    if isinstance(status_log, dict) and status_log.get("enabled"):
        lines.append(f"- Source: {status_log.get('source', 'Google Sheets')}")
        file_name = str(status_log.get("file_name", "")).strip()
        if file_name:
            lines.append(f"- File: {file_name}")
        selected_sheets = status_log.get("selected_sheets", [])
        if isinstance(selected_sheets, list) and selected_sheets:
            lines.append(f"- Year sheets used: {', '.join(str(item) for item in selected_sheets[:4])}")

        entries = status_log.get("entries", [])
        if isinstance(entries, list) and entries:
            lines.append("| Date | Sheet | Topic | Summary |")
            lines.append("|---|---|---|---|")
            for row in entries[:10]:
                if not isinstance(row, dict):
                    continue
                lines.append(
                    f"| {row.get('date', '')} | {row.get('sheet', '')} | {row.get('topic', '')} | {row.get('summary', '')} |"
                )
        else:
            lines.append("- No dated status entries found in selected sheets.")

        status_errors = status_log.get("errors", [])
        if isinstance(status_errors, list):
            for err in status_errors[:5]:
                lines.append(f"- Status warning: {err}")
    else:
        lines.append("- No status-log context in this run.")

    lines.append("")
    lines.append("## Winter-break regional context (PL)")
    lines.append(
        "- Regional GMV proxy is estimated as: `population x average salary` (normalized to 100% total)."
    )
    source_name = str(ferie_context.get("source", "")).strip()
    source_url = str(ferie_context.get("source_url", "")).strip()
    if source_name:
        if source_url:
            lines.append(f"- Source: {source_name} ({source_url})")
        else:
            lines.append(f"- Source: {source_name}")
    ferie_errors = ferie_context.get("errors")
    if isinstance(ferie_errors, list):
        for error in ferie_errors[:3]:
            lines.append(f"- Data warning: {error}")

    missing_years = ferie_context.get("missing_years", [])
    if isinstance(missing_years, list) and missing_years:
        missing_years_label = ", ".join(str(year) for year in missing_years)
        lines.append(f"- Missing winter-break calendar for years: {missing_years_label}.")

    lines.append("| Window | Days with winter-break overlap | Avg daily GMV proxy exposure | Peak daily GMV proxy exposure |")
    lines.append("|---|---:|---:|---:|")
    ferie_window_stats = ferie_context.get("window_stats", {})
    for key in ["current_28d", "previous_28d", "yoy_52w"]:
        row = ferie_window_stats.get(key, {}) if isinstance(ferie_window_stats, dict) else {}
        lines.append(
            f"| {windows[key].name} | {int(row.get('days_with_ferie', 0))} | {float(row.get('avg_daily_gmv_share', 0.0)) * 100:.2f}% | {float(row.get('peak_daily_gmv_share', 0.0)) * 100:.2f}% |"
        )

    lines.append("")
    lines.append("### Strongest regions by GMV proxy")
    lines.append("| Voivodeship | Population (m) | Avg salary PLN | GMV proxy share |")
    lines.append("|---|---:|---:|---:|")
    profiles_ranked = ferie_context.get("profiles_ranked", [])
    if isinstance(profiles_ranked, list) and profiles_ranked:
        for profile in profiles_ranked[:8]:
            if not isinstance(profile, dict):
                continue
            lines.append(
                f"| {profile.get('name', '')} | {float(profile.get('population_m', 0.0)):.2f} | {_fmt_int(profile.get('avg_salary_pln', 0.0))} | {float(profile.get('gmv_share', 0.0)) * 100:.2f}% |"
            )
    else:
        lines.append("| (no data) | 0.00 | 0 | 0.00% |")

    lines.append("")
    lines.append("### Winter-break YoY differences by region (current window vs 52W)")
    lines.append("| Voivodeship | Current break days | YoY break days | Delta days | Contribution to YoY exposure delta (pp) |")
    lines.append("|---|---:|---:|---:|---:|")
    yoy_comparison = ferie_context.get("yoy_comparison", {})
    yoy_rows = yoy_comparison.get("rows", []) if isinstance(yoy_comparison, dict) else []
    if isinstance(yoy_rows, list) and yoy_rows:
        for row in yoy_rows:
            if not isinstance(row, dict):
                continue
            lines.append(
                f"| {row.get('name', '')} | {int(row.get('current_days', 0))} | {int(row.get('yoy_days', 0))} | {int(row.get('delta_days', 0)):+d} | {float(row.get('contribution_pp', 0.0)):+.2f} |"
            )
        lines.append(
            f"- Net YoY delta in avg daily winter-break exposure: {float(yoy_comparison.get('avg_daily_delta_pp', 0.0)):+.2f} pp."
        )
    else:
        lines.append("| (no YoY differences) | 0 | 0 | +0 | +0.00 |")

    lines.append("")
    lines.append("## Integrated reasoning")
    lines.append("Reasoning combines KPI, GSC segments, winter breaks/holidays, weather, and external sources.")
    for reasoning_line in _build_integrated_reasoning(
        totals=totals,
        scope_results=scope_results,
        external_signals=external_signals,
        weather_summary=weather_summary,
        ferie_context=ferie_context,
        segment_diagnostics=segment_diagnostics,
        additional_context=additional_context,
        senuto_summary=senuto_summary,
        senuto_error=senuto_error,
    ):
        lines.append(reasoning_line)

    lines.append("")
    lines.append("## Root cause matrix")
    lines.append("| Root cause category | Confidence | Hypothesis | Key evidence | Owner |")
    lines.append("|---|---:|---|---|---|")
    for row in hypotheses[:8]:
        evidence = row.get("evidence", [])
        evidence_text = ""
        if isinstance(evidence, list):
            evidence_text = "; ".join(str(item) for item in evidence[:2])
        lines.append(
            f"| {row.get('category', '')} | {int(row.get('confidence', 0))}/100 | {row.get('thesis', '')} | {evidence_text} | {row.get('owner', '')} |"
        )

    lines.append("")
    lines.append("## Upcoming trends (next 60 days)")
    trends = _build_upcoming_trends(
        run_date=run_date,
        external_signals=external_signals,
        ferie_trends=upcoming_ferie_trends,
        additional_context=additional_context,
    )
    if trends:
        lines.append("| Date | Trend | Business/SEO impact | Suggested action |")
        lines.append("|---|---|---|---|")
        for trend_date, title, impact, action in trends:
            lines.append(f"| {trend_date} | {title} | {impact} | {action} |")
    else:
        lines.append("- No major upcoming trend detected in the next 60 days.")

    lines.append("")
    lines.append("## Query filtering")
    if query_filter_stats:
        lines.append("| Scope | Dropped current | Dropped previous | Dropped YoY |")
        lines.append("|---|---:|---:|---:|")
        for scope_name, stats in query_filter_stats.items():
            lines.append(
                f"| {scope_name} | {stats.get('current', 0)} | {stats.get('previous', 0)} | {stats.get('yoy', 0)} |"
            )
    else:
        lines.append("- Query filtering disabled or no query-based scopes.")

    lines.append("")
    lines.append("## Segment diagnostics")

    def _segment_table(
        title: str,
        rows: list[dict[str, float | str]] | None,
    ) -> None:
        lines.append(f"### {title}")
        lines.append("| Segment | Current clicks | Previous clicks | Delta vs prev | Delta vs YoY | CTR current | Pos current |")
        lines.append("|---|---:|---:|---:|---:|---:|---:|")
        if not rows:
            lines.append("| (no data) | 0 | 0 | +0 (0.00%) | +0 (0.00%) | 0.00% | 0.00 |")
            return
        for row in rows:
            lines.append(
                f"| {row.get('segment', '')} | {_fmt_int(row.get('current_clicks', 0.0))} | {_fmt_int(row.get('previous_clicks', 0.0))} | {_fmt_signed_int(row.get('delta_vs_previous', 0.0))} ({_signed_pct(float(row.get('delta_pct_vs_previous', 0.0)))}) | {_fmt_signed_int(row.get('delta_vs_yoy', 0.0))} ({_signed_pct(float(row.get('delta_pct_vs_yoy', 0.0)))}) | {_pct(float(row.get('ctr_current', 0.0)))} | {float(row.get('position_current', 0.0)):.2f} |"
            )

    _segment_table(
        "Brand vs non-brand",
        (segment_diagnostics or {}).get("brand_non_brand"),
    )
    lines.append("")
    _segment_table(
        "By device",
        (segment_diagnostics or {}).get("device"),
    )
    lines.append("")
    _segment_table(
        "By page name",
        (segment_diagnostics or {}).get("page_template"),
    )

    lines.append("")
    lines.append("## Further analysis flags")
    flags = _follow_up_flags(
        totals=totals,
        segment_diagnostics=segment_diagnostics,
        external_signals=external_signals,
        weather_summary=weather_summary,
    )
    if flags:
        for flag in flags:
            lines.append(f"- {flag}")
    else:
        lines.append("- No high-priority follow-up flags triggered by current thresholds.")

    lines.append("")
    lines.append("## Findings")
    finding_rows = _flatten_findings(scope_results)
    if finding_rows:
        for scope_name, finding in finding_rows:
            lines.append(
                f"- [{finding.severity.upper()}] ({scope_name}) {finding.title}: {finding.details} Action: {finding.recommendation}"
            )
    else:
        lines.append("- No critical findings from automated thresholds.")

    lines.append("")
    lines.append("## External signals (deduped, quality-scored)")
    quality_rows = _signal_quality_rows(external_signals, limit=12)
    lines.append("| Date | Signal theme | Source coverage | Severity | Source quality |")
    lines.append("|---|---|---:|---|---|")
    if quality_rows:
        for row in quality_rows:
            day = row.get("date")
            day_label = day.isoformat() if isinstance(day, date) else str(day or "")
            lines.append(
                f"| {day_label} | {row.get('title', '')} | {int(row.get('source_count', 0))} sources / {int(row.get('mentions', 0))} mentions | {str(row.get('severity', 'info')).upper()} | {str(row.get('quality', 'low')).upper()} |"
            )
    else:
        lines.append("| - | No external signals detected in this run | 0 | INFO | LOW |")

    lines.append("")
    lines.append("### News impact by GMV category")
    category_rows = _news_category_impact_rows(external_signals, limit=10)
    lines.append("| Category | Signals | Weighted impact | Latest date | Example headlines |")
    lines.append("|---|---:|---:|---|---|")
    if category_rows:
        for row in category_rows:
            latest_day = row.get("latest_day")
            latest_label = latest_day.isoformat() if isinstance(latest_day, date) else "-"
            examples = row.get("examples", [])
            if isinstance(examples, list):
                examples_text = "; ".join(str(item).strip() for item in examples[:2] if str(item).strip())
            else:
                examples_text = ""
            lines.append(
                f"| {row.get('category', '')} | {_fmt_int(row.get('signals', 0))} | {float(row.get('weighted_impact', 0.0)):.1f} | {latest_label} | {examples_text or '-'} |"
            )
    else:
        lines.append("| No categorized news signals | 0 | 0.0 | - | - |")

    lines.append("")
    lines.append("### Raw external signals (appendix log)")
    if external_signals:
        for signal in external_signals[:30]:
            url_suffix = f" ({signal.url})" if signal.url else ""
            lines.append(
                f"- [{signal.severity.upper()}] {signal.day.isoformat()} | {signal.source} | {signal.title}: {signal.details}{url_suffix}"
            )
    else:
        lines.append("- No raw external signals in this run.")

    lines.append("")
    lines.append("## Weather context")
    lines.append(
        f"- Avg temp current vs previous: {weather_summary.get('avg_temp_current_c', 0.0):.1f}C vs {weather_summary.get('avg_temp_previous_c', 0.0):.1f}C"
    )
    lines.append(
        f"- Total precipitation current vs previous: {weather_summary.get('precip_current_mm', 0.0):.1f}mm vs {weather_summary.get('precip_previous_mm', 0.0):.1f}mm"
    )
    forecast_start = str(weather_summary.get("forecast_start", "")).strip()
    forecast_end = str(weather_summary.get("forecast_end", "")).strip()
    if forecast_start and forecast_end:
        lines.append(
            "- Forecast next 7 days "
            f"({forecast_start} to {forecast_end}): avg temp {float(weather_summary.get('forecast_avg_temp_c', 0.0)):+.1f}C, "
            f"range {float(weather_summary.get('forecast_min_temp_c', 0.0)):+.1f}C to {float(weather_summary.get('forecast_max_temp_c', 0.0)):+.1f}C, "
            f"total precipitation {float(weather_summary.get('forecast_precip_mm', 0.0)):.1f}mm."
        )

    lines.append("")
    lines.append("## Senuto visibility")
    if senuto_summary:
        lines.append(
            f"- Average visibility current vs previous: {senuto_summary.get('avg_current', 0.0):.3f} vs {senuto_summary.get('avg_previous', 0.0):.3f} ({senuto_summary.get('avg_delta_pct', 0.0):+.2f}%)"
        )
        lines.append(
            f"- Latest visibility current vs previous: {senuto_summary.get('latest_current', 0.0):.3f} vs {senuto_summary.get('latest_previous', 0.0):.3f} ({senuto_summary.get('latest_delta_pct', 0.0):+.2f}%)"
        )
        if "avg_yoy" in senuto_summary:
            lines.append(
                f"- Average visibility YoY: {senuto_summary.get('avg_yoy', 0.0):.3f} ({senuto_summary.get('avg_delta_vs_yoy_pct', 0.0):+.2f}% vs current)"
            )
    elif senuto_error:
        lines.append(f"- Senuto unavailable in this run: {senuto_error}")
    else:
        lines.append("- Senuto not configured.")

    for scope_name, analysis in scope_results:
        lines.append("")
        lines.append(f"## Scope: {scope_name}")
        lines.extend(_scope_table("Top losers", analysis.top_losers))
        lines.append("")
        lines.extend(_scope_table("Top winners", analysis.top_winners))

    return "\n".join(lines)


def _scope_table(title: str, rows: list[KeyDelta]) -> list[str]:
    out = [f"### {title}"]
    out.append("| Key | Current clicks | Prev clicks | Delta vs prev | Delta vs YoY | CTR current | Pos current |")
    out.append("|---|---:|---:|---:|---:|---:|---:|")
    if not rows:
        out.append("| (no data) | 0 | 0 | 0 | 0 | 0.00% | 0.00 |")
        return out

    for row in rows:
        key = row.key.replace("|", "/")
        out.append(
            f"| {key[:80]} | {_fmt_int(row.current_clicks)} | {_fmt_int(row.previous_clicks)} | {_fmt_signed_int(row.click_delta_vs_previous)} ({_signed_pct(row.click_delta_pct_vs_previous)}) | {_fmt_signed_int(row.click_delta_vs_yoy)} ({_signed_pct(row.click_delta_pct_vs_yoy)}) | {_pct(row.current_ctr)} | {row.current_position:.2f} |"
        )
    return out


def _split_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_separator(line: str) -> bool:
    compact = line.replace("|", "").replace(":", "").replace("-", "").strip()
    return compact == ""


def _add_markdown_runs(paragraph, text: str) -> None:
    def _append_colored_run(value: str, *, bold: bool) -> None:
        last = 0
        for signed in SIGNED_VALUE_RE.finditer(value):
            start, end = signed.span()
            if start > last:
                base = paragraph.add_run(value[last:start])
                base.bold = bold
            token = signed.group(1)
            run = paragraph.add_run(token)
            run.bold = bold
            run.font.color.rgb = DARK_GREEN if token.startswith("+") else DARK_RED
            last = end
        if last < len(value):
            tail = paragraph.add_run(value[last:])
            tail.bold = bold

    last = 0
    for match in BOLD_MARKDOWN_RE.finditer(text):
        start, end = match.span()
        if start > last:
            _append_colored_run(text[last:start], bold=False)

        _append_colored_run(match.group(1), bold=True)
        last = end

    if last < len(text):
        _append_colored_run(text[last:], bold=False)


def _set_cell_markdown(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _add_markdown_runs(paragraph, text)


def _resolve_style_name(doc: Document, style_candidates: list[str], fallback: str = "Normal") -> str:
    for name in style_candidates:
        if not name:
            continue
        try:
            _ = doc.styles[name]
            return name
        except KeyError:
            continue
    return fallback


def _list_level(raw_line: str) -> int:
    leading = raw_line[: len(raw_line) - len(raw_line.lstrip(" \t"))]
    spaces = leading.count(" ") + (leading.count("\t") * 4)
    return min(4, max(0, spaces // 2))


def _apply_paragraph_spacing(paragraph, *, compact: bool = False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3 if compact else 6)
    pf.line_spacing = 1.15


def write_docx(path: Path, title: str, content: str) -> None:
    doc = Document()
    doc.core_properties.title = title
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
    except KeyError:
        pass

    lines = content.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped_line = raw_line.lstrip(" \t")

        if not raw_line:
            doc.add_paragraph("")
            index += 1
            continue

        if (
            raw_line.startswith("|")
            and index + 1 < len(lines)
            and lines[index + 1].startswith("|")
            and _is_markdown_separator(lines[index + 1])
        ):
            headers = _split_markdown_row(raw_line)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                _set_cell_markdown(table.rows[0].cells[col], header)

            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                row_cells = _split_markdown_row(lines[index])
                row = table.add_row().cells
                for col in range(len(headers)):
                    value = row_cells[col] if col < len(row_cells) else ""
                    _set_cell_markdown(row[col], value)
                index += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", raw_line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            paragraph = doc.add_heading("", level=level)
            _add_markdown_runs(paragraph, heading_match.group(2))
            _apply_paragraph_spacing(paragraph, compact=True)
        elif re.match(r"^\d+\.\s+", stripped_line):
            level = _list_level(raw_line)
            style_name = _resolve_style_name(
                doc,
                [f"List Number {level + 1}", "List Number"],
                fallback="Normal",
            )
            paragraph = doc.add_paragraph("", style=style_name)
            _add_markdown_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped_line))
            _apply_paragraph_spacing(paragraph, compact=True)
        elif stripped_line.startswith("- "):
            level = _list_level(raw_line)
            style_name = _resolve_style_name(
                doc,
                [f"List Bullet {level + 1}", "List Bullet"],
                fallback="Normal",
            )
            paragraph = doc.add_paragraph("", style=style_name)
            _add_markdown_runs(paragraph, stripped_line[2:])
            _apply_paragraph_spacing(paragraph, compact=True)
        else:
            paragraph = doc.add_paragraph("")
            _add_markdown_runs(paragraph, raw_line)
            _apply_paragraph_spacing(paragraph, compact=False)

        index += 1

    doc.save(path)
